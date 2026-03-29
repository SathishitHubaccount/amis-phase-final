"""
Generate Updated Demand Intelligence SQL Validation Document
=============================================================
This script creates a comprehensive Word document explaining:
- NEW: Database-stored historical data (not Python-generated)
- Which database tables store the data
- SQL queries to validate every step
- Complete data traceability
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create document
doc = Document()

# Configure default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================================
# TITLE PAGE
# ============================================================================
title = doc.add_heading('Demand Intelligence Module', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

subtitle = doc.add_paragraph('SQL Validation & Data Traceability Guide')
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle.runs[0].font.size = Pt(18)
subtitle.runs[0].font.color.rgb = RGBColor(68, 114, 196)

doc.add_paragraph()
info = doc.add_paragraph('AMIS - Autonomous Manufacturing Intelligence System')
info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
info.runs[0].font.italic = True

date = doc.add_paragraph('Updated: March 2026')
date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
date.runs[0].font.italic = True

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
doc.add_heading('Table of Contents', 1)
toc_items = [
    "1. Overview - Database Migration",
    "2. Database Architecture",
    "3. Historical Demand Data Table",
    "4. Market Context Data Table",
    "5. SQL Validation Queries",
    "6. Step-by-Step Data Flow with SQL",
    "7. Validating AI Tool Inputs",
    "8. Validating AI Tool Outputs",
    "9. Complete Traceability Workflow",
    "10. Troubleshooting",
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ============================================================================
# SECTION 1: OVERVIEW
# ============================================================================
doc.add_heading('1. Overview - Database Migration', 1)

doc.add_paragraph(
    'This document explains how historical demand data is now stored in the '
    'database (not generated on-the-fly), and how you can validate every step '
    'of the Demand Intelligence module using SQL queries.'
)

doc.add_heading('What Changed', 2)
doc.add_paragraph('BEFORE (Old Approach):', style='List Bullet').bold = True
bullets = [
    'Historical data generated on-the-fly by Python functions',
    'Data changed every time you ran the pipeline',
    'No way to validate with SQL',
    'No traceability or auditability',
    'Hard to debug issues',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph('AFTER (New Approach):', style='List Bullet').bold = True
bullets = [
    'Historical data stored in database tables',
    'Data is persistent and repeatable',
    'Full SQL validation capabilities',
    'Complete traceability and audit trail',
    'Easy to debug and verify results',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Key Benefits', 2)
p = doc.add_paragraph()
p.add_run('Repeatability: ').bold = True
p.add_run('Run the pipeline 100 times, get the same input data every time.')

p = doc.add_paragraph()
p.add_run('Traceability: ').bold = True
p.add_run('Every number the AI sees can be traced to a database row.')

p = doc.add_paragraph()
p.add_run('Validation: ').bold = True
p.add_run('Use SQL queries to verify AI inputs and outputs.')

p = doc.add_paragraph()
p.add_run('Professionalism: ').bold = True
p.add_run('Production-ready data architecture for hackathon demo.')

doc.add_page_break()

# ============================================================================
# SECTION 2: DATABASE ARCHITECTURE
# ============================================================================
doc.add_heading('2. Database Architecture', 1)

doc.add_paragraph(
    'The Demand Intelligence module uses 3 main database tables:'
)

# Table 1
doc.add_heading('Table 1: historical_demand_data', 2)
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('Stores 52 weeks of historical demand data for all products.')

p = doc.add_paragraph()
p.add_run('Records: ').bold = True
p.add_run('156 total (52 weeks × 3 products: PROD-A, PROD-B, PROD-C)')

p = doc.add_paragraph()
p.add_run('Data Source: ').bold = True
p.add_run('Populated by migrate_historical_data.py script')

# Table 2
doc.add_heading('Table 2: market_context_data', 2)
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('Stores macroeconomic and market condition data.')

p = doc.add_paragraph()
p.add_run('Records: ').bold = True
p.add_run('1 record with current market conditions')

p = doc.add_paragraph()
p.add_run('Data Source: ').bold = True
p.add_run('Populated by migrate_historical_data.py script')

# Table 3
doc.add_heading('Table 3: demand_forecasts', 2)
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('Stores AI-generated demand forecasts (OUTPUT).')

p = doc.add_paragraph()
p.add_run('Records: ').bold = True
p.add_run('Created dynamically by the AI pipeline')

p = doc.add_paragraph()
p.add_run('Data Source: ').bold = True
p.add_run('Generated by Demand Forecasting Agent')

doc.add_page_break()

# ============================================================================
# SECTION 3: HISTORICAL DEMAND DATA TABLE
# ============================================================================
doc.add_heading('3. Historical Demand Data Table', 1)

doc.add_heading('Table Schema', 2)
doc.add_paragraph('SQL Definition:')
code = doc.add_paragraph(
    'CREATE TABLE IF NOT EXISTS historical_demand_data (\n'
    '    id INTEGER PRIMARY KEY AUTOINCREMENT,\n'
    '    product_id TEXT NOT NULL,\n'
    '    week_start_date DATE NOT NULL,\n'
    '    week_number INTEGER NOT NULL,\n'
    '    year INTEGER NOT NULL,\n'
    '    demand_units INTEGER NOT NULL,\n'
    '    avg_price REAL,\n'
    '    promotions_active BOOLEAN DEFAULT 0,\n'
    '    competitor_price REAL,\n'
    '    is_anomaly BOOLEAN DEFAULT 0,\n'
    '    anomaly_reason TEXT,\n'
    '    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,\n'
    '    FOREIGN KEY (product_id) REFERENCES products(id),\n'
    '    UNIQUE(product_id, week_start_date)\n'
    ');',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

doc.add_heading('Column Descriptions', 2)

columns = [
    ('product_id', 'TEXT', 'Product identifier (PROD-A, PROD-B, PROD-C)'),
    ('week_start_date', 'DATE', 'Start date of the week (e.g., 2026-02-25)'),
    ('week_number', 'INTEGER', 'ISO week number (1-52)'),
    ('year', 'INTEGER', 'Year (e.g., 2026)'),
    ('demand_units', 'INTEGER', 'Actual demand in units for that week'),
    ('avg_price', 'REAL', 'Average selling price during that week'),
    ('promotions_active', 'BOOLEAN', '1 if promotions were running, 0 otherwise'),
    ('competitor_price', 'REAL', 'Competitor price during that week'),
    ('is_anomaly', 'BOOLEAN', '1 if week had demand anomaly, 0 otherwise'),
    ('anomaly_reason', 'TEXT', 'Explanation of anomaly (e.g., "Viral social media spike")'),
]

for col_name, col_type, col_desc in columns:
    p = doc.add_paragraph()
    p.add_run(f'{col_name} ').bold = True
    p.add_run(f'({col_type}): ')
    p.add_run(col_desc)

doc.add_heading('Sample Data', 2)
doc.add_paragraph('Query to view sample data:')
code = doc.add_paragraph(
    'SELECT week_start_date, demand_units, promotions_active, \n'
    '       is_anomaly, anomaly_reason\n'
    'FROM historical_demand_data\n'
    'WHERE product_id = \'PROD-A\'\n'
    'ORDER BY week_start_date DESC\n'
    'LIMIT 5;',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

doc.add_paragraph('Example Output:')
table = doc.add_table(rows=6, cols=5)
table.style = 'Light Grid Accent 1'

# Header
header_cells = table.rows[0].cells
header_cells[0].text = 'Date'
header_cells[1].text = 'Demand'
header_cells[2].text = 'Promo'
header_cells[3].text = 'Anomaly'
header_cells[4].text = 'Reason'

# Data rows
data_rows = [
    ('2026-02-25', '1789', 'No', 'No', '-'),
    ('2026-02-18', '1687', 'No', 'No', '-'),
    ('2026-02-11', '1686', 'No', 'No', '-'),
    ('2026-02-04', '2081', 'No', 'Yes', 'Viral social media spike'),
    ('2026-01-28', '1491', 'No', 'No', '-'),
]

for i, row_data in enumerate(data_rows, start=1):
    row_cells = table.rows[i].cells
    for j, cell_data in enumerate(row_data):
        row_cells[j].text = cell_data

doc.add_page_break()

# ============================================================================
# SECTION 4: MARKET CONTEXT DATA TABLE
# ============================================================================
doc.add_heading('4. Market Context Data Table', 1)

doc.add_heading('Table Schema', 2)
doc.add_paragraph('SQL Definition:')
code = doc.add_paragraph(
    'CREATE TABLE IF NOT EXISTS market_context_data (\n'
    '    id INTEGER PRIMARY KEY AUTOINCREMENT,\n'
    '    date_recorded DATE NOT NULL UNIQUE,\n'
    '    industry_growth_rate REAL,\n'
    '    economic_indicator TEXT,\n'
    '    competitor_activity TEXT,\n'
    '    raw_material_price_trend TEXT,\n'
    '    trade_show_date DATE,\n'
    '    major_client_contract_renewal_date DATE,\n'
    '    seasonal_pattern TEXT,\n'
    '    market_sentiment TEXT,\n'
    '    supply_chain_status TEXT,\n'
    '    notes TEXT,\n'
    '    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP\n'
    ');',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

doc.add_heading('Column Descriptions', 2)

columns = [
    ('industry_growth_rate', 'REAL', 'Industry growth rate in % (e.g., 7.2)'),
    ('economic_indicator', 'TEXT', 'Economic indicator description'),
    ('competitor_activity', 'TEXT', 'Recent competitor actions'),
    ('raw_material_price_trend', 'TEXT', 'Raw material price trends'),
    ('trade_show_date', 'DATE', 'Upcoming trade show date'),
    ('major_client_contract_renewal_date', 'DATE', 'Contract renewal date'),
    ('seasonal_pattern', 'TEXT', 'Seasonal demand patterns'),
    ('market_sentiment', 'TEXT', 'Overall market sentiment'),
    ('supply_chain_status', 'TEXT', 'Supply chain condition'),
]

for col_name, col_type, col_desc in columns:
    p = doc.add_paragraph()
    p.add_run(f'{col_name} ').bold = True
    p.add_run(f'({col_type}): ')
    p.add_run(col_desc)

doc.add_heading('Current Market Context', 2)
doc.add_paragraph('Query to view current context:')
code = doc.add_paragraph(
    'SELECT * FROM market_context_data\n'
    'ORDER BY date_recorded DESC\n'
    'LIMIT 1;',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

doc.add_paragraph('Current Values:')
context_data = [
    ('Industry Growth Rate', '7.2%'),
    ('Economic Indicator', 'Manufacturing PMI at 52.3 (expansion)'),
    ('Competitor Activity', 'Competitor launched new product targeting enterprise market'),
    ('Raw Material Price Trend', 'Raw material costs up 4% YoY due to supply chain constraints'),
    ('Seasonal Pattern', 'Q1 typically sees 15% higher demand (industrial buying cycle)'),
    ('Market Sentiment', 'Positive - Economic recovery driving capital investments'),
]

for label, value in context_data:
    p = doc.add_paragraph()
    p.add_run(f'{label}: ').bold = True
    p.add_run(value)

doc.add_page_break()

# ============================================================================
# SECTION 5: SQL VALIDATION QUERIES
# ============================================================================
doc.add_heading('5. SQL Validation Queries', 1)

doc.add_paragraph(
    'This section contains the most important SQL queries to validate the '
    'Demand Intelligence module.'
)

# Query 1
doc.add_heading('Query 1: View Historical Demand', 2)
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('View the raw historical data that the AI uses as input.')

doc.add_paragraph('SQL Query:')
code = doc.add_paragraph(
    'SELECT week_start_date, demand_units, promotions_active,\n'
    '       is_anomaly, anomaly_reason\n'
    'FROM historical_demand_data\n'
    'WHERE product_id = \'PROD-A\'\n'
    'ORDER BY week_start_date DESC\n'
    'LIMIT 12;',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('What to look for: ').bold = True
p.add_run('12 weeks of historical demand, demand values around 1,200-1,800 units, '
          'check for anomaly flags and promotion weeks.')

# Query 2
doc.add_heading('Query 2: Calculate Average Weekly Demand', 2)
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('Verify the "Average Weekly Demand" metric shown in the UI.')

doc.add_paragraph('SQL Query:')
code = doc.add_paragraph(
    'SELECT\n'
    '    ROUND(AVG(demand_units), 0) as avg_demand,\n'
    '    MIN(demand_units) as min_demand,\n'
    '    MAX(demand_units) as max_demand\n'
    'FROM historical_demand_data\n'
    'WHERE product_id = \'PROD-A\';',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Expected Result: ').bold = True
p.add_run('Average: ~1,344 units/week, Min: ~947 units, Max: ~2,081 units')

p = doc.add_paragraph()
p.add_run('Validation: ').bold = True
p.add_run('Compare this result to the "Average Weekly Demand" shown in the UI. '
          'They should match.')

# Query 3
doc.add_heading('Query 3: Find Demand Anomalies', 2)
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('Verify that the AI correctly identifies demand anomalies.')

doc.add_paragraph('SQL Query:')
code = doc.add_paragraph(
    'SELECT week_start_date, demand_units, anomaly_reason\n'
    'FROM historical_demand_data\n'
    'WHERE product_id = \'PROD-A\'\n'
    '  AND is_anomaly = 1\n'
    'ORDER BY week_start_date DESC;',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Expected Anomalies: ').bold = True

anomalies = [
    ('2026-02-04', '2081 units', 'Viral social media campaign spike'),
    ('2025-12-03', '1194 units', 'Competitor supply chain disruption'),
    ('2025-09-03', '1643 units', 'Black Friday promotional campaign'),
]

for date, demand, reason in anomalies:
    p = doc.add_paragraph(f'• {date}: {demand} - {reason}', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Validation: ').bold = True
p.add_run('Check if the AI mentions these anomalies in its analysis.')

# Query 4
doc.add_heading('Query 4: Week-over-Week Changes', 2)
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('Understand demand volatility and identify spikes.')

doc.add_paragraph('SQL Query:')
code = doc.add_paragraph(
    'WITH weekly_data AS (\n'
    '    SELECT week_start_date, demand_units,\n'
    '           LAG(demand_units) OVER (ORDER BY week_start_date) as prev\n'
    '    FROM historical_demand_data\n'
    '    WHERE product_id = \'PROD-A\'\n'
    ')\n'
    'SELECT week_start_date, demand_units, prev,\n'
    '       ROUND(((demand_units - prev) * 100.0 / prev), 2) as change_pct\n'
    'FROM weekly_data\n'
    'WHERE prev IS NOT NULL\n'
    'ORDER BY week_start_date DESC\n'
    'LIMIT 10;',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('What to look for: ').bold = True
p.add_run('Large spikes (>30%) indicate anomalies. Consistent positive changes = upward trend.')

# Query 5
doc.add_heading('Query 5: Compare Historical to Forecast', 2)
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('Validate that AI forecasts are reasonable based on historical data.')

doc.add_paragraph('SQL Query:')
code = doc.add_paragraph(
    '-- Historical average (AI input)\n'
    'SELECT \'Historical Avg\' as metric,\n'
    '       ROUND(AVG(demand_units), 0) as value\n'
    'FROM historical_demand_data\n'
    'WHERE product_id = \'PROD-A\'\n\n'
    'UNION ALL\n\n'
    '-- AI forecast (AI output)\n'
    'SELECT \'AI Forecast\' as metric, base_case as value\n'
    'FROM demand_forecasts\n'
    'WHERE product_id = \'PROD-A\'\n'
    'ORDER BY forecast_date DESC\n'
    'LIMIT 1;',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Validation: ').bold = True
p.add_run('AI forecast should be close to historical average. If different, check if AI '
          'adjusted for trends or anomalies.')

doc.add_page_break()

# ============================================================================
# SECTION 6: STEP-BY-STEP DATA FLOW
# ============================================================================
doc.add_heading('6. Step-by-Step Data Flow with SQL', 1)

doc.add_paragraph(
    'This section shows the complete data flow from database to UI, with SQL '
    'queries to validate each step.'
)

# Step 1
doc.add_heading('Step 1: User Clicks "Run Analysis"', 2)
doc.add_paragraph('User clicks the "Run Analysis" button in the UI.')
p = doc.add_paragraph()
p.add_run('What happens: ').bold = True
p.add_run('Frontend sends POST request to /api/pipeline/run')

# Step 2
doc.add_heading('Step 2: Backend Starts Demand Forecasting Agent', 2)
doc.add_paragraph('Backend invokes the Demand Forecasting Agent.')
p = doc.add_paragraph()
p.add_run('Agent decides: ').bold = True
p.add_run('Call Tool 1 (get_demand_data_summary) to gather context data.')

# Step 3
doc.add_heading('Step 3: Tool 1 Queries Database', 2)
doc.add_paragraph('Tool 1 executes database query to get historical demand.')

doc.add_paragraph('SQL Query (executed by Tool 1):')
code = doc.add_paragraph(
    'SELECT product_id, week_start_date as date, demand_units,\n'
    '       avg_price, promotions_active, competitor_price\n'
    'FROM historical_demand_data\n'
    'WHERE product_id = \'PROD-A\'\n'
    'ORDER BY week_start_date DESC\n'
    'LIMIT 12;',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Validation: ').bold = True
p.add_run('Run this query yourself to see the exact data the AI receives.')

p = doc.add_paragraph()
p.add_run('Result: ').bold = True
p.add_run('Tool 1 returns 12 weeks of historical demand data to the AI.')

# Step 4
doc.add_heading('Step 4: AI Analyzes Data and Calls Tool 2', 2)
doc.add_paragraph('AI analyzes the historical data and decides to call Tool 2.')
p = doc.add_paragraph()
p.add_run('What Tool 2 does: ').bold = True
p.add_run('Generates 3 demand scenarios (optimistic, base case, pessimistic) '
          'using statistical formulas based on historical data.')

p = doc.add_paragraph()
p.add_run('Input to Tool 2: ').bold = True
p.add_run('Historical demand from database (retrieved in Step 3)')

p = doc.add_paragraph()
p.add_run('Validation: ').bold = True
p.add_run('Tool 2 calculates average, trend, and volatility from the historical data. '
          'You can verify these calculations with SQL (see Query 2 and Query 4).')

# Step 5
doc.add_heading('Step 5: AI Calls Tool 3 for Trend Analysis', 2)
doc.add_paragraph('AI calls Tool 3 to perform statistical trend analysis.')
p = doc.add_paragraph()
p.add_run('What Tool 3 does: ').bold = True
p.add_run('Calculates linear regression slope, detects anomalies using z-scores, '
          'and computes correlations.')

p = doc.add_paragraph()
p.add_run('Input to Tool 3: ').bold = True
p.add_run('Same historical demand data from database')

p = doc.add_paragraph()
p.add_run('Validation: ').bold = True
p.add_run('Use Query 3 to verify that Tool 3 correctly identified the anomalies '
          'flagged in the database.')

# Step 6
doc.add_heading('Step 6: AI Generates Final Forecast', 2)
doc.add_paragraph('AI combines all tool results and generates final forecast.')
p = doc.add_paragraph()
p.add_run('AI reasoning: ').bold = True
p.add_run('Considers historical average, trend direction, anomalies, and market context.')

doc.add_paragraph('AI inserts forecast into database:')
code = doc.add_paragraph(
    'INSERT INTO demand_forecasts\n'
    '(product_id, week_number, forecast_date,\n'
    ' optimistic, base_case, pessimistic)\n'
    'VALUES (\'PROD-A\', 10, \'2026-03-04\',\n'
    '        1950, 1800, 1650);',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

# Step 7
doc.add_heading('Step 7: UI Displays Results', 2)
doc.add_paragraph('Frontend queries the database and displays results to user.')

doc.add_paragraph('SQL Query (executed by Frontend):')
code = doc.add_paragraph(
    'SELECT * FROM demand_forecasts\n'
    'WHERE product_id = \'PROD-A\'\n'
    'ORDER BY created_at DESC\n'
    'LIMIT 1;',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('Validation: ').bold = True
p.add_run('Run this query to see the exact forecast displayed in the UI.')

doc.add_page_break()

# ============================================================================
# SECTION 7: VALIDATING AI TOOL INPUTS
# ============================================================================
doc.add_heading('7. Validating AI Tool Inputs', 1)

doc.add_paragraph(
    'This section shows how to validate that the AI tools receive correct input data.'
)

doc.add_heading('Tool 1: get_demand_data_summary', 2)
p = doc.add_paragraph()
p.add_run('What it does: ').bold = True
p.add_run('Gathers 5 pieces of context data including historical demand.')

p = doc.add_paragraph()
p.add_run('Database query executed: ').bold = True

code = doc.add_paragraph(
    'SELECT * FROM historical_demand_data\n'
    'WHERE product_id = \'PROD-A\'\n'
    'ORDER BY week_start_date DESC LIMIT 12;',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('How to validate: ').bold = True
p.add_run('Run the query above and compare to the data shown in Tool 1 output. '
          'They should match exactly.')

doc.add_heading('Tool 2: simulate_demand_scenarios', 2)
p = doc.add_paragraph()
p.add_run('What it does: ').bold = True
p.add_run('Generates 3 probabilistic demand scenarios.')

p = doc.add_paragraph()
p.add_run('Input validation: ').bold = True

doc.add_paragraph('Step 1: Get historical average used by Tool 2:')
code = doc.add_paragraph(
    'SELECT AVG(demand_units) as historical_avg\n'
    'FROM historical_demand_data\n'
    'WHERE product_id = \'PROD-A\';',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

doc.add_paragraph('Step 2: Check Tool 2 output mentions this average.')

doc.add_heading('Tool 3: analyze_demand_trends', 2)
p = doc.add_paragraph()
p.add_run('What it does: ').bold = True
p.add_run('Performs statistical analysis (regression, anomaly detection).')

p = doc.add_paragraph()
p.add_run('Input validation: ').bold = True

doc.add_paragraph('Verify Tool 3 detected the anomalies:')
code = doc.add_paragraph(
    'SELECT week_start_date, demand_units, anomaly_reason\n'
    'FROM historical_demand_data\n'
    'WHERE product_id = \'PROD-A\' AND is_anomaly = 1;',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

doc.add_paragraph('Tool 3 should mention these anomalies in its output.')

doc.add_page_break()

# ============================================================================
# SECTION 8: VALIDATING AI TOOL OUTPUTS
# ============================================================================
doc.add_heading('8. Validating AI Tool Outputs', 1)

doc.add_paragraph(
    'This section shows how to validate the forecasts generated by the AI.'
)

doc.add_heading('Validate Base Case Forecast', 2)

doc.add_paragraph('Step 1: Get historical average:')
code = doc.add_paragraph(
    'SELECT AVG(demand_units) FROM historical_demand_data\n'
    'WHERE product_id = \'PROD-A\';',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)
doc.add_paragraph('Result: ~1,344 units/week')

doc.add_paragraph('Step 2: Get AI forecast:')
code = doc.add_paragraph(
    'SELECT base_case FROM demand_forecasts\n'
    'WHERE product_id = \'PROD-A\'\n'
    'ORDER BY created_at DESC LIMIT 1;',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

doc.add_paragraph('Step 3: Compare values:')
p = doc.add_paragraph('• If AI forecast ≈ historical average: AI is using base case logic')
p = doc.add_paragraph('• If AI forecast > historical average: AI detected upward trend')
p = doc.add_paragraph('• If AI forecast < historical average: AI detected downward trend')

doc.add_heading('Validate Anomaly Detection', 2)

doc.add_paragraph('Step 1: Check database for anomalies:')
code = doc.add_paragraph(
    'SELECT COUNT(*) FROM historical_demand_data\n'
    'WHERE product_id = \'PROD-A\' AND is_anomaly = 1;',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)
doc.add_paragraph('Result: 3 anomalies')

doc.add_paragraph('Step 2: Check if AI mentioned anomalies in analysis.')
doc.add_paragraph('The AI should reference the viral social media spike in its reasoning.')

doc.add_heading('Validate Trend Direction', 2)

doc.add_paragraph('SQL query to calculate trend:')
code = doc.add_paragraph(
    'WITH recent AS (\n'
    '    SELECT demand_units FROM historical_demand_data\n'
    '    WHERE product_id = \'PROD-A\'\n'
    '    ORDER BY week_start_date DESC LIMIT 4\n'
    '),\n'
    'all_data AS (\n'
    '    SELECT demand_units FROM historical_demand_data\n'
    '    WHERE product_id = \'PROD-A\'\n'
    '    ORDER BY week_start_date DESC LIMIT 12\n'
    ')\n'
    'SELECT\n'
    '    (SELECT AVG(demand_units) FROM recent) as recent_avg,\n'
    '    (SELECT AVG(demand_units) FROM all_data) as overall_avg;',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

doc.add_paragraph('If recent_avg > overall_avg: Upward trend')
doc.add_paragraph('If recent_avg < overall_avg: Downward trend')
doc.add_paragraph('Compare to the "Trend Direction" shown in UI.')

doc.add_page_break()

# ============================================================================
# SECTION 9: COMPLETE TRACEABILITY WORKFLOW
# ============================================================================
doc.add_heading('9. Complete Traceability Workflow', 1)

doc.add_paragraph(
    'This is the complete workflow to trace any number from the UI back to the database.'
)

doc.add_heading('Example: Tracing "Average Weekly Demand" Metric', 2)

doc.add_paragraph('Step 1: User sees "Average Weekly Demand: 1,344 units" in UI')

doc.add_paragraph('Step 2: Find where this is calculated in backend')
p = doc.add_paragraph('File: tools/forecasting.py')
p = doc.add_paragraph('Function: get_demand_data_summary()')
p = doc.add_paragraph('Code: Queries historical_demand_data table')

doc.add_paragraph('Step 3: Run the exact SQL query:')
code = doc.add_paragraph(
    'SELECT AVG(demand_units) FROM historical_demand_data\n'
    'WHERE product_id = \'PROD-A\';',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

doc.add_paragraph('Step 4: Verify result matches UI')
doc.add_paragraph('Result: 1343.58 units → Rounded to 1,344 in UI ✓')

doc.add_heading('Example: Tracing Anomaly Detection', 2)

doc.add_paragraph('Step 1: AI says "Detected demand spike of 2,081 units on 2026-02-04"')

doc.add_paragraph('Step 2: Verify this in database:')
code = doc.add_paragraph(
    'SELECT * FROM historical_demand_data\n'
    'WHERE product_id = \'PROD-A\'\n'
    '  AND week_start_date = \'2026-02-04\';',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'
code.runs[0].font.size = Pt(9)

doc.add_paragraph('Step 3: Check result:')
table = doc.add_table(rows=2, cols=4)
table.style = 'Light Grid Accent 1'

header = table.rows[0].cells
header[0].text = 'Date'
header[1].text = 'Demand'
header[2].text = 'Anomaly'
header[3].text = 'Reason'

data = table.rows[1].cells
data[0].text = '2026-02-04'
data[1].text = '2081'
data[2].text = 'Yes'
data[3].text = 'Viral social media spike'

doc.add_paragraph('Step 4: Confirm AI correctly identified this anomaly ✓')

doc.add_page_break()

# ============================================================================
# SECTION 10: TROUBLESHOOTING
# ============================================================================
doc.add_heading('10. Troubleshooting', 1)

doc.add_heading('Issue 1: No data returned from SQL query', 2)
p = doc.add_paragraph()
p.add_run('Problem: ').bold = True
p.add_run('Query returns empty result set.')

p = doc.add_paragraph()
p.add_run('Solution: ').bold = True
p.add_run('Run the migration script to populate data:')

code = doc.add_paragraph(
    'cd backend\n'
    'python migrate_historical_data.py',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'

doc.add_heading('Issue 2: AI forecast seems incorrect', 2)
p = doc.add_paragraph()
p.add_run('Problem: ').bold = True
p.add_run('AI forecast doesn\'t match expectations.')

p = doc.add_paragraph()
p.add_run('Solution: ').bold = True
p.add_run('Trace the calculation:')

steps = [
    'Step 1: Check historical average with Query 2',
    'Step 2: Check for anomalies with Query 3',
    'Step 3: Check trend direction with Query 4',
    'Step 4: Verify AI is considering all these factors',
]
for step in steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('Issue 3: UI shows different value than SQL', 2)
p = doc.add_paragraph()
p.add_run('Problem: ').bold = True
p.add_run('UI metric doesn\'t match SQL query result.')

p = doc.add_paragraph()
p.add_run('Solution: ').bold = True
p.add_run('Check for:')

items = [
    'Rounding differences (SQL: 1343.58, UI: 1,344)',
    'Time period mismatch (SQL: 52 weeks, UI: 12 weeks)',
    'Cached data in frontend (refresh browser)',
    'Multiple forecast records (check created_at timestamp)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Issue 4: How to reset all data', 2)
p = doc.add_paragraph()
p.add_run('Solution: ').bold = True
p.add_run('Delete database and re-run migration:')

code = doc.add_paragraph(
    'cd backend\n'
    'rm amis.db\n'
    'python main.py  # Recreates database\n'
    'python migrate_historical_data.py  # Populates data',
    style='No Spacing'
)
code.runs[0].font.name = 'Courier New'

doc.add_heading('Getting Help', 2)
doc.add_paragraph('If you encounter issues:')
items = [
    'Check DATABASE_MIGRATION_COMPLETE.md for detailed documentation',
    'Run test_database_queries.py to validate migration',
    'Check SQL_VALIDATION_QUERIES.md for quick reference queries',
    'Review backend logs for error messages',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================================
# SAVE DOCUMENT
# ============================================================================
output_path = 'DEMAND_INTELLIGENCE_SQL_VALIDATION_UPDATED.docx'
doc.save(output_path)
print(f'[SUCCESS] Word document created: {output_path}')
print(f'Document contains comprehensive SQL validation guide with database-stored data.')
