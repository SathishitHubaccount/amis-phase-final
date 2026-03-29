"""
Generate Complete Inventory Management Documentation
=====================================================
This script creates a comprehensive Word document that explains EVERYTHING
about the Inventory Management module - crystal clear and complete.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_title_page(doc):
    """Add professional title page"""
    title = doc.add_heading('Inventory Management Module', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.size = Pt(32)
        run.font.color.rgb = RGBColor(0, 102, 51)

    subtitle = doc.add_paragraph('Complete Technical Documentation')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(20)
    subtitle.runs[0].font.color.rgb = RGBColor(68, 114, 196)

    doc.add_paragraph()
    doc.add_paragraph()

    desc = doc.add_paragraph('Everything You Need to Know:')
    desc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    desc.runs[0].font.size = Pt(14)
    desc.runs[0].font.bold = True

    features = [
        'What the AI Agent Does',
        'All 6 Tools Explained in Detail',
        'Every Calculation with Examples',
        'All Database Tables & SQL Queries',
        'Complete Step-by-Step Flow',
        'ROP, EOQ, Safety Stock Formulas',
    ]

    for feature in features:
        p = doc.add_paragraph(f'• {feature}')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.runs[0].font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph('AMIS - Autonomous Manufacturing Intelligence System')
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    info.runs[0].font.italic = True

    doc.add_page_break()

def add_toc(doc):
    """Add table of contents"""
    doc.add_heading('Table of Contents', 1)

    sections = [
        ('1', 'What is Inventory Management?', 3),
        ('2', 'How the AI Agent Works', 4),
        ('3', 'Database Tables Used', 6),
        ('4', 'Tool 1: Get Inventory Status', 10),
        ('5', 'Tool 2: Calculate Reorder Point (ROP)', 14),
        ('6', 'Tool 3: Optimize Safety Stock', 19),
        ('7', 'Tool 4: Simulate Stockout Risk (Monte Carlo)', 24),
        ('8', 'Tool 5: Evaluate Holding Costs (EOQ)', 29),
        ('9', 'Tool 6: Generate Replenishment Plan', 34),
        ('10', 'Complete Execution Flow with Examples', 39),
        ('11', 'All SQL Queries Reference', 44),
        ('12', 'Validation & Verification', 47),
        ('13', 'Formulas & Calculations Cheat Sheet', 50),
    ]

    for num, title, page in sections:
        p = doc.add_paragraph()
        p.add_run(f'{num}. {title}').bold = True
        p.add_run(f' {"." * (65 - len(title))} {page}')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

def add_section_1(doc):
    """Section 1: What is Inventory Management?"""
    doc.add_heading('1. What is Inventory Management?', 1)

    doc.add_paragraph(
        'Inventory Management is an AI-powered module that optimizes inventory levels '
        'by calculating when to reorder, how much to order, and how much safety stock '
        'to maintain. It balances the cost of holding inventory against the risk of stockouts.'
    )

    doc.add_heading('What Problem Does It Solve?', 2)
    doc.add_paragraph(
        'Manufacturing companies face a critical tradeoff:'
    )

    p = doc.add_paragraph()
    p.add_run('Too Much Inventory: ').bold = True
    p.add_run('Money tied up, high holding costs, risk of obsolescence')

    p = doc.add_paragraph()
    p.add_run('Too Little Inventory: ').bold = True
    p.add_run('Stockouts, lost sales, unhappy customers, production delays')

    doc.add_paragraph(
        'The Inventory Management Agent uses advanced formulas (ROP, EOQ, Safety Stock) '
        'and Monte Carlo simulation to find the optimal balance.'
    )

    doc.add_heading('The 6 Key Questions It Answers', 2)
    questions = [
        ('When to order?', 'Reorder Point (ROP) calculation'),
        ('How much to order?', 'Economic Order Quantity (EOQ) calculation'),
        ('How much safety stock?', 'Safety Stock optimization'),
        ('What\'s the stockout risk?', 'Monte Carlo simulation'),
        ('What are the costs?', 'Holding cost vs stockout cost analysis'),
        ('What\'s the plan?', 'Week-by-week replenishment schedule'),
    ]

    for i, (question, answer) in enumerate(questions, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {question} ').bold = True
        p.add_run(f'→ {answer}')

    doc.add_heading('How Does It Work?', 2)
    steps = [
        ('Step 1', 'User clicks "Run Analysis" in the UI'),
        ('Step 2', 'AI Agent starts and calls Tool 1 to gather inventory status'),
        ('Step 3', 'AI calls Tool 2 to calculate Reorder Point (ROP)'),
        ('Step 4', 'AI calls Tool 3 to optimize Safety Stock levels'),
        ('Step 5', 'AI calls Tool 4 to simulate stockout risk (Monte Carlo)'),
        ('Step 6', 'AI calls Tool 5 to evaluate holding costs (EOQ)'),
        ('Step 7', 'AI calls Tool 6 to generate week-by-week replenishment plan'),
        ('Step 8', 'AI combines all results and generates recommendations'),
        ('Step 9', 'Recommendations displayed in UI'),
    ]

    for step_num, step_desc in steps:
        p = doc.add_paragraph()
        p.add_run(f'{step_num}: ').bold = True
        p.add_run(step_desc)

    doc.add_page_break()

def add_section_2(doc):
    """Section 2: How the AI Agent Works"""
    doc.add_heading('2. How the AI Agent Works', 1)

    doc.add_heading('The AI Agent (Claude 3.5 Sonnet)', 2)
    doc.add_paragraph(
        'The Inventory Management Agent is powered by Claude 3.5 Sonnet. It autonomously '
        'decides which tools to call and interprets their results to generate recommendations.'
    )

    doc.add_heading('Agent Configuration', 2)
    p = doc.add_paragraph()
    p.add_run('Framework: ').bold = True
    p.add_run('LangChain (Python library for AI agents)')

    p = doc.add_paragraph()
    p.add_run('Model: ').bold = True
    p.add_run('Claude 3.5 Sonnet')

    p = doc.add_paragraph()
    p.add_run('Tools Available: ').bold = True
    p.add_run('6 specialized inventory tools')

    p = doc.add_paragraph()
    p.add_run('Temperature: ').bold = True
    p.add_run('0.3 (low randomness for consistency)')

    doc.add_heading('Agent System Prompt', 2)
    doc.add_paragraph('The AI receives this instruction:')

    prompt = doc.add_paragraph(
        'You are an Inventory Management Agent for a manufacturing company. '
        'Your job is to optimize inventory levels by balancing holding costs and stockout risk. '
        'You have access to 6 tools: get inventory status, calculate reorder point, optimize safety stock, '
        'simulate stockout risk, evaluate holding costs, and generate replenishment plan. '
        'Use these tools to analyze the situation and provide actionable recommendations.',
        style='Intense Quote'
    )
    prompt.runs[0].font.italic = True

    doc.add_heading('How the Agent Decides Which Tools to Call', 2)
    doc.add_paragraph('The AI uses reasoning:')

    p = doc.add_paragraph()
    p.add_run('1. "I need to see current inventory status" ').bold = True
    p.add_run('→ Calls Tool 1 (get_inventory_status)')

    p = doc.add_paragraph()
    p.add_run('2. "I should calculate when to reorder" ').bold = True
    p.add_run('→ Calls Tool 2 (calculate_reorder_point)')

    p = doc.add_paragraph()
    p.add_run('3. "Is safety stock optimal?" ').bold = True
    p.add_run('→ Calls Tool 3 (optimize_safety_stock)')

    p = doc.add_paragraph()
    p.add_run('4. "What\'s the stockout risk?" ').bold = True
    p.add_run('→ Calls Tool 4 (simulate_stockout_risk)')

    p = doc.add_paragraph()
    p.add_run('5. "How much should we order?" ').bold = True
    p.add_run('→ Calls Tool 5 (evaluate_holding_costs / EOQ)')

    p = doc.add_paragraph()
    p.add_run('6. "What\'s the action plan?" ').bold = True
    p.add_run('→ Calls Tool 6 (generate_replenishment_plan)')

    doc.add_heading('Agent Output', 2)
    doc.add_paragraph('The AI generates:')
    doc.add_paragraph('• Current inventory health assessment', style='List Bullet')
    doc.add_paragraph('• Reorder recommendations (when and how much)', style='List Bullet')
    doc.add_paragraph('• Safety stock adjustments', style='List Bullet')
    doc.add_paragraph('• Stockout risk probability', style='List Bullet')
    doc.add_paragraph('• Cost analysis (holding vs stockout)', style='List Bullet')
    doc.add_paragraph('• 4-week replenishment plan', style='List Bullet')

    doc.add_page_break()

def add_section_3(doc):
    """Section 3: Database Tables Used"""
    doc.add_heading('3. Database Tables Used', 1)

    doc.add_paragraph(
        'The Inventory Management module uses 1 primary database table. '
        'Note: Currently most data comes from sample_data.py (Python functions), '
        'but the table structure exists for future migration.'
    )

    # TABLE 1: inventory
    doc.add_heading('Table 1: inventory', 2)

    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Stores current inventory levels, reorder points, and safety stock settings.')

    p = doc.add_paragraph()
    p.add_run('Total Records: ').bold = True
    p.add_run('1 record per product (3 products)')

    p = doc.add_paragraph()
    p.add_run('Updated: ').bold = True
    p.add_run('Real-time updates as inventory changes')

    doc.add_heading('Table Schema', 3)

    table = doc.add_table(rows=9, cols=3)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Column Name'
    header[1].text = 'Data Type'
    header[2].text = 'Description'

    schema_data = [
        ('product_id', 'TEXT', 'Product identifier (primary key)'),
        ('current_stock', 'INTEGER', 'Current inventory level in units'),
        ('safety_stock', 'INTEGER', 'Safety stock target in units'),
        ('reorder_point', 'INTEGER', 'Reorder point in units (when to order)'),
        ('avg_daily_usage', 'REAL', 'Average daily demand in units'),
        ('lead_time', 'INTEGER', 'Supplier lead time in days'),
        ('stockout_risk', 'REAL', 'Current stockout risk probability'),
        ('last_updated', 'TIMESTAMP', 'Last update timestamp'),
    ]

    for i, (col, dtype, desc) in enumerate(schema_data, start=1):
        row = table.rows[i].cells
        row[0].text = col
        row[1].text = dtype
        row[2].text = desc

    doc.add_heading('SQL to View Data', 3)
    code = doc.add_paragraph(
        'SELECT * FROM inventory\n'
        'WHERE product_id = \'PROD-A\';',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(10)

    doc.add_heading('Example Data', 3)

    table = doc.add_table(rows=2, cols=6)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Product'
    header[1].text = 'Stock'
    header[2].text = 'Safety Stock'
    header[3].text = 'ROP'
    header[4].text = 'Daily Usage'
    header[5].text = 'Lead Time'

    data = table.rows[1].cells
    data[0].text = 'PROD-A'
    data[1].text = '1,850'
    data[2].text = '300'
    data[3].text = '710'
    data[4].text = '142'
    data[5].text = '5 days'

    # TABLE 2: inventory_history
    doc.add_heading('Table 2: inventory_history', 2)

    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Tracks historical inventory levels over time for trend analysis.')

    p = doc.add_paragraph()
    p.add_run('Total Records: ').bold = True
    p.add_run('Daily snapshots (dynamic)')

    doc.add_heading('Table Schema', 3)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Column Name'
    header[1].text = 'Data Type'
    header[2].text = 'Description'

    schema_data = [
        ('id', 'INTEGER', 'Auto-increment primary key'),
        ('product_id', 'TEXT', 'Product identifier'),
        ('date', 'DATE', 'Snapshot date'),
        ('stock_level', 'INTEGER', 'Inventory level on that date'),
        ('stock_value', 'REAL', 'Dollar value of inventory'),
    ]

    for i, (col, dtype, desc) in enumerate(schema_data, start=1):
        row = table.rows[i].cells
        row[0].text = col
        row[1].text = dtype
        row[2].text = desc

    doc.add_heading('SQL to View Data', 3)
    code = doc.add_paragraph(
        'SELECT * FROM inventory_history\n'
        'WHERE product_id = \'PROD-A\'\n'
        'ORDER BY date DESC\n'
        'LIMIT 30;  -- Last 30 days',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(10)

    # Data Sources Note
    doc.add_heading('Note: Current Data Sources', 2)
    p = doc.add_paragraph()
    p.add_run('IMPORTANT: ').bold = True
    p.add_run(
        'Currently, most inventory data comes from Python functions in sample_data.py, not the database. '
        'This includes:'
    )

    sources = [
        'Warehouse details (zones, capacity)',
        'Supplier performance (lead times, on-time delivery)',
        'Stockout history',
        'Reorder history',
        'Incoming orders',
    ]

    for source in sources:
        doc.add_paragraph(f'• {source}', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Future Enhancement: ').bold = True
    p.add_run(
        'These can be migrated to database tables (similar to how we migrated '
        'historical_demand_data for the Demand Intelligence module).'
    )

    doc.add_page_break()

# Continue with remaining sections...
def generate_tool_sections(doc):
    """Generate sections 4-9 for all 6 tools"""

    # TOOL 1
    doc.add_heading('4. Tool 1: Get Inventory Status', 1)

    doc.add_heading('What This Tool Does', 2)
    doc.add_paragraph(
        'Tool 1 gathers a comprehensive snapshot of the inventory situation. '
        'It queries 6 different data sources and calculates 8 health indicators.'
    )

    doc.add_heading('Tool Definition', 2)
    p = doc.add_paragraph()
    p.add_run('Function Name: ').bold = True
    p.add_run('get_inventory_status(product_id)')

    p = doc.add_paragraph()
    p.add_run('Input: ').bold = True
    p.add_run('product_id (string, e.g., "PROD-A")')

    p = doc.add_paragraph()
    p.add_run('Output: ').bold = True
    p.add_run('JSON object with 6 data sections + 8 health indicators')

    doc.add_heading('The 6 Data Sources', 2)

    # Data Source 1
    p = doc.add_paragraph()
    p.add_run('1. Current Inventory (from sample_data.py)').bold = True

    doc.add_paragraph('Python Function Call:')
    code = doc.add_paragraph(
        'inventory = get_current_inventory(product_id)',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Returns:')
    returns = [
        'current_stock: 1,850 units',
        'safety_stock: 300 units',
        'warehouse_capacity: 5,000 units',
        'avg_daily_consumption: 142 units/day',
        'days_of_supply: 13 days',
        'incoming_orders: [Order 1, Order 2]',
        'unit_holding_cost: $2.30/unit/year',
        'unit_stockout_cost: $45/unit',
    ]
    for item in returns:
        doc.add_paragraph(f'• {item}', style='List Bullet')

    # Data Source 2
    p = doc.add_paragraph()
    p.add_run('2. Warehouse Details (from sample_data.py)').bold = True

    doc.add_paragraph('Python Function Call:')
    code = doc.add_paragraph(
        'warehouse = get_warehouse_details(product_id)',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Returns:')
    returns = [
        'Total capacity: 5,000 units',
        'Zone A (Primary): 2,500 units capacity, 1,200 units current',
        'Zone B (Fast-Pick): 1,000 units capacity, 450 units current',
        'Zone C (Overflow): 1,500 units capacity, 200 units current',
        'Total utilization: 37%',
    ]
    for item in returns:
        doc.add_paragraph(f'• {item}', style='List Bullet')

    # Data Source 3
    p = doc.add_paragraph()
    p.add_run('3. Supplier Performance (from sample_data.py)').bold = True

    doc.add_paragraph('Python Function Call:')
    code = doc.add_paragraph(
        'suppliers = get_supplier_performance(product_id)',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Returns data for 2 suppliers:')

    table = doc.add_table(rows=3, cols=5)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Supplier'
    header[1].text = 'Lead Time'
    header[2].text = 'On-Time %'
    header[3].text = 'Unit Cost'
    header[4].text = 'Quality Reject %'

    row1 = table.rows[1].cells
    row1[0].text = 'Supplier A'
    row1[1].text = '4.2 days (±0.8)'
    row1[2].text = '92.5%'
    row1[3].text = '$52.00'
    row1[4].text = '1.2%'

    row2 = table.rows[2].cells
    row2[0].text = 'Supplier B'
    row2[1].text = '6.8 days (±1.5)'
    row2[2].text = '85.0%'
    row2[3].text = '$49.50'
    row2[4].text = '2.1%'

    # Data Sources 4-6
    p = doc.add_paragraph()
    p.add_run('4. Stockout History').bold = True
    doc.add_paragraph('Returns: Last 3 stockout events with revenue lost, duration, root cause')

    p = doc.add_paragraph()
    p.add_run('5. Reorder History').bold = True
    doc.add_paragraph('Returns: Last 8 reorder events with quantities, dates, suppliers')

    p = doc.add_paragraph()
    p.add_run('6. Product Info').bold = True
    doc.add_paragraph('Returns: Product name, unit cost ($52), unit price ($89.50), margin (41.9%)')

    doc.add_heading('The 8 Health Indicators Calculated', 2)

    doc.add_paragraph('Tool 1 calculates these metrics:')

    table = doc.add_table(rows=9, cols=3)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Indicator'
    header[1].text = 'Formula'
    header[2].text = 'Example Value'

    indicators = [
        ('Days of Supply', 'current_stock / avg_daily_consumption', '13.0 days'),
        ('Above Safety Stock?', 'current_stock > safety_stock', 'Yes (1,850 > 300)'),
        ('Buffer Above Safety', 'current_stock - safety_stock', '1,550 units'),
        ('Incoming Pipeline', 'SUM(incoming_orders.quantity)', '800 units'),
        ('Effective Days Supply', '(stock + incoming) / daily_consumption', '18.7 days'),
        ('Warehouse Utilization', 'current_stock / capacity * 100', '37%'),
        ('Stockout Events (12mo)', 'COUNT(stockout_history)', '3 events'),
        ('Total Revenue Lost', 'SUM(stockout_history.revenue_lost)', '$111,875'),
    ]

    for i, (indicator, formula, value) in enumerate(indicators, start=1):
        row = table.rows[i].cells
        row[0].text = indicator
        row[1].text = formula
        row[2].text = value

    doc.add_heading('Example Calculation: Effective Days of Supply', 2)

    calc = doc.add_paragraph(
        'Current Stock: 1,850 units\n'
        'Incoming Orders: 500 + 300 = 800 units\n'
        'Avg Daily Consumption: 142 units/day\n\n'
        'Effective Days of Supply = (1,850 + 800) / 142\n'
        '                         = 2,650 / 142\n'
        '                         = 18.7 days',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    doc.add_paragraph(
        'Interpretation: We have 18.7 days of supply before stockout (if no new orders placed).'
    )

    doc.add_page_break()

    # TOOL 2
    doc.add_heading('5. Tool 2: Calculate Reorder Point (ROP)', 1)

    doc.add_heading('What This Tool Does', 2)
    doc.add_paragraph(
        'Tool 2 calculates the Reorder Point - the inventory level at which you should '
        'place a new order. It uses the industry-standard ROP formula that accounts for '
        'both demand variability and lead time variability.'
    )

    doc.add_heading('Tool Definition', 2)
    p = doc.add_paragraph()
    p.add_run('Function Name: ').bold = True
    p.add_run('calculate_reorder_point(...)')

    p = doc.add_paragraph()
    p.add_run('Inputs: ').bold = True

    inputs = [
        'expected_daily_demand: 142 units/day (from demand forecast)',
        'demand_std_dev_daily: 25 units (demand volatility)',
        'lead_time_days: 5 days (supplier lead time)',
        'lead_time_std_dev_days: 1 day (lead time variability)',
        'service_level_pct: 95% (target service level)',
    ]
    for item in inputs:
        doc.add_paragraph(f'• {item}', style='List Bullet')

    doc.add_heading('The ROP Formula', 2)

    doc.add_paragraph('Step 1: Calculate Average Demand During Lead Time')
    formula = doc.add_paragraph(
        'avg_demand_during_LT = daily_demand × lead_time',
        style='No Spacing'
    )
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True

    doc.add_paragraph('Example:')
    calc = doc.add_paragraph(
        'avg_demand_during_LT = 142 × 5\n'
        '                     = 710 units',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    doc.add_paragraph('Step 2: Calculate Safety Stock')
    formula = doc.add_paragraph(
        'SS = Z × sqrt(LT × σ_demand² + demand² × σ_LT²)\n\n'
        'where:\n'
        '  Z = Z-score for service level (1.645 for 95%)\n'
        '  LT = lead time in days\n'
        '  σ_demand = demand standard deviation\n'
        '  σ_LT = lead time standard deviation',
        style='No Spacing'
    )
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True

    doc.add_heading('Z-Score Table (Service Level)', 3)

    table = doc.add_table(rows=9, cols=2)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Service Level'
    header[1].text = 'Z-Score'

    z_data = [
        ('85%', '1.04'),
        ('90%', '1.28'),
        ('92%', '1.41'),
        ('95%', '1.645'),
        ('97%', '1.88'),
        ('98%', '2.05'),
        ('99%', '2.326'),
        ('99.9%', '3.09'),
    ]

    for i, (level, z) in enumerate(z_data, start=1):
        row = table.rows[i].cells
        row[0].text = level
        row[1].text = z

    doc.add_paragraph('Example Safety Stock Calculation:')
    calc = doc.add_paragraph(
        'Z = 1.645  (for 95% service level)\n'
        'LT = 5 days\n'
        'σ_demand = 25 units\n'
        'demand = 142 units\n'
        'σ_LT = 1 day\n\n'
        'Demand variance = 25² = 625\n'
        'LT variance = 1² = 1\n\n'
        'SS = 1.645 × sqrt(5 × 625 + 142² × 1)\n'
        '   = 1.645 × sqrt(3,125 + 20,164)\n'
        '   = 1.645 × sqrt(23,289)\n'
        '   = 1.645 × 152.6\n'
        '   = 251 units',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    doc.add_paragraph('Step 3: Calculate Reorder Point')
    formula = doc.add_paragraph(
        'ROP = avg_demand_during_LT + safety_stock',
        style='No Spacing'
    )
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True

    doc.add_paragraph('Example:')
    calc = doc.add_paragraph(
        'ROP = 710 + 251\n'
        '    = 961 units',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    doc.add_heading('Interpretation', 2)
    p = doc.add_paragraph()
    p.add_run('Result: ROP = 961 units').bold = True

    doc.add_paragraph('What this means:')
    doc.add_paragraph(
        'When inventory drops to 961 units, place a new order. At 95% service level, '
        'you will have enough stock to cover demand during the 5-day lead time.'
    )

    p = doc.add_paragraph()
    p.add_run('Current Status Check:').bold = True
    doc.add_paragraph('Current stock: 1,850 units')
    doc.add_paragraph('ROP: 961 units')
    doc.add_paragraph('Result: Stock is ABOVE reorder point - no order needed yet')
    doc.add_paragraph('Days until ROP: (1850 - 961) / 142 = 6.3 days')

    doc.add_page_break()

    # TOOL 3
    doc.add_heading('6. Tool 3: Optimize Safety Stock', 1)

    doc.add_heading('What This Tool Does', 2)
    doc.add_paragraph(
        'Tool 3 finds the optimal safety stock level by testing multiple service levels '
        '(85% to 99%) and calculating the total cost (holding cost + stockout cost) for each. '
        'It finds the service level that minimizes total cost.'
    )

    doc.add_heading('Tool Definition', 2)
    p = doc.add_paragraph()
    p.add_run('Function Name: ').bold = True
    p.add_run('optimize_safety_stock(...)')

    p = doc.add_paragraph()
    p.add_run('Inputs: ').bold = True

    inputs = [
        'expected_daily_demand: 142 units/day',
        'demand_std_dev_daily: 25 units',
        'lead_time_days: 5 days',
        'lead_time_std_dev_days: 1 day',
        'holding_cost_per_unit_per_year: $2.30',
        'stockout_cost_per_unit: $45.00',
    ]
    for item in inputs:
        doc.add_paragraph(f'• {item}', style='List Bullet')

    doc.add_heading('The Optimization Process', 2)

    doc.add_paragraph('Step 1: For Each Service Level (85% to 99%)')
    doc.add_paragraph('Calculate safety stock using same formula as Tool 2:')

    formula = doc.add_paragraph(
        'SS = Z × sqrt(LT × σ_demand² + demand² × σ_LT²)',
        style='No Spacing'
    )
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True

    doc.add_paragraph('Step 2: Calculate Annual Holding Cost')
    formula = doc.add_paragraph(
        'annual_holding_cost = safety_stock × holding_cost_per_unit_per_year',
        style='No Spacing'
    )
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True

    doc.add_paragraph('Step 3: Calculate Expected Annual Stockout Cost')
    formula = doc.add_paragraph(
        'stockout_prob = 1 - (service_level / 100)\n'
        'cycles_per_year = annual_demand / (daily_demand × LT × 2)\n'
        'shortage_per_cycle = σ_demand × sqrt(LT) × stockout_prob × 3\n'
        'annual_stockout_cost = shortage_per_cycle × cycles_per_year × stockout_cost',
        style='No Spacing'
    )
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True
    formula.runs[0].font.size = Pt(8)

    doc.add_paragraph('Step 4: Calculate Total Cost')
    formula = doc.add_paragraph(
        'total_cost = annual_holding_cost + annual_stockout_cost',
        style='No Spacing'
    )
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True

    doc.add_heading('Example Calculation for 95% Service Level', 2)

    calc = doc.add_paragraph(
        '1. Safety Stock:\n'
        '   SS = 1.645 × sqrt(5 × 625 + 142² × 1)\n'
        '      = 251 units\n\n'
        '2. Annual Holding Cost:\n'
        '   holding = 251 × $2.30\n'
        '           = $577.30\n\n'
        '3. Annual Stockout Cost:\n'
        '   stockout_prob = 1 - 0.95 = 0.05\n'
        '   annual_demand = 142 × 365 = 51,830\n'
        '   cycles_per_year = 51,830 / (142 × 5 × 2) = 36.5\n'
        '   shortage_per_cycle = 25 × sqrt(5) × 0.05 × 3 = 8.38\n'
        '   stockout_cost = 8.38 × 36.5 × $45 = $13,765\n\n'
        '4. Total Cost:\n'
        '   total = $577.30 + $13,765 = $14,342.30',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    doc.add_heading('Complete Analysis Table', 2)

    table = doc.add_table(rows=8, cols=5)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Service Level'
    header[1].text = 'Safety Stock'
    header[2].text = 'Holding Cost'
    header[3].text = 'Stockout Cost'
    header[4].text = 'Total Cost'

    analysis_data = [
        ('85%', '168', '$386', '$27,450', '$27,836'),
        ('90%', '207', '$476', '$20,587', '$21,063'),
        ('92%', '228', '$524', '$16,470', '$16,994'),
        ('95%', '251', '$577', '$13,765', '$14,342'),
        ('97%', '304', '$699', '$8,235', '$8,934'),
        ('98%', '331', '$761', '$5,490', '$6,251'),
        ('99%', '376', '$865', '$2,745', '$3,610'),
    ]

    for i, (level, ss, hold, stock, total) in enumerate(analysis_data, start=1):
        row = table.rows[i].cells
        row[0].text = level
        row[1].text = ss
        row[2].text = hold
        row[3].text = stock
        row[4].text = total

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('OPTIMAL: 99% service level with safety stock of 376 units').bold = True
    p.add_run('\nTotal annual cost: $3,610 (lowest total cost)')

    doc.add_heading('Current vs Optimal Comparison', 2)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Metric'
    header[1].text = 'Current'
    header[2].text = 'Optimal'

    comparison = [
        ('Safety Stock', '300 units', '376 units'),
        ('Service Level', '~96%', '99%'),
        ('Annual Cost', '$690', '$3,610'),
    ]

    for i, (metric, current, optimal) in enumerate(comparison, start=1):
        row = table.rows[i].cells
        row[0].text = metric
        row[1].text = current
        row[2].text = optimal

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Recommendation: ').bold = True
    p.add_run('Increase safety stock from 300 to 376 units to achieve 99% service level.')

    doc.add_page_break()

    # TOOL 4
    doc.add_heading('7. Tool 4: Simulate Stockout Risk (Monte Carlo)', 1)

    doc.add_heading('What This Tool Does', 2)
    doc.add_paragraph(
        'Tool 4 runs Monte Carlo simulation to calculate the probability of stockout '
        'over the next 14 days. It simulates 1,000 different scenarios of daily demand '
        'and incoming orders to quantify stockout risk.'
    )

    doc.add_heading('Tool Definition', 2)
    p = doc.add_paragraph()
    p.add_run('Function Name: ').bold = True
    p.add_run('simulate_stockout_risk(...)')

    p = doc.add_paragraph()
    p.add_run('Inputs: ').bold = True

    inputs = [
        'current_stock: 1,850 units',
        'expected_daily_demand: 142 units/day',
        'demand_std_dev_daily: 25 units',
        'horizon_days: 14 days',
        'simulations: 1,000 runs',
    ]
    for item in inputs:
        doc.add_paragraph(f'• {item}', style='List Bullet')

    doc.add_heading('Monte Carlo Simulation Algorithm', 2)

    doc.add_paragraph('For each of 1,000 simulations:')

    steps = [
        '1. Start with current stock (1,850 units)',
        '2. For each day (1 to 14):',
        '   a. Generate random daily demand from normal distribution',
        '      demand ~ N(142, 25)',
        '   b. Subtract demand from stock',
        '   c. Check if any incoming orders arrive (with ±1 day jitter)',
        '   d. Add incoming orders to stock',
        '   e. Record stock level',
        '   f. If stock <= 0, record stockout',
        '3. Repeat for all days',
        '4. Track stockout probability for each day',
    ]

    for step in steps:
        doc.add_paragraph(step)

    doc.add_heading('Example Simulation Run', 2)

    table = doc.add_table(rows=8, cols=6)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Day'
    header[1].text = 'Start Stock'
    header[2].text = 'Demand'
    header[3].text = 'Incoming'
    header[4].text = 'End Stock'
    header[5].text = 'Stockout?'

    sim_data = [
        ('1', '1,850', '138', '0', '1,712', 'No'),
        ('2', '1,712', '155', '0', '1,557', 'No'),
        ('3', '1,557', '142', '0', '1,415', 'No'),
        ('4', '1,415', '149', '500', '1,766', 'No'),
        ('5', '1,766', '168', '0', '1,598', 'No'),
        ('6', '1,598', '132', '0', '1,466', 'No'),
        ('7', '1,466', '151', '300', '1,615', 'No'),
    ]

    for i, (day, start, demand, incoming, end, stockout) in enumerate(sim_data, start=1):
        row = table.rows[i].cells
        row[0].text = day
        row[1].text = start
        row[2].text = demand
        row[3].text = incoming
        row[4].text = end
        row[5].text = stockout

    doc.add_heading('Simulation Results (1,000 runs)', 2)

    table = doc.add_table(rows=8, cols=5)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Day'
    header[1].text = 'Avg Stock'
    header[2].text = 'Min Stock'
    header[3].text = 'Max Stock'
    header[4].text = 'Stockout %'

    results_data = [
        ('1', '1,708', '1,520', '1,890', '0.0%'),
        ('2', '1,566', '1,210', '1,750', '0.0%'),
        ('3', '1,424', '890', '1,610', '0.1%'),
        ('4', '1,782', '1,250', '2,090', '0.0%'),
        ('5', '1,640', '980', '1,950', '0.1%'),
        ('10', '1,126', '120', '1,680', '2.3%'),
        ('14', '784', '-180', '1,420', '8.7%'),
    ]

    for i, (day, avg, min_s, max_s, stockout) in enumerate(results_data, start=1):
        row = table.rows[i].cells
        row[0].text = day
        row[1].text = avg
        row[2].text = min_s
        row[3].text = max_s
        row[4].text = stockout

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Overall Stockout Probability: 8.7%').bold = True
    p.add_run(' (87 out of 1,000 simulations had at least one stockout)')

    p = doc.add_paragraph()
    p.add_run('Critical Days: ').bold = True
    p.add_run('Day 10 (2.3% risk), Day 14 (8.7% risk)')

    p = doc.add_paragraph()
    p.add_run('Worst Case: ').bold = True
    p.add_run('Day 14, stock = -180 units (180 units short)')

    doc.add_heading('Interpretation', 2)
    doc.add_paragraph(
        'There is an 8.7% chance of running out of stock within the next 14 days. '
        'The risk increases significantly after day 10. Recommend placing order soon to '
        'reduce stockout risk.'
    )

    doc.add_page_break()

    # TOOL 5
    doc.add_heading('8. Tool 5: Evaluate Holding Costs (EOQ)', 1)

    doc.add_heading('What This Tool Does', 2)
    doc.add_paragraph(
        'Tool 5 calculates the Economic Order Quantity (EOQ) - the optimal order size '
        'that minimizes total inventory costs (ordering cost + holding cost). It also '
        'evaluates costs at different inventory levels.'
    )

    doc.add_heading('Tool Definition', 2)
    p = doc.add_paragraph()
    p.add_run('Function Name: ').bold = True
    p.add_run('evaluate_holding_costs(...)')

    p = doc.add_paragraph()
    p.add_run('Inputs: ').bold = True

    inputs = [
        'current_stock: 1,850 units',
        'safety_stock: 300 units',
        'holding_cost_per_unit_per_year: $2.30',
        'stockout_cost_per_unit: $45.00',
        'avg_daily_demand: 142 units/day',
        'warehouse_capacity: 5,000 units',
    ]
    for item in inputs:
        doc.add_paragraph(f'• {item}', style='List Bullet')

    doc.add_heading('EOQ Formula', 2)

    formula = doc.add_paragraph(
        'EOQ = sqrt((2 × annual_demand × ordering_cost) / holding_cost_per_unit_per_year)\n\n'
        'where:\n'
        '  annual_demand = daily_demand × 365\n'
        '  ordering_cost = fixed cost per order ($150)',
        style='No Spacing'
    )
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True

    doc.add_heading('Example EOQ Calculation', 2)

    calc = doc.add_paragraph(
        'Annual Demand:\n'
        'annual_demand = 142 × 365\n'
        '              = 51,830 units\n\n'
        'EOQ Calculation:\n'
        'EOQ = sqrt((2 × 51,830 × $150) / $2.30)\n'
        '    = sqrt(15,549,000 / 2.30)\n'
        '    = sqrt(6,760,435)\n'
        '    = 2,600 units\n\n'
        'Reorder Frequency:\n'
        'frequency = EOQ / daily_demand\n'
        '          = 2,600 / 142\n'
        '          = 18.3 days\n\n'
        'Annual Ordering Cost:\n'
        'orders_per_year = 51,830 / 2,600 = 19.9 orders\n'
        'ordering_cost = 19.9 × $150 = $2,985\n\n'
        'Annual Holding Cost:\n'
        'avg_inventory = EOQ / 2 = 2,600 / 2 = 1,300 units\n'
        'holding_cost = 1,300 × $2.30 = $2,990',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Optimal: Order 2,600 units every 18.3 days').bold = True
    p.add_run('\nTotal annual cost: $5,975 (ordering + holding)')

    doc.add_heading('Cost Comparison at Different Inventory Levels', 2)

    table = doc.add_table(rows=7, cols=5)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Stock Level'
    header[1].text = 'Days Supply'
    header[2].text = 'Annual Holding Cost'
    header[3].text = 'Stockout Risk %'
    header[4].text = 'Warehouse Use %'

    comparison_data = [
        ('300 (safety)', '2.1', '$690', '75%', '6%'),
        ('994 (1 week)', '7.0', '$2,286', '15%', '20%'),
        ('1,350 (EOQ/2)', '9.5', '$3,105', '8%', '27%'),
        ('1,850 (current)', '13.0', '$4,255', '2%', '37%'),
        ('1,988 (2 weeks)', '14.0', '$4,572', '1.5%', '40%'),
        ('2,350 (current+500)', '16.5', '$5,405', '0.5%', '47%'),
    ]

    for i, (level, days, cost, risk, util) in enumerate(comparison_data, start=1):
        row = table.rows[i].cells
        row[0].text = level
        row[1].text = days
        row[2].text = cost
        row[3].text = risk
        row[4].text = util

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Current Position Analysis:').bold = True
    doc.add_paragraph('• Current stock: 1,850 units (13 days of supply)')
    doc.add_paragraph('• Annual holding cost: $4,255')
    doc.add_paragraph('• Capital tied up: 1,850 × $52 = $96,200')
    doc.add_paragraph('• Warehouse utilization: 37%')
    doc.add_paragraph('• Stockout risk: 2% (low)')

    doc.add_heading('Historical Stockout Impact', 2)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Metric'
    header[1].text = 'Last 12 Months'

    stockout_data = [
        ('Stockout Events', '3'),
        ('Total Revenue Lost', '$111,875'),
        ('Avg Cost Per Event', '$37,292'),
    ]

    for i, (metric, value) in enumerate(stockout_data, start=1):
        row = table.rows[i].cells
        row[0].text = metric
        row[1].text = value

    doc.add_paragraph()
    doc.add_paragraph(
        'The 3 stockout events cost $111,875 in lost revenue. This justifies maintaining '
        'adequate safety stock despite the $4,255 annual holding cost.'
    )

    doc.add_page_break()

    # TOOL 6
    doc.add_heading('9. Tool 6: Generate Replenishment Plan', 1)

    doc.add_heading('What This Tool Does', 2)
    doc.add_paragraph(
        'Tool 6 generates a week-by-week replenishment schedule for the next 4 weeks. '
        'It determines when to order, how much to order, and which suppliers to use. '
        'It allocates orders across suppliers (65% Supplier A, 35% Supplier B) based '
        'on reliability and cost.'
    )

    doc.add_heading('Tool Definition', 2)
    p = doc.add_paragraph()
    p.add_run('Function Name: ').bold = True
    p.add_run('generate_replenishment_plan(...)')

    p = doc.add_paragraph()
    p.add_run('Inputs: ').bold = True

    inputs = [
        'expected_weekly_demand: 1,050 units/week (from demand forecast)',
        'demand_std_dev_weekly: 120 units',
        'planning_horizon_weeks: 4 weeks',
        'target_service_level_pct: 95%',
        'current_stock: 1,850 units',
    ]
    for item in inputs:
        doc.add_paragraph(f'• {item}', style='List Bullet')

    doc.add_heading('Planning Algorithm', 2)

    doc.add_paragraph('For each week in the planning horizon:')

    steps = [
        '1. Start with current stock level',
        '2. Add any incoming orders scheduled for this week',
        '3. Subtract expected weekly demand (with random variation)',
        '4. Check if projected stock for next week < safety stock × 1.5',
        '5. If yes, place order:',
        '   a. Calculate target stock = 2 weeks demand + safety stock',
        '   b. Order quantity = target - current stock',
        '   c. Split order: 65% Supplier A, 35% Supplier B',
        '   d. Respect min/max order quantities',
        '6. Record week plan',
        '7. Repeat for next week',
    ]

    for step in steps:
        doc.add_paragraph(step)

    doc.add_heading('Supplier Allocation Logic', 2)

    table = doc.add_table(rows=3, cols=4)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Supplier'
    header[1].text = 'Allocation %'
    header[2].text = 'Why?'
    header[3].text = 'Unit Cost'

    row1 = table.rows[1].cells
    row1[0].text = 'Supplier A'
    row1[1].text = '65%'
    row1[2].text = 'Higher reliability (92.5% on-time), faster (4.2 days)'
    row1[3].text = '$52.00'

    row2 = table.rows[2].cells
    row2[0].text = 'Supplier B'
    row2[1].text = '35%'
    row2[2].text = 'Lower cost ($49.50), diversification'
    row2[3].text = '$49.50'

    doc.add_heading('Example 4-Week Plan', 2)

    table = doc.add_table(rows=5, cols=6)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Week'
    header[1].text = 'Start Stock'
    header[2].text = 'Incoming'
    header[3].text = 'Demand'
    header[4].text = 'End Stock'
    header[5].text = 'Order Placed?'

    plan_data = [
        ('1', '1,850', '800', '1,050', '1,600', 'No'),
        ('2', '1,600', '0', '1,050', '550', 'Yes - 1,800 units'),
        ('3', '550', '1,200', '1,050', '700', 'Yes - 1,800 units'),
        ('4', '700', '1,200', '1,050', '850', 'No'),
    ]

    for i, (week, start, incoming, demand, end, order) in enumerate(plan_data, start=1):
        row = table.rows[i].cells
        row[0].text = week
        row[1].text = start
        row[2].text = incoming
        row[3].text = demand
        row[4].text = end
        row[5].text = order

    doc.add_heading('Week 2 Order Details (Example)', 2)

    calc = doc.add_paragraph(
        'Target Stock = 2 × 1,050 + 376 = 2,476 units\n'
        'Current Stock = 550 units\n'
        'Order Quantity = 2,476 - 550 = 1,926 units\n\n'
        'Supplier A (65%):\n'
        '  Quantity = 1,926 × 0.65 = 1,252 → rounded to 1,300 (min 100)\n'
        '  Cost = 1,300 × $52.00 = $67,600\n'
        '  ETA = 4 days (arrives Week 3)\n\n'
        'Supplier B (35%):\n'
        '  Quantity = 1,926 × 0.35 = 674 → rounded to 700 (min 200)\n'
        '  Cost = 700 × $49.50 = $34,650\n'
        '  ETA = 7 days (arrives Week 3)\n\n'
        'Total Order: 2,000 units for $102,250',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    doc.add_heading('Plan Summary', 2)

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Metric'
    header[1].text = 'Value'

    summary_data = [
        ('Total Units Ordered', '3,600 units (2 orders)'),
        ('Total Procurement Cost', '$187,200'),
        ('Avg Projected Stock', '925 units'),
        ('Min Projected Stock', '550 units (Week 2)'),
        ('Weeks Below Safety Stock', '2 out of 4 weeks'),
    ]

    for i, (metric, value) in enumerate(summary_data, start=1):
        row = table.rows[i].cells
        row[0].text = metric
        row[1].text = value

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Action Required:').bold = True
    doc.add_paragraph('• Week 2: Place order for 2,000 units ($102,250)')
    doc.add_paragraph('• Week 3: Place order for 1,800 units ($93,600)')
    doc.add_paragraph('• Monitor stock levels closely in Week 2 (only 550 units)')

    doc.add_page_break()

def add_remaining_sections(doc):
    """Add sections 10-13"""

    # SECTION 10: Complete Execution Flow
    doc.add_heading('10. Complete Execution Flow with Examples', 1)

    doc.add_paragraph(
        'This section shows the COMPLETE flow from start to finish with REAL NUMBERS.'
    )

    # STEP 1
    doc.add_heading('STEP 1: User Clicks "Run Analysis"', 2)
    p = doc.add_paragraph()
    p.add_run('User action: ').bold = True
    p.add_run('Clicks "Run Analysis" in Inventory Management UI for PROD-A')

    # STEP 2
    doc.add_heading('STEP 2: Agent Calls Tool 1 - Get Inventory Status', 2)
    p = doc.add_paragraph()
    p.add_run('Tool called: ').bold = True
    p.add_run('get_inventory_status(product_id="PROD-A")')

    doc.add_paragraph('Returns:')
    doc.add_paragraph('• Current stock: 1,850 units')
    doc.add_paragraph('• Safety stock: 300 units')
    doc.add_paragraph('• Days of supply: 13 days')
    doc.add_paragraph('• Incoming orders: 800 units (500 in 4 days, 300 in 7 days)')
    doc.add_paragraph('• 3 stockout events in last 12 months ($111,875 lost)')
    doc.add_paragraph('• Supplier A: 92.5% on-time, 4.2 day lead time, $52/unit')
    doc.add_paragraph('• Supplier B: 85% on-time, 6.8 day lead time, $49.50/unit')

    # STEP 3
    doc.add_heading('STEP 3: Agent Calls Tool 2 - Calculate ROP', 2)
    p = doc.add_paragraph()
    p.add_run('Agent thinks: ').bold = True
    p.add_run('"I see current stock is 1,850. I should calculate when to reorder."')

    p = doc.add_paragraph()
    p.add_run('Tool called: ').bold = True
    p.add_run('calculate_reorder_point(daily_demand=142, lead_time=5, service_level=95%)')

    doc.add_paragraph('Calculation:')
    calc = doc.add_paragraph(
        'SS = 1.645 × sqrt(5 × 625 + 20,164) = 251 units\n'
        'ROP = (142 × 5) + 251 = 961 units\n\n'
        'Current: 1,850 > ROP of 961 ✓\n'
        'Days until ROP: (1850 - 961) / 142 = 6.3 days',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    # STEP 4
    doc.add_heading('STEP 4: Agent Calls Tool 3 - Optimize Safety Stock', 2)
    p = doc.add_paragraph()
    p.add_run('Agent thinks: ').bold = True
    p.add_run('"Safety stock is 300 units. Is this optimal?"')

    p = doc.add_paragraph()
    p.add_run('Tool called: ').bold = True
    p.add_run('optimize_safety_stock(holding_cost=$2.30, stockout_cost=$45)')

    doc.add_paragraph('Result:')
    doc.add_paragraph('• Current safety stock: 300 units (~96% service level)')
    doc.add_paragraph('• Optimal safety stock: 376 units (99% service level)')
    doc.add_paragraph('• Current annual cost: $690')
    doc.add_paragraph('• Optimal annual cost: $3,610')
    doc.add_paragraph('• Recommendation: Increase safety stock by 76 units')

    # STEP 5
    doc.add_heading('STEP 5: Agent Calls Tool 4 - Simulate Stockout Risk', 2)
    p = doc.add_paragraph()
    p.add_run('Tool called: ').bold = True
    p.add_run('simulate_stockout_risk(current_stock=1850, horizon_days=14)')

    doc.add_paragraph('Monte Carlo Result:')
    doc.add_paragraph('• Overall stockout probability: 8.7% (over 14 days)')
    doc.add_paragraph('• Critical days: Day 10 (2.3% risk), Day 14 (8.7% risk)')
    doc.add_paragraph('• Worst case: -180 units on Day 14')
    doc.add_paragraph('• Recommendation: Place order soon to reduce risk')

    # STEP 6
    doc.add_heading('STEP 6: Agent Calls Tool 5 - Evaluate Holding Costs', 2)
    p = doc.add_paragraph()
    p.add_run('Tool called: ').bold = True
    p.add_run('evaluate_holding_costs(current_stock=1850)')

    doc.add_paragraph('EOQ Result:')
    doc.add_paragraph('• Economic Order Quantity: 2,600 units')
    doc.add_paragraph('• Optimal reorder frequency: Every 18.3 days')
    doc.add_paragraph('• Current holding cost: $4,255/year')
    doc.add_paragraph('• Capital tied up: $96,200')

    # STEP 7
    doc.add_heading('STEP 7: Agent Calls Tool 6 - Generate Plan', 2)
    p = doc.add_paragraph()
    p.add_run('Tool called: ').bold = True
    p.add_run('generate_replenishment_plan(weekly_demand=1050, horizon_weeks=4)')

    doc.add_paragraph('4-Week Plan Result:')
    doc.add_paragraph('• Week 1: Start with 1,850, receive 800, demand 1,050 → End with 1,600')
    doc.add_paragraph('• Week 2: Start with 1,600, demand 1,050 → End with 550 (LOW!) → ORDER 2,000 units')
    doc.add_paragraph('• Week 3: Start with 550, receive 1,200, demand 1,050 → End with 700 → ORDER 1,800 units')
    doc.add_paragraph('• Week 4: Start with 700, receive 1,200, demand 1,050 → End with 850')
    doc.add_paragraph()
    doc.add_paragraph('Total orders: 3,600 units for $187,200')

    # STEP 8
    doc.add_heading('STEP 8: Agent Generates Recommendations', 2)

    p = doc.add_paragraph()
    p.add_run('Agent reasoning:').bold = True

    reasoning = doc.add_paragraph(
        '"Based on all 6 tools:\n\n'
        '1. Current stock (1,850) is ABOVE reorder point (961) - safe for now\n'
        '2. But stockout risk is 8.7% over next 14 days - concerning\n'
        '3. Safety stock should be increased from 300 to 376 units\n'
        '4. EOQ is 2,600 units - order sizes should be around this amount\n'
        '5. Week 2 will drop to 550 units - CRITICAL LOW POINT\n'
        '6. Need to place 2 orders in next 4 weeks\n\n'
        'RECOMMENDATIONS:\n'
        '• IMMEDIATE: Place order for 2,000 units this week (split 65/35 between suppliers)\n'
        '• Increase safety stock target from 300 to 376 units\n'
        '• Place second order for 1,800 units in Week 3\n'
        '• Monitor stock closely in Week 2 (projected 550 units)\n'
        '• Consider increasing order frequency to match EOQ of 2,600 units"',
        style='Intense Quote'
    )
    reasoning.runs[0].font.italic = True
    reasoning.runs[0].font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Confidence: HIGH ').bold = True
    p.add_run('(based on comprehensive 6-tool analysis)')

    doc.add_page_break()

    # SECTION 11: All SQL Queries
    doc.add_heading('11. All SQL Queries Reference', 1)

    doc.add_paragraph(
        'This section lists ALL SQL queries used in the Inventory Management module.'
    )

    queries = [
        (
            'Query 1: Get Current Inventory',
            'SELECT * FROM inventory\n'
            'WHERE product_id = \'PROD-A\';'
        ),
        (
            'Query 2: Get All Products Inventory',
            'SELECT product_id, current_stock, safety_stock,\n'
            '       reorder_point, avg_daily_usage\n'
            'FROM inventory\n'
            'ORDER BY product_id;'
        ),
        (
            'Query 3: Check Stock Levels',
            'SELECT product_id, current_stock, safety_stock,\n'
            '       (current_stock - safety_stock) as buffer,\n'
            '       CASE\n'
            '         WHEN current_stock < safety_stock THEN \'CRITICAL\'\n'
            '         WHEN current_stock < reorder_point THEN \'LOW\'\n'
            '         ELSE \'OK\'\n'
            '       END as status\n'
            'FROM inventory;'
        ),
        (
            'Query 4: Calculate Days of Supply',
            'SELECT product_id,\n'
            '       current_stock,\n'
            '       avg_daily_usage,\n'
            '       ROUND(current_stock / avg_daily_usage, 1) as days_of_supply\n'
            'FROM inventory\n'
            'WHERE avg_daily_usage > 0;'
        ),
        (
            'Query 5: View Inventory History (Last 30 Days)',
            'SELECT date, stock_level, stock_value\n'
            'FROM inventory_history\n'
            'WHERE product_id = \'PROD-A\'\n'
            'ORDER BY date DESC\n'
            'LIMIT 30;'
        ),
        (
            'Query 6: Calculate Inventory Turnover',
            'SELECT product_id,\n'
            '       (avg_daily_usage * 365) as annual_demand,\n'
            '       current_stock,\n'
            '       ROUND((avg_daily_usage * 365) / current_stock, 2) as turnover_ratio\n'
            'FROM inventory\n'
            'WHERE current_stock > 0;'
        ),
        (
            'Query 7: Update Inventory Level',
            'UPDATE inventory\n'
            'SET current_stock = 2000,\n'
            '    last_updated = CURRENT_TIMESTAMP\n'
            'WHERE product_id = \'PROD-A\';'
        ),
        (
            'Query 8: Update Reorder Point',
            'UPDATE inventory\n'
            'SET reorder_point = 961,\n'
            '    safety_stock = 376\n'
            'WHERE product_id = \'PROD-A\';'
        ),
    ]

    for i, (title, query) in enumerate(queries, start=1):
        doc.add_heading(title, 3)
        code = doc.add_paragraph(query, style='No Spacing')
        code.runs[0].font.name = 'Courier New'
        code.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    doc.add_page_break()

    # SECTION 12: Validation
    doc.add_heading('12. Validation & Verification', 1)

    doc.add_paragraph(
        'This section shows how to validate that the Inventory Management module is working correctly.'
    )

    doc.add_heading('Validation 1: Verify Inventory Data Exists', 2)
    code = doc.add_paragraph(
        'SELECT COUNT(*) FROM inventory;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    doc.add_paragraph('Expected: 3 (one record per product)')

    doc.add_heading('Validation 2: Verify ROP Calculation', 2)
    doc.add_paragraph('Step 1: Get current stock and daily usage:')
    code = doc.add_paragraph(
        'SELECT current_stock, avg_daily_usage, lead_time\n'
        'FROM inventory WHERE product_id = \'PROD-A\';',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Step 2: Calculate ROP manually:')
    doc.add_paragraph('ROP = (142 × 5) + 251 = 961 units')

    doc.add_paragraph('Step 3: Compare to Tool 2 output')
    doc.add_paragraph('Tool 2 should return ROP = 961 units')

    doc.add_heading('Validation 3: Verify EOQ Calculation', 2)
    doc.add_paragraph('Manual calculation:')
    calc = doc.add_paragraph(
        'annual_demand = 142 × 365 = 51,830\n'
        'EOQ = sqrt((2 × 51,830 × 150) / 2.30)\n'
        '    = sqrt(6,760,435)\n'
        '    = 2,600 units',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    doc.add_paragraph('Tool 5 should return EOQ = 2,600 units')

    doc.add_heading('Validation 4: Verify Stockout Simulation', 2)
    doc.add_paragraph('Run Tool 4 twice with same inputs')
    doc.add_paragraph('Results should be similar but not identical (Monte Carlo randomness)')
    doc.add_paragraph('Stockout probability should be within ±2% between runs')

    doc.add_page_break()

    # SECTION 13: Formulas Cheat Sheet
    doc.add_heading('13. Formulas & Calculations Cheat Sheet', 1)

    formulas = [
        ('Reorder Point (ROP)', 'ROP = (avg_daily_demand × lead_time) + safety_stock'),
        ('Safety Stock',
         'SS = Z × sqrt(LT × σ_demand² + demand² × σ_LT²)'),
        ('Economic Order Quantity (EOQ)',
         'EOQ = sqrt((2 × annual_demand × ordering_cost) / holding_cost)'),
        ('Days of Supply', 'days_of_supply = current_stock / avg_daily_usage'),
        ('Effective Days of Supply',
         'effective_days = (current_stock + incoming) / avg_daily_usage'),
        ('Annual Holding Cost',
         'annual_holding = inventory_level × holding_cost_per_unit_per_year'),
        ('Inventory Turnover', 'turnover = annual_demand / avg_inventory'),
        ('Service Level from Z-Score',
         '95% → Z=1.645, 97% → Z=1.88, 99% → Z=2.326'),
        ('Stockout Probability', 'stockout_prob = 1 - (service_level / 100)'),
        ('Total Inventory Cost',
         'total_cost = holding_cost + ordering_cost + stockout_cost'),
    ]

    for formula_name, formula in formulas:
        doc.add_heading(formula_name, 3)
        code = doc.add_paragraph(formula, style='No Spacing')
        code.runs[0].font.name = 'Courier New'
        code.runs[0].font.size = Pt(10)
        code.runs[0].font.italic = True
        doc.add_paragraph()

    doc.add_heading('Quick Reference Table', 2)

    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Metric'
    header[1].text = 'Example Value'
    header[2].text = 'What It Means'

    data = [
        ('ROP', '961 units', 'Order when stock drops to this level'),
        ('Safety Stock', '376 units', 'Buffer against uncertainty'),
        ('EOQ', '2,600 units', 'Optimal order quantity'),
        ('Lead Time', '5 days', 'Time from order to delivery'),
        ('Service Level', '95%', 'Probability of not stocking out'),
        ('Z-Score', '1.645', 'Statistical multiplier for 95% service'),
    ]

    for i, (metric, value, meaning) in enumerate(data, start=1):
        row = table.rows[i].cells
        row[0].text = metric
        row[1].text = value
        row[2].text = meaning

def generate_document():
    """Generate the complete Word document"""
    doc = Document()

    # Configure default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    print('[1/8] Adding title page...')
    add_title_page(doc)

    print('[2/8] Adding table of contents...')
    add_toc(doc)

    print('[3/8] Adding Sections 1-3...')
    add_section_1(doc)
    add_section_2(doc)
    add_section_3(doc)

    print('[4/8] Adding Sections 4-9 (All 6 Tools)...')
    generate_tool_sections(doc)

    print('[5/8] Adding Sections 10-13...')
    add_remaining_sections(doc)

    # Save document
    output_path = 'INVENTORY_MANAGEMENT_COMPLETE_GUIDE.docx'
    doc.save(output_path)

    print(f'\n[SUCCESS] Document created: {output_path}')
    print('Document contains:')
    print('  - 13 comprehensive sections')
    print('  - All 6 tools explained in detail')
    print('  - Every calculation with examples')
    print('  - ROP, EOQ, Safety Stock formulas')
    print('  - Monte Carlo simulation explained')
    print('  - All SQL queries')
    print('  - Complete step-by-step flow')
    print('  - Validation guides')
    print('  - Formula cheat sheet')

if __name__ == '__main__':
    generate_document()
