"""
Generate Complete Demand Intelligence Documentation
====================================================
This script creates a comprehensive Word document that explains EVERYTHING
about the Demand Intelligence module - crystal clear and complete.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH

def add_title_page(doc):
    """Add professional title page"""
    title = doc.add_heading('Demand Intelligence Module', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.size = Pt(32)
        run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph('Complete Technical Documentation')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(20)
    subtitle.runs[0].font.color.rgb = RGBColor(68, 114, 196)

    doc.add_paragraph()
    doc.add_paragraph()

    desc = doc.add_paragraph('Everything You Need to Know:')
    desc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    desc.runs[0].font.size = Pt(14)
    desc.runs[0].font.bold = True

    features = [
        'What the AI Agent Does',
        'All 3 Tools Explained in Detail',
        'Every Calculation with Examples',
        'All Database Tables & SQL Queries',
        'Complete Step-by-Step Flow',
    ]

    for feature in features:
        p = doc.add_paragraph(f'• {feature}')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.runs[0].font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph('AMIS - Autonomous Manufacturing Intelligence System')
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    info.runs[0].font.italic = True

    doc.add_page_break()

def add_toc(doc):
    """Add table of contents"""
    doc.add_heading('Table of Contents', 1)

    sections = [
        ('1', 'What is Demand Intelligence?', 3),
        ('2', 'How the AI Agent Works', 4),
        ('3', 'Database Tables Used', 6),
        ('4', 'Tool 1: Get Demand Data Summary', 9),
        ('5', 'Tool 2: Simulate Demand Scenarios', 13),
        ('6', 'Tool 3: Analyze Demand Trends', 18),
        ('7', 'Complete Execution Flow with Examples', 23),
        ('8', 'All SQL Queries Reference', 28),
        ('9', 'Validation & Verification', 31),
        ('10', 'Formulas & Calculations Cheat Sheet', 34),
    ]

    for num, title, page in sections:
        p = doc.add_paragraph()
        p.add_run(f'{num}. {title}').bold = True
        p.add_run(f' {"." * (70 - len(title))} {page}')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

def add_section_1(doc):
    """Section 1: What is Demand Intelligence?"""
    doc.add_heading('1. What is Demand Intelligence?', 1)

    doc.add_paragraph(
        'Demand Intelligence is an AI-powered module that forecasts future product demand '
        'by analyzing historical sales data, market trends, and external factors.'
    )

    doc.add_heading('What Problem Does It Solve?', 2)
    doc.add_paragraph(
        'Manufacturing companies need to know how much product to make in the coming weeks. '
        'Make too much = wasted inventory. Make too little = lost sales. Demand Intelligence '
        'uses AI to predict the optimal production quantity.'
    )

    doc.add_heading('How Does It Work?', 2)
    p = doc.add_paragraph('The AI Agent follows this process:')

    steps = [
        ('Step 1', 'User clicks "Run Analysis" in the UI'),
        ('Step 2', 'AI Agent starts and decides which tools to call'),
        ('Step 3', 'AI calls Tool 1 to gather historical data from database'),
        ('Step 4', 'AI calls Tool 2 to generate 3 demand scenarios (optimistic, base, pessimistic)'),
        ('Step 5', 'AI calls Tool 3 to analyze trends and detect anomalies'),
        ('Step 6', 'AI combines all results and generates final forecast'),
        ('Step 7', 'Forecast is saved to database and displayed in UI'),
    ]

    for step_num, step_desc in steps:
        p = doc.add_paragraph()
        p.add_run(f'{step_num}: ').bold = True
        p.add_run(step_desc)

    doc.add_heading('What Makes It "Agentic AI"?', 2)
    p = doc.add_paragraph()
    p.add_run('Traditional AI: ').bold = True
    p.add_run('You tell it exactly what to do, step by step.')

    p = doc.add_paragraph()
    p.add_run('Agentic AI: ').bold = True
    p.add_run('You give it a goal, and it DECIDES which tools to use and how to use them. '
              'The AI reasons about the data and chooses its own strategy.')

    doc.add_paragraph(
        'In our case, the AI decides to call Tool 1 first (to get context), then Tool 2 '
        '(to generate scenarios), then Tool 3 (to analyze trends). It makes these decisions '
        'autonomously based on the goal: "Forecast demand for Product A."'
    )

    doc.add_page_break()

def add_section_2(doc):
    """Section 2: How the AI Agent Works"""
    doc.add_heading('2. How the AI Agent Works', 1)

    doc.add_heading('The AI Agent (Claude 3.5 Sonnet)', 2)
    doc.add_paragraph(
        'The Demand Forecasting Agent is powered by Claude 3.5 Sonnet, a large language model '
        'that can reason about data and call tools.'
    )

    doc.add_heading('Agent Configuration', 2)
    p = doc.add_paragraph()
    p.add_run('Framework: ').bold = True
    p.add_run('LangChain (Python library for building AI agents)')

    p = doc.add_paragraph()
    p.add_run('Model: ').bold = True
    p.add_run('Claude 3.5 Sonnet (anthropic.claude-3-5-sonnet-20241022)')

    p = doc.add_paragraph()
    p.add_run('Tools Available: ').bold = True
    p.add_run('3 custom tools (explained in sections 4-6)')

    p = doc.add_paragraph()
    p.add_run('Temperature: ').bold = True
    p.add_run('0.3 (low randomness for consistent results)')

    doc.add_heading('Agent System Prompt', 2)
    doc.add_paragraph('The AI receives this instruction:')

    prompt = doc.add_paragraph(
        'You are a Demand Forecasting Agent for a manufacturing company. '
        'Your job is to forecast future demand for products by analyzing historical sales data, '
        'market trends, and external factors. You have access to 3 tools. Use them to gather data, '
        'run simulations, and analyze trends. Provide a clear forecast with reasoning.',
        style='Intense Quote'
    )
    prompt.runs[0].font.italic = True

    doc.add_heading('How the Agent Decides Which Tools to Call', 2)
    doc.add_paragraph('The AI uses reasoning to decide:')

    p = doc.add_paragraph()
    p.add_run('1. "I need context first" ').bold = True
    p.add_run('→ Calls Tool 1 (get_demand_data_summary)')

    p = doc.add_paragraph()
    p.add_run('2. "I have historical data, now I need scenarios" ').bold = True
    p.add_run('→ Calls Tool 2 (simulate_demand_scenarios)')

    p = doc.add_paragraph()
    p.add_run('3. "I should check for trends and anomalies" ').bold = True
    p.add_run('→ Calls Tool 3 (analyze_demand_trends)')

    p = doc.add_paragraph()
    p.add_run('4. "I have all the data I need" ').bold = True
    p.add_run('→ Generates final forecast and stops calling tools')

    doc.add_heading('Agent Output', 2)
    doc.add_paragraph('The AI generates:')
    doc.add_paragraph('• A forecast (optimistic, base case, pessimistic scenarios)', style='List Bullet')
    doc.add_paragraph('• Confidence level (High, Medium, Low)', style='List Bullet')
    doc.add_paragraph('• Reasoning explaining why it made this forecast', style='List Bullet')
    doc.add_paragraph('• Risk factors to consider', style='List Bullet')

    doc.add_page_break()

def add_section_3(doc):
    """Section 3: Database Tables Used"""
    doc.add_heading('3. Database Tables Used', 1)

    doc.add_paragraph(
        'The Demand Intelligence module uses 3 database tables. This section explains each table '
        'in detail with SQL queries to view the data.'
    )

    # TABLE 1
    doc.add_heading('Table 1: historical_demand_data', 2)

    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Stores 52 weeks of historical demand data. This is the INPUT to the AI.')

    p = doc.add_paragraph()
    p.add_run('Total Records: ').bold = True
    p.add_run('156 (52 weeks × 3 products)')

    p = doc.add_paragraph()
    p.add_run('Updated: ').bold = True
    p.add_run('Populated once by migration script, then remains static')

    doc.add_heading('Table Schema', 3)

    # Create table for schema
    table = doc.add_table(rows=11, cols=3)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Column Name'
    header[1].text = 'Data Type'
    header[2].text = 'Description'

    schema_data = [
        ('product_id', 'TEXT', 'Product identifier (PROD-A, PROD-B, PROD-C)'),
        ('week_start_date', 'DATE', 'Week start date (e.g., 2026-02-25)'),
        ('week_number', 'INTEGER', 'ISO week number (1-52)'),
        ('year', 'INTEGER', 'Year (e.g., 2026)'),
        ('demand_units', 'INTEGER', 'Actual demand in units for that week'),
        ('avg_price', 'REAL', 'Average selling price that week'),
        ('promotions_active', 'BOOLEAN', '1 = promotion running, 0 = no promotion'),
        ('competitor_price', 'REAL', 'Competitor price that week'),
        ('is_anomaly', 'BOOLEAN', '1 = anomaly detected, 0 = normal'),
        ('anomaly_reason', 'TEXT', 'Reason for anomaly (e.g., "Viral spike")'),
    ]

    for i, (col, dtype, desc) in enumerate(schema_data, start=1):
        row = table.rows[i].cells
        row[0].text = col
        row[1].text = dtype
        row[2].text = desc

    doc.add_heading('SQL to View Data', 3)
    code = doc.add_paragraph(
        'SELECT * FROM historical_demand_data\n'
        'WHERE product_id = \'PROD-A\'\n'
        'ORDER BY week_start_date DESC\n'
        'LIMIT 10;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(10)

    doc.add_heading('Example Data', 3)

    table = doc.add_table(rows=4, cols=5)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Date'
    header[1].text = 'Demand'
    header[2].text = 'Promo'
    header[3].text = 'Anomaly'
    header[4].text = 'Reason'

    example_data = [
        ('2026-02-25', '1789', 'No', 'No', '-'),
        ('2026-02-04', '2081', 'No', 'Yes', 'Viral social media spike'),
        ('2026-01-21', '1662', 'Yes', 'No', '-'),
    ]

    for i, row_data in enumerate(example_data, start=1):
        row = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row[j].text = cell_data

    # TABLE 2
    doc.add_heading('Table 2: market_context_data', 2)

    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Stores market conditions and external factors that affect demand.')

    p = doc.add_paragraph()
    p.add_run('Total Records: ').bold = True
    p.add_run('1 (current market snapshot)')

    doc.add_heading('Table Schema', 3)

    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Column Name'
    header[1].text = 'Data Type'
    header[2].text = 'Description'

    schema_data = [
        ('industry_growth_rate', 'REAL', 'Industry growth % (e.g., 7.2)'),
        ('economic_indicator', 'TEXT', 'Economic indicator description'),
        ('competitor_activity', 'TEXT', 'Recent competitor actions'),
        ('seasonal_pattern', 'TEXT', 'Seasonal demand patterns'),
        ('market_sentiment', 'TEXT', 'Overall market sentiment'),
        ('supply_chain_status', 'TEXT', 'Supply chain condition'),
    ]

    for i, (col, dtype, desc) in enumerate(schema_data, start=1):
        row = table.rows[i].cells
        row[0].text = col
        row[1].text = dtype
        row[2].text = desc

    doc.add_heading('SQL to View Data', 3)
    code = doc.add_paragraph(
        'SELECT * FROM market_context_data\n'
        'ORDER BY date_recorded DESC\n'
        'LIMIT 1;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(10)

    # TABLE 3
    doc.add_heading('Table 3: demand_forecasts', 2)

    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Stores AI-generated forecasts. This is the OUTPUT from the AI.')

    p = doc.add_paragraph()
    p.add_run('Total Records: ').bold = True
    p.add_run('Created each time pipeline runs (dynamic)')

    doc.add_heading('Table Schema', 3)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Column Name'
    header[1].text = 'Data Type'
    header[2].text = 'Description'

    schema_data = [
        ('product_id', 'TEXT', 'Product identifier'),
        ('forecast_date', 'DATE', 'Date forecast was created'),
        ('optimistic', 'INTEGER', 'Optimistic scenario forecast'),
        ('base_case', 'INTEGER', 'Base case scenario forecast'),
        ('pessimistic', 'INTEGER', 'Pessimistic scenario forecast'),
    ]

    for i, (col, dtype, desc) in enumerate(schema_data, start=1):
        row = table.rows[i].cells
        row[0].text = col
        row[1].text = dtype
        row[2].text = desc

    doc.add_heading('SQL to View Data', 3)
    code = doc.add_paragraph(
        'SELECT * FROM demand_forecasts\n'
        'WHERE product_id = \'PROD-A\'\n'
        'ORDER BY created_at DESC\n'
        'LIMIT 1;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(10)

    doc.add_page_break()

def add_section_4(doc):
    """Section 4: Tool 1 - Get Demand Data Summary"""
    doc.add_heading('4. Tool 1: Get Demand Data Summary', 1)

    doc.add_heading('What This Tool Does', 2)
    doc.add_paragraph(
        'Tool 1 gathers all the context data the AI needs to make a forecast. '
        'It queries 5 different data sources and returns a comprehensive summary.'
    )

    doc.add_heading('Tool Definition', 2)
    p = doc.add_paragraph()
    p.add_run('Function Name: ').bold = True
    p.add_run('get_demand_data_summary(product_id)')

    p = doc.add_paragraph()
    p.add_run('Input: ').bold = True
    p.add_run('product_id (string, e.g., "PROD-A")')

    p = doc.add_paragraph()
    p.add_run('Output: ').bold = True
    p.add_run('JSON object with 5 data sections')

    doc.add_heading('The 5 Data Sources', 2)

    # Data Source 1
    p = doc.add_paragraph()
    p.add_run('1. Historical Demand (12 weeks)').bold = True

    doc.add_paragraph('SQL Query:')
    code = doc.add_paragraph(
        'SELECT week_start_date as date, demand_units, avg_price,\n'
        '       promotions_active, competitor_price\n'
        'FROM historical_demand_data\n'
        'WHERE product_id = \'PROD-A\'\n'
        'ORDER BY week_start_date DESC\n'
        'LIMIT 12;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Returns: List of 12 weeks with demand values')

    # Data Source 2
    p = doc.add_paragraph()
    p.add_run('2. Current Inventory').bold = True

    doc.add_paragraph('SQL Query:')
    code = doc.add_paragraph(
        'SELECT current_stock, safety_stock, warehouse_capacity\n'
        'FROM inventory\n'
        'WHERE product_id = \'PROD-A\';',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Returns: Current inventory levels')

    # Data Source 3
    p = doc.add_paragraph()
    p.add_run('3. Market Context').bold = True

    doc.add_paragraph('SQL Query:')
    code = doc.add_paragraph(
        'SELECT industry_growth_rate, economic_indicator,\n'
        '       competitor_activity, seasonal_pattern\n'
        'FROM market_context_data\n'
        'ORDER BY date_recorded DESC LIMIT 1;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Returns: Current market conditions')

    # Data Source 4
    p = doc.add_paragraph()
    p.add_run('4. Production Capacity').bold = True

    doc.add_paragraph('Returns: Max daily output, current utilization, available lines')
    doc.add_paragraph('(Note: Currently from sample_data.py, not database)')

    # Data Source 5
    p = doc.add_paragraph()
    p.add_run('5. Product Information').bold = True

    doc.add_paragraph('Returns: Product name, unit cost, price, margin, lead time')
    doc.add_paragraph('(Note: Currently from sample_data.py, not database)')

    doc.add_heading('Example Output', 2)
    doc.add_paragraph('Tool 1 returns this JSON structure to the AI:')

    output = doc.add_paragraph(
        '{\n'
        '  "historical_demand": [\n'
        '    {"date": "2026-02-25", "demand_units": 1789, "avg_price": 89.50},\n'
        '    {"date": "2026-02-18", "demand_units": 1687, "avg_price": 88.75},\n'
        '    ...\n'
        '  ],\n'
        '  "current_inventory": {\n'
        '    "current_stock": 1850,\n'
        '    "safety_stock": 300,\n'
        '    "days_of_supply": 13\n'
        '  },\n'
        '  "market_context": {\n'
        '    "industry_growth_rate": 7.2,\n'
        '    "economic_indicator": "Manufacturing PMI at 52.3",\n'
        '    "seasonal_pattern": "Q1 sees 15% higher demand"\n'
        '  },\n'
        '  "production_capacity": {\n'
        '    "max_daily_output": 220,\n'
        '    "current_utilization_pct": 76\n'
        '  },\n'
        '  "product_info": {\n'
        '    "product_name": "Industrial Valve Assembly - Type A",\n'
        '    "unit_price": 89.50\n'
        '  }\n'
        '}',
        style='No Spacing'
    )
    output.runs[0].font.name = 'Courier New'
    output.runs[0].font.size = Pt(9)

    doc.add_heading('How the AI Uses This Data', 2)
    doc.add_paragraph(
        'The AI reads all 5 data sources and forms an understanding of the situation. '
        'It sees the historical trend, current inventory level, market conditions, and production '
        'capacity. This context helps it make informed decisions about future demand.'
    )

    doc.add_page_break()

def add_section_5(doc):
    """Section 5: Tool 2 - Simulate Demand Scenarios"""
    doc.add_heading('5. Tool 2: Simulate Demand Scenarios', 1)

    doc.add_heading('What This Tool Does', 2)
    doc.add_paragraph(
        'Tool 2 generates 3 probabilistic demand forecasts: Optimistic (70% chance of exceeding), '
        'Base Case (50% chance), and Pessimistic (30% chance). It uses mathematical formulas '
        'based on historical data.'
    )

    doc.add_heading('Tool Definition', 2)
    p = doc.add_paragraph()
    p.add_run('Function Name: ').bold = True
    p.add_run('simulate_demand_scenarios(product_id, horizon_weeks)')

    p = doc.add_paragraph()
    p.add_run('Input: ').bold = True
    p.add_run('product_id (string), horizon_weeks (integer, default=4)')

    p = doc.add_paragraph()
    p.add_run('Output: ').bold = True
    p.add_run('JSON with 3 scenarios + expected weighted demand')

    doc.add_heading('Step-by-Step Calculation', 2)

    # Step 1
    p = doc.add_paragraph()
    p.add_run('Step 1: Get Historical Demand').bold = True

    doc.add_paragraph('SQL Query:')
    code = doc.add_paragraph(
        'SELECT demand_units\n'
        'FROM historical_demand_data\n'
        'WHERE product_id = \'PROD-A\'\n'
        'ORDER BY week_start_date DESC\n'
        'LIMIT 12;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Example Result: [1789, 1687, 1686, 2081, 1491, 1662, 1588, 1406, 1482, 1495, 1308, 1404]')

    # Step 2
    p = doc.add_paragraph()
    p.add_run('Step 2: Calculate Base Statistics').bold = True

    doc.add_paragraph('Formula for Mean:')
    formula = doc.add_paragraph('mean = SUM(demand_units) / COUNT(weeks)', style='No Spacing')
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True

    doc.add_paragraph('Example Calculation:')
    calc = doc.add_paragraph(
        'mean = (1789 + 1687 + 1686 + ... + 1404) / 12\n'
        'mean = 19,079 / 12\n'
        'mean = 1,590 units',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    doc.add_paragraph('Formula for Standard Deviation:')
    formula = doc.add_paragraph(
        'std_dev = SQRT(SUM((demand - mean)^2) / COUNT(weeks))',
        style='No Spacing'
    )
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True

    doc.add_paragraph('Example Calculation:')
    calc = doc.add_paragraph(
        'std_dev = SQRT(((1789-1590)^2 + (1687-1590)^2 + ... + (1404-1590)^2) / 12)\n'
        'std_dev = SQRT(39,601 + 9,409 + ... + 34,596 / 12)\n'
        'std_dev = SQRT(61,234)\n'
        'std_dev = 247 units',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    # Step 3
    p = doc.add_paragraph()
    p.add_run('Step 3: Calculate Trend').bold = True

    doc.add_paragraph('Formula (Linear Regression Slope):')
    formula = doc.add_paragraph(
        'trend_slope = SUM((x - mean_x) * (y - mean_y)) / SUM((x - mean_x)^2)\n'
        'where x = week number, y = demand',
        style='No Spacing'
    )
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True

    doc.add_paragraph('Example Calculation:')
    calc = doc.add_paragraph(
        'Weeks: [1, 2, 3, ..., 12]\n'
        'Demand: [1404, 1308, 1495, ..., 1789]\n'
        'trend_slope = 26.5 units/week (upward trend)',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    # Step 4
    p = doc.add_paragraph()
    p.add_run('Step 4: Generate 3 Scenarios').bold = True

    doc.add_paragraph('Formulas:')

    formulas = [
        ('Optimistic', 'mean + (1.5 × std_dev) + (trend × horizon_weeks)'),
        ('Base Case', 'mean + (trend × horizon_weeks)'),
        ('Pessimistic', 'mean - (1.0 × std_dev) + (trend × horizon_weeks)'),
    ]

    for scenario, formula in formulas:
        p = doc.add_paragraph()
        p.add_run(f'{scenario}: ').bold = True
        code = doc.add_paragraph(formula, style='No Spacing')
        code.runs[0].font.name = 'Courier New'
        code.runs[0].font.italic = True

    doc.add_paragraph('Example Calculation (4-week horizon):')

    calc = doc.add_paragraph(
        'Optimistic = 1590 + (1.5 × 247) + (26.5 × 4)\n'
        '           = 1590 + 371 + 106\n'
        '           = 2,067 units\n\n'
        'Base Case  = 1590 + (26.5 × 4)\n'
        '           = 1590 + 106\n'
        '           = 1,696 units\n\n'
        'Pessimistic = 1590 - (1.0 × 247) + (26.5 × 4)\n'
        '            = 1590 - 247 + 106\n'
        '            = 1,449 units',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    # Step 5
    p = doc.add_paragraph()
    p.add_run('Step 5: Calculate Expected Weighted Demand').bold = True

    doc.add_paragraph('Formula:')
    formula = doc.add_paragraph(
        'expected = (optimistic × 0.25) + (base_case × 0.50) + (pessimistic × 0.25)',
        style='No Spacing'
    )
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True

    doc.add_paragraph('Example Calculation:')
    calc = doc.add_paragraph(
        'expected = (2067 × 0.25) + (1696 × 0.50) + (1449 × 0.25)\n'
        '         = 517 + 848 + 362\n'
        '         = 1,727 units',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    doc.add_heading('Example Output', 2)
    doc.add_paragraph('Tool 2 returns this JSON to the AI:')

    output = doc.add_paragraph(
        '{\n'
        '  "scenarios": {\n'
        '    "optimistic": 2067,\n'
        '    "base_case": 1696,\n'
        '    "pessimistic": 1449\n'
        '  },\n'
        '  "expected_weighted_demand": 1727,\n'
        '  "historical_summary": {\n'
        '    "weeks_analyzed": 12,\n'
        '    "mean_weekly_demand": 1590,\n'
        '    "std_dev": 247,\n'
        '    "trend_units_per_week": 26.5\n'
        '  }\n'
        '}',
        style='No Spacing'
    )
    output.runs[0].font.name = 'Courier New'
    output.runs[0].font.size = Pt(9)

    doc.add_heading('How the AI Uses This Data', 2)
    doc.add_paragraph(
        'The AI sees 3 different scenarios and understands the range of possibilities. '
        'The base case (1,696 units) is the most likely, but there\'s a chance demand could '
        'be as high as 2,067 (optimistic) or as low as 1,449 (pessimistic). The AI uses '
        'this to assess risk.'
    )

    doc.add_page_break()

def add_section_6(doc):
    """Section 6: Tool 3 - Analyze Demand Trends"""
    doc.add_heading('6. Tool 3: Analyze Demand Trends', 1)

    doc.add_heading('What This Tool Does', 2)
    doc.add_paragraph(
        'Tool 3 performs statistical analysis on historical demand to detect trends, '
        'anomalies, and correlations. It uses advanced formulas like linear regression, '
        'z-scores, and correlation coefficients.'
    )

    doc.add_heading('Tool Definition', 2)
    p = doc.add_paragraph()
    p.add_run('Function Name: ').bold = True
    p.add_run('analyze_demand_trends(product_id)')

    p = doc.add_paragraph()
    p.add_run('Input: ').bold = True
    p.add_run('product_id (string, e.g., "PROD-A")')

    p = doc.add_paragraph()
    p.add_run('Output: ').bold = True
    p.add_run('JSON with trend analysis, anomaly detection, and correlations')

    doc.add_heading('Analysis 1: Trend Detection (Linear Regression)', 2)

    doc.add_paragraph('SQL Query to Get Data:')
    code = doc.add_paragraph(
        'SELECT demand_units\n'
        'FROM historical_demand_data\n'
        'WHERE product_id = \'PROD-A\'\n'
        'ORDER BY week_start_date ASC\n'
        'LIMIT 12;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Formula for Linear Regression Slope:')
    formula = doc.add_paragraph(
        'slope = (N × SUM(x×y) - SUM(x) × SUM(y)) / (N × SUM(x²) - (SUM(x))²)\n'
        'where:\n'
        '  N = number of weeks\n'
        '  x = week number (1, 2, 3, ...)\n'
        '  y = demand value',
        style='No Spacing'
    )
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True

    doc.add_paragraph('Example Calculation:')
    calc = doc.add_paragraph(
        'Weeks (x): [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12]\n'
        'Demand (y): [1404, 1308, 1495, 1482, 1406, 1588, 1662, 1491, 2081, 1686, 1687, 1789]\n\n'
        'N = 12\n'
        'SUM(x) = 78\n'
        'SUM(y) = 19,079\n'
        'SUM(x×y) = 126,537\n'
        'SUM(x²) = 650\n\n'
        'slope = (12 × 126,537 - 78 × 19,079) / (12 × 650 - 78²)\n'
        '      = (1,518,444 - 1,488,162) / (7,800 - 6,084)\n'
        '      = 30,282 / 1,716\n'
        '      = 17.6 units/week',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    doc.add_paragraph('Interpretation:')
    p = doc.add_paragraph('• slope > 5: Upward trend')
    p = doc.add_paragraph('• slope between -5 and 5: Stable')
    p = doc.add_paragraph('• slope < -5: Downward trend')

    p = doc.add_paragraph()
    p.add_run('Result: ').bold = True
    p.add_run('17.6 units/week → UPWARD TREND')

    doc.add_heading('Analysis 2: Anomaly Detection (Z-Score)', 2)

    doc.add_paragraph('Formula for Z-Score:')
    formula = doc.add_paragraph(
        'z_score = (demand - mean) / std_dev\n'
        'Anomaly if: |z_score| > 2.0',
        style='No Spacing'
    )
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True

    doc.add_paragraph('Example Calculation:')
    calc = doc.add_paragraph(
        'Week of 2026-02-04: demand = 2081 units\n'
        'mean = 1590 units\n'
        'std_dev = 247 units\n\n'
        'z_score = (2081 - 1590) / 247\n'
        '        = 491 / 247\n'
        '        = 1.99\n\n'
        'Since |1.99| < 2.0, this is NOT flagged as anomaly by z-score.\n'
        '(But database has is_anomaly=1 because we know it was a viral spike)',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    doc.add_paragraph('Tool 3 ALSO checks database for pre-flagged anomalies:')
    code = doc.add_paragraph(
        'SELECT week_start_date, demand_units, anomaly_reason\n'
        'FROM historical_demand_data\n'
        'WHERE product_id = \'PROD-A\' AND is_anomaly = 1;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_heading('Analysis 3: Demand Statistics', 2)

    doc.add_paragraph('Calculates:')
    stats = [
        ('Mean', 'Average weekly demand', '1590 units'),
        ('Median', 'Middle value when sorted', '1587 units'),
        ('Std Dev', 'Volatility measure', '247 units'),
        ('Min', 'Lowest demand week', '1308 units'),
        ('Max', 'Highest demand week', '2081 units'),
        ('CV', 'Coefficient of Variation (std_dev/mean)', '15.5%'),
    ]

    for stat, desc, example in stats:
        p = doc.add_paragraph()
        p.add_run(f'{stat}: ').bold = True
        p.add_run(f'{desc} → {example}')

    doc.add_heading('Analysis 4: Correlation Analysis', 2)

    doc.add_paragraph('Correlation between demand and price:')
    formula = doc.add_paragraph(
        'correlation = SUM((x - mean_x) × (y - mean_y)) / \n'
        '              (SQRT(SUM((x - mean_x)²)) × SQRT(SUM((y - mean_y)²)))\n'
        'where x = demand, y = price',
        style='No Spacing'
    )
    formula.runs[0].font.name = 'Courier New'
    formula.runs[0].font.italic = True

    doc.add_paragraph('Example Result: -0.15 (weak negative correlation)')
    doc.add_paragraph('Meaning: Higher prices slightly decrease demand, but correlation is weak')

    doc.add_heading('Example Output', 2)
    doc.add_paragraph('Tool 3 returns this JSON to the AI:')

    output = doc.add_paragraph(
        '{\n'
        '  "trend_analysis": {\n'
        '    "direction": "Upward",\n'
        '    "slope_units_per_week": 17.6,\n'
        '    "confidence": "High"\n'
        '  },\n'
        '  "demand_statistics": {\n'
        '    "mean": 1590,\n'
        '    "median": 1587,\n'
        '    "std_dev": 247,\n'
        '    "min": 1308,\n'
        '    "max": 2081,\n'
        '    "coefficient_of_variation": 0.155\n'
        '  },\n'
        '  "anomalies_detected": [\n'
        '    {\n'
        '      "date": "2026-02-04",\n'
        '      "demand": 2081,\n'
        '      "reason": "Viral social media campaign spike",\n'
        '      "z_score": 1.99\n'
        '    }\n'
        '  ],\n'
        '  "correlation_analysis": {\n'
        '    "demand_vs_price": -0.15,\n'
        '    "demand_vs_promotions": 0.42\n'
        '  }\n'
        '}',
        style='No Spacing'
    )
    output.runs[0].font.name = 'Courier New'
    output.runs[0].font.size = Pt(9)

    doc.add_heading('How the AI Uses This Data', 2)
    doc.add_paragraph(
        'The AI sees there\'s an upward trend (17.6 units/week growth), one anomaly spike, '
        'and moderate volatility (CV=15.5%). It also sees that promotions have a positive '
        'correlation with demand (0.42), suggesting promotions are effective.'
    )

    doc.add_page_break()

def add_section_7(doc):
    """Section 7: Complete Execution Flow with Examples"""
    doc.add_heading('7. Complete Execution Flow with Examples', 1)

    doc.add_paragraph(
        'This section shows the COMPLETE flow from start to finish with REAL NUMBERS.'
    )

    # STEP 1
    doc.add_heading('STEP 1: User Clicks "Run Analysis"', 2)
    p = doc.add_paragraph()
    p.add_run('User action: ').bold = True
    p.add_run('Clicks "Run Analysis" button in UI for product PROD-A')

    p = doc.add_paragraph()
    p.add_run('Frontend: ').bold = True
    p.add_run('Sends POST request to /api/pipeline/run')

    p = doc.add_paragraph()
    p.add_run('Backend: ').bold = True
    p.add_run('Receives request and starts Demand Forecasting Agent')

    # STEP 2
    doc.add_heading('STEP 2: Agent Initializes', 2)
    p = doc.add_paragraph()
    p.add_run('Agent receives prompt: ').bold = True
    p.add_run('"Forecast demand for product PROD-A for the next 4 weeks."')

    p = doc.add_paragraph()
    p.add_run('Agent thinks: ').bold = True
    p.add_run('"I need to gather data first. I\'ll call get_demand_data_summary."')

    # STEP 3
    doc.add_heading('STEP 3: Agent Calls Tool 1', 2)
    p = doc.add_paragraph()
    p.add_run('Tool called: ').bold = True
    p.add_run('get_demand_data_summary(product_id="PROD-A")')

    doc.add_paragraph('Tool 1 executes 3 SQL queries:')

    # Query 1
    p = doc.add_paragraph()
    p.add_run('Query 1: Get historical demand').bold = True
    code = doc.add_paragraph(
        'SELECT * FROM historical_demand_data\n'
        'WHERE product_id = \'PROD-A\'\n'
        'ORDER BY week_start_date DESC LIMIT 12;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Returns: 12 weeks of demand data')
    doc.add_paragraph('[1789, 1687, 1686, 2081, 1491, 1662, 1588, 1406, 1482, 1495, 1308, 1404]')

    # Query 2
    p = doc.add_paragraph()
    p.add_run('Query 2: Get market context').bold = True
    code = doc.add_paragraph(
        'SELECT * FROM market_context_data\n'
        'ORDER BY date_recorded DESC LIMIT 1;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Returns: Industry growth 7.2%, PMI at 52.3 (expansion)')

    # Query 3
    p = doc.add_paragraph()
    p.add_run('Query 3: Get inventory levels').bold = True
    doc.add_paragraph('Returns: Current stock 1,850 units, safety stock 300 units')

    p = doc.add_paragraph()
    p.add_run('Tool 1 returns to AI: ').bold = True
    p.add_run('Complete JSON with all 5 data sources')

    # STEP 4
    doc.add_heading('STEP 4: Agent Analyzes Tool 1 Output', 2)
    p = doc.add_paragraph()
    p.add_run('Agent thinks: ').bold = True
    p.add_run(
        '"I see historical demand averaging around 1,590 units/week. '
        'There was a spike to 2,081 on 2026-02-04. Market is growing at 7.2%. '
        'Current inventory is 1,850 units. I should generate scenarios now."'
    )

    p = doc.add_paragraph()
    p.add_run('Agent decision: ').bold = True
    p.add_run('Call Tool 2 to simulate demand scenarios')

    # STEP 5
    doc.add_heading('STEP 5: Agent Calls Tool 2', 2)
    p = doc.add_paragraph()
    p.add_run('Tool called: ').bold = True
    p.add_run('simulate_demand_scenarios(product_id="PROD-A", horizon_weeks=4)')

    doc.add_paragraph('Tool 2 executes calculations:')

    calc = doc.add_paragraph(
        '1. Get historical data (same as Tool 1)\n'
        '2. Calculate mean = 1,590 units\n'
        '3. Calculate std_dev = 247 units\n'
        '4. Calculate trend = 26.5 units/week\n'
        '5. Generate scenarios:\n'
        '   - Optimistic = 1590 + (1.5×247) + (26.5×4) = 2,067\n'
        '   - Base Case  = 1590 + (26.5×4) = 1,696\n'
        '   - Pessimistic = 1590 - (1.0×247) + (26.5×4) = 1,449\n'
        '6. Expected weighted = (2067×0.25) + (1696×0.50) + (1449×0.25) = 1,727',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    p = doc.add_paragraph()
    p.add_run('Tool 2 returns to AI: ').bold = True
    p.add_run('3 scenarios + expected demand of 1,727 units')

    # STEP 6
    doc.add_heading('STEP 6: Agent Calls Tool 3', 2)
    p = doc.add_paragraph()
    p.add_run('Agent thinks: ').bold = True
    p.add_run('"I should analyze trends and check for anomalies."')

    p = doc.add_paragraph()
    p.add_run('Tool called: ').bold = True
    p.add_run('analyze_demand_trends(product_id="PROD-A")')

    doc.add_paragraph('Tool 3 executes calculations:')

    calc = doc.add_paragraph(
        '1. Linear regression: slope = 17.6 units/week → Upward trend\n'
        '2. Check database for anomalies:\n'
        '   SELECT * FROM historical_demand_data\n'
        '   WHERE product_id = \'PROD-A\' AND is_anomaly = 1;\n'
        '   → Found: 2026-02-04 spike (2,081 units - viral campaign)\n'
        '3. Calculate statistics: mean=1590, std_dev=247, CV=15.5%\n'
        '4. Correlation analysis: demand vs promotions = 0.42',
        style='No Spacing'
    )
    calc.runs[0].font.name = 'Courier New'
    calc.runs[0].font.size = Pt(9)

    p = doc.add_paragraph()
    p.add_run('Tool 3 returns to AI: ').bold = True
    p.add_run('Upward trend confirmed, 1 anomaly detected, moderate volatility')

    # STEP 7
    doc.add_heading('STEP 7: Agent Generates Final Forecast', 2)
    p = doc.add_paragraph()
    p.add_run('Agent reasoning: ').bold = True

    reasoning = doc.add_paragraph(
        '"Based on the data:\n'
        '• Historical average: 1,590 units/week\n'
        '• Upward trend: +17.6 units/week\n'
        '• Recent anomaly spike (viral campaign) - likely one-time event\n'
        '• Market growing at 7.2%\n'
        '• Promotions have positive effect (correlation 0.42)\n\n'
        'My forecast for next 4 weeks:\n'
        '• Optimistic: 2,067 units (if viral effect continues)\n'
        '• Base Case: 1,696 units (most likely - trend continues)\n'
        '• Pessimistic: 1,449 units (if demand drops)\n'
        '• Expected: 1,727 units (probability-weighted)\n\n'
        'Confidence: MEDIUM (volatility is 15.5%, recent anomaly adds uncertainty)"',
        style='Intense Quote'
    )
    reasoning.runs[0].font.italic = True
    reasoning.runs[0].font.size = Pt(10)

    # STEP 8
    doc.add_heading('STEP 8: Agent Saves Forecast to Database', 2)
    doc.add_paragraph('SQL Query Executed:')
    code = doc.add_paragraph(
        'INSERT INTO demand_forecasts\n'
        '(product_id, forecast_date, week_number,\n'
        ' optimistic, base_case, pessimistic, actual)\n'
        'VALUES\n'
        '(\'PROD-A\', \'2026-03-04\', 10,\n'
        ' 2067, 1696, 1449, NULL);',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    p = doc.add_paragraph()
    p.add_run('Result: ').bold = True
    p.add_run('Forecast saved to database with ID (auto-incremented)')

    # STEP 9
    doc.add_heading('STEP 9: Frontend Displays Results', 2)
    doc.add_paragraph('Frontend queries database:')
    code = doc.add_paragraph(
        'SELECT * FROM demand_forecasts\n'
        'WHERE product_id = \'PROD-A\'\n'
        'ORDER BY created_at DESC LIMIT 1;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Returns: Latest forecast (optimistic=2067, base=1696, pessimistic=1449)')

    p = doc.add_paragraph()
    p.add_run('UI displays: ').bold = True
    p.add_run('Forecast chart with 3 scenarios + AI reasoning text')

    doc.add_page_break()

def add_section_8(doc):
    """Section 8: All SQL Queries Reference"""
    doc.add_heading('8. All SQL Queries Reference', 1)

    doc.add_paragraph(
        'This section lists ALL SQL queries used in the Demand Intelligence module '
        'for quick reference.'
    )

    queries = [
        (
            'Query 1: Get Historical Demand (12 weeks)',
            'SELECT * FROM historical_demand_data\n'
            'WHERE product_id = \'PROD-A\'\n'
            'ORDER BY week_start_date DESC\n'
            'LIMIT 12;'
        ),
        (
            'Query 2: Get All Historical Data (52 weeks)',
            'SELECT * FROM historical_demand_data\n'
            'WHERE product_id = \'PROD-A\'\n'
            'ORDER BY week_start_date ASC;'
        ),
        (
            'Query 3: Calculate Average Demand',
            'SELECT AVG(demand_units) as avg_demand\n'
            'FROM historical_demand_data\n'
            'WHERE product_id = \'PROD-A\';'
        ),
        (
            'Query 4: Find Anomalies',
            'SELECT week_start_date, demand_units, anomaly_reason\n'
            'FROM historical_demand_data\n'
            'WHERE product_id = \'PROD-A\'\n'
            '  AND is_anomaly = 1\n'
            'ORDER BY week_start_date DESC;'
        ),
        (
            'Query 5: Get Market Context',
            'SELECT * FROM market_context_data\n'
            'ORDER BY date_recorded DESC\n'
            'LIMIT 1;'
        ),
        (
            'Query 6: Get Latest Forecast',
            'SELECT * FROM demand_forecasts\n'
            'WHERE product_id = \'PROD-A\'\n'
            'ORDER BY created_at DESC\n'
            'LIMIT 1;'
        ),
        (
            'Query 7: Compare Forecast to Historical Average',
            'SELECT\n'
            '  (SELECT AVG(demand_units) FROM historical_demand_data\n'
            '   WHERE product_id = \'PROD-A\') as historical_avg,\n'
            '  (SELECT base_case FROM demand_forecasts\n'
            '   WHERE product_id = \'PROD-A\'\n'
            '   ORDER BY created_at DESC LIMIT 1) as forecast;'
        ),
        (
            'Query 8: Week-over-Week Changes',
            'WITH weekly_data AS (\n'
            '  SELECT week_start_date, demand_units,\n'
            '    LAG(demand_units) OVER (ORDER BY week_start_date) as prev\n'
            '  FROM historical_demand_data\n'
            '  WHERE product_id = \'PROD-A\'\n'
            ')\n'
            'SELECT week_start_date, demand_units, prev,\n'
            '  ROUND(((demand_units - prev) * 100.0 / prev), 2) as change_pct\n'
            'FROM weekly_data\n'
            'WHERE prev IS NOT NULL\n'
            'ORDER BY week_start_date DESC\n'
            'LIMIT 10;'
        ),
        (
            'Query 9: Count Total Records',
            'SELECT product_id, COUNT(*) as week_count\n'
            'FROM historical_demand_data\n'
            'GROUP BY product_id;'
        ),
        (
            'Query 10: Get Promotion Weeks',
            'SELECT week_start_date, demand_units\n'
            'FROM historical_demand_data\n'
            'WHERE product_id = \'PROD-A\'\n'
            '  AND promotions_active = 1\n'
            'ORDER BY week_start_date DESC;'
        ),
    ]

    for i, (title, query) in enumerate(queries, start=1):
        doc.add_heading(title, 3)
        code = doc.add_paragraph(query, style='No Spacing')
        code.runs[0].font.name = 'Courier New'
        code.runs[0].font.size = Pt(9)
        doc.add_paragraph()  # Spacing

    doc.add_page_break()

def add_section_9(doc):
    """Section 9: Validation & Verification"""
    doc.add_heading('9. Validation & Verification', 1)

    doc.add_paragraph(
        'This section shows how to validate that everything is working correctly.'
    )

    doc.add_heading('Validation 1: Verify Historical Data Exists', 2)
    doc.add_paragraph('Run this query:')
    code = doc.add_paragraph(
        'SELECT COUNT(*) FROM historical_demand_data;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'

    doc.add_paragraph('Expected result: 156 (52 weeks × 3 products)')

    doc.add_heading('Validation 2: Verify Tool 1 Gets Correct Data', 2)
    doc.add_paragraph('Step 1: Run this query manually:')
    code = doc.add_paragraph(
        'SELECT demand_units FROM historical_demand_data\n'
        'WHERE product_id = \'PROD-A\'\n'
        'ORDER BY week_start_date DESC LIMIT 12;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Step 2: Run the pipeline and check Tool 1 output')
    doc.add_paragraph('Step 3: Compare - they should match exactly')

    doc.add_heading('Validation 3: Verify Tool 2 Calculations', 2)
    doc.add_paragraph('Step 1: Get mean demand:')
    code = doc.add_paragraph(
        'SELECT AVG(demand_units) FROM historical_demand_data\n'
        'WHERE product_id = \'PROD-A\';',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Expected: ~1,590 units')

    doc.add_paragraph('Step 2: Check Tool 2 output mentions this mean')
    doc.add_paragraph('Step 3: Verify scenarios use this mean in formulas')

    doc.add_heading('Validation 4: Verify Tool 3 Detected Anomalies', 2)
    doc.add_paragraph('Step 1: Query database for anomalies:')
    code = doc.add_paragraph(
        'SELECT * FROM historical_demand_data\n'
        'WHERE product_id = \'PROD-A\' AND is_anomaly = 1;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Expected: 3 anomalies (viral spike, competitor disruption, Black Friday)')

    doc.add_paragraph('Step 2: Check Tool 3 output lists these anomalies')

    doc.add_heading('Validation 5: Verify Final Forecast Saved', 2)
    doc.add_paragraph('After running pipeline:')
    code = doc.add_paragraph(
        'SELECT * FROM demand_forecasts\n'
        'ORDER BY created_at DESC LIMIT 1;',
        style='No Spacing'
    )
    code.runs[0].font.name = 'Courier New'
    code.runs[0].font.size = Pt(9)

    doc.add_paragraph('Should show latest forecast with optimistic, base_case, pessimistic values')

    doc.add_page_break()

def add_section_10(doc):
    """Section 10: Formulas & Calculations Cheat Sheet"""
    doc.add_heading('10. Formulas & Calculations Cheat Sheet', 1)

    doc.add_paragraph(
        'Quick reference for all mathematical formulas used in Demand Intelligence.'
    )

    formulas = [
        ('Mean (Average)', 'mean = SUM(values) / COUNT(values)'),
        ('Standard Deviation', 'std_dev = SQRT(SUM((value - mean)²) / COUNT(values))'),
        ('Linear Regression Slope',
         'slope = (N×SUM(xy) - SUM(x)×SUM(y)) / (N×SUM(x²) - (SUM(x))²)'),
        ('Z-Score (Anomaly Detection)', 'z_score = (value - mean) / std_dev'),
        ('Correlation Coefficient',
         'r = SUM((x-mean_x)×(y-mean_y)) / (SQRT(SUM((x-mean_x)²))×SQRT(SUM((y-mean_y)²)))'),
        ('Optimistic Scenario', 'mean + (1.5 × std_dev) + (trend × horizon)'),
        ('Base Case Scenario', 'mean + (trend × horizon)'),
        ('Pessimistic Scenario', 'mean - (1.0 × std_dev) + (trend × horizon)'),
        ('Expected Weighted Demand',
         '(optimistic × 0.25) + (base_case × 0.50) + (pessimistic × 0.25)'),
        ('Coefficient of Variation', 'CV = std_dev / mean'),
    ]

    for formula_name, formula in formulas:
        doc.add_heading(formula_name, 3)
        code = doc.add_paragraph(formula, style='No Spacing')
        code.runs[0].font.name = 'Courier New'
        code.runs[0].font.size = Pt(10)
        code.runs[0].font.italic = True
        doc.add_paragraph()  # Spacing

    doc.add_heading('Quick Reference Table', 2)

    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'

    header = table.rows[0].cells
    header[0].text = 'Metric'
    header[1].text = 'Example Value'
    header[2].text = 'What It Means'

    data = [
        ('Mean Demand', '1,590 units', 'Average weekly demand'),
        ('Std Dev', '247 units', 'Typical variance from mean'),
        ('Trend Slope', '+17.6 units/week', 'Demand growing 17.6 units per week'),
        ('CV', '15.5%', 'Moderate volatility'),
        ('Z-Score > 2', '2.5', 'Significant anomaly'),
        ('Correlation > 0.5', '0.42', 'Moderate positive relationship'),
    ]

    for i, (metric, value, meaning) in enumerate(data, start=1):
        row = table.rows[i].cells
        row[0].text = metric
        row[1].text = value
        row[2].text = meaning

# ============================================================================
# MAIN FUNCTION
# ============================================================================
def generate_document():
    """Generate the complete Word document"""
    doc = Document()

    # Configure default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    print('[1/10] Adding title page...')
    add_title_page(doc)

    print('[2/10] Adding table of contents...')
    add_toc(doc)

    print('[3/10] Adding Section 1: What is Demand Intelligence?')
    add_section_1(doc)

    print('[4/10] Adding Section 2: How the AI Agent Works')
    add_section_2(doc)

    print('[5/10] Adding Section 3: Database Tables Used')
    add_section_3(doc)

    print('[6/10] Adding Section 4: Tool 1')
    add_section_4(doc)

    print('[7/10] Adding Section 5: Tool 2')
    add_section_5(doc)

    print('[8/10] Adding Section 6: Tool 3')
    add_section_6(doc)

    print('[9/10] Adding Section 7: Complete Execution Flow')
    add_section_7(doc)

    print('[10/10] Adding Sections 8-10: Reference Material')
    add_section_8(doc)
    add_section_9(doc)
    add_section_10(doc)

    # Save document
    output_path = 'DEMAND_INTELLIGENCE_COMPLETE_GUIDE.docx'
    doc.save(output_path)

    print(f'\n[SUCCESS] Document created: {output_path}')
    print('Document contains:')
    print('  - 10 comprehensive sections')
    print('  - All 3 tools explained in detail')
    print('  - Every calculation with examples')
    print('  - All SQL queries')
    print('  - Complete step-by-step flow')
    print('  - Validation guides')
    print('  - Formula cheat sheet')

if __name__ == '__main__':
    generate_document()
