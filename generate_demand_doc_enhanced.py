"""
Enhanced Word Document Generator for Demand Intelligence Module
Includes detailed "What Each Tool Does" section
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_number(paragraph):
    """Add page numbers to footer"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def create_enhanced_doc():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ============================================================
    # TITLE PAGE
    # ============================================================
    title = doc.add_heading('DEMAND INTELLIGENCE MODULE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Complete Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.bold = True

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('AMIS - Autonomous Manufacturing Intelligence System\n').font.size = Pt(12)
    info.add_run('Version 1.0 | March 2026\n').font.size = Pt(11)
    info.add_run('Complete Tool-by-Tool Explanation').font.size = Pt(11)

    doc.add_page_break()

    # ============================================================
    # NEW SECTION: DEMAND AGENT TOOL EXECUTION (DETAILED)
    # ============================================================
    doc.add_heading('DEMAND FORECASTING AGENT - TOOL EXECUTION FLOW', 1)

    p = doc.add_paragraph(
        'This section explains EXACTLY what happens when the Demand Forecasting Agent runs. '
        'We break down each tool call, what it does, and what it returns.'
    )

    # Tool Execution Overview
    doc.add_heading('Overview: 3 Tools Called in Sequence', 2)

    p = doc.add_paragraph()
    p.add_run('When the pipeline runs, the Demand Forecasting Agent calls ').bold = False
    p.add_run('3 main tools').bold = True
    p.add_run(' in this order:')

    doc.add_paragraph('TOOL 1: get_demand_data_summary() - Gathers all context data', style='List Number')
    doc.add_paragraph('TOOL 2: simulate_demand_scenarios() - Generates 3 forecast scenarios', style='List Number')
    doc.add_paragraph('TOOL 3: analyze_demand_trends() - Performs statistical trend analysis', style='List Number')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Think of it like a detective investigating a case:').italic = True
    doc.add_paragraph('Tool 1: Gather all evidence (historical data, market context)', style='List Bullet')
    doc.add_paragraph('Tool 2: Build 3 possible theories (optimistic, base, pessimistic scenarios)', style='List Bullet')
    doc.add_paragraph('Tool 3: Analyze patterns and trends (statistical analysis)', style='List Bullet')

    doc.add_page_break()

    # ============================================================
    # TOOL 1: GET DEMAND DATA SUMMARY
    # ============================================================
    doc.add_heading('TOOL 1: get_demand_data_summary()', 1)

    # What it does
    doc.add_heading('What This Tool Does', 2)
    p = doc.add_paragraph(
        'This is the FIRST tool called. It gathers all the context data the AI agent needs '
        'to make intelligent forecasts. Think of it as "getting the full picture before making decisions."'
    )

    # Inputs
    doc.add_heading('Inputs', 2)
    table1 = doc.add_table(rows=2, cols=3)
    table1.style = 'Light Grid Accent 1'

    hdr = table1.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Type'
    hdr[2].text = 'Example Value'
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True

    row1 = table1.rows[1].cells
    row1[0].text = 'product_id'
    row1[1].text = 'string'
    row1[2].text = '"PROD-A"'

    # What it does step by step
    doc.add_heading('What Happens Inside (Step-by-Step)', 2)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Step 1: ').bold = True
    p.add_run('Load Historical Demand (12 weeks)')

    code1 = """Calls: get_historical_demand("PROD-A", weeks=12)
Returns: 12 weeks of past demand data
Example:
  Week 1: 920 units
  Week 2: 985 units
  Week 3: 1450 units (spike!)
  ...
  Week 12: 1240 units"""

    p = doc.add_paragraph(code1)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Step 2: ').bold = True
    p.add_run('Get Current Inventory Status')

    code2 = """Calls: get_current_inventory("PROD-A")
Returns:
  - Current stock: 1,850 units
  - Safety stock: 300 units
  - Days of supply: 13 days
  - Incoming orders: 800 units (arriving in 4-7 days)"""

    p = doc.add_paragraph(code2)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Step 3: ').bold = True
    p.add_run('Get Market Context (External Factors)')

    code3 = """Calls: get_market_context()
Returns:
  - Season: Q1 2026
  - Industry trend: Growing at 4.2% per year
  - Social media: 2,400 mentions (up 180%!)
  - Sentiment: 78% positive
  - Top topic: "Viral TikTok video featuring our product"
  - Upcoming: Trade show in 3 weeks, client contract renewal"""

    p = doc.add_paragraph(code3)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Step 4: ').bold = True
    p.add_run('Get Production Capacity')

    code4 = """Calls: get_production_capacity()
Returns:
  - Max daily output: 220 units/day
  - Current utilization: 76%
  - Available lines: 4 out of 5 (Line 4 under maintenance)
  - Overtime available: Yes (35% cost premium)
  - Contract manufacturer: Available (42% cost premium)"""

    p = doc.add_paragraph(code4)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Step 5: ').bold = True
    p.add_run('Get Product Information')

    code5 = """Calls: get_product_info("PROD-A")
Returns:
  - Product name: "Automotive Sensor Unit"
  - Unit cost: $52.00
  - Unit price: $89.50
  - Margin: 41.9%
  - Lead time: 5 days"""

    p = doc.add_paragraph(code5)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # Output
    doc.add_heading('Output (Combined JSON)', 2)

    p = doc.add_paragraph(
        'All 5 pieces of data are combined into one big JSON object:'
    )

    output1 = """{
  "product_info": {...},
  "historical_demand_last_12_weeks": [...],
  "current_inventory": {...},
  "market_context": {...},
  "production_capacity": {...}
}"""

    p = doc.add_paragraph(output1)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 100, 0)

    # Why it matters
    doc.add_heading('Why This Tool Matters', 2)

    doc.add_paragraph(
        'Context is EVERYTHING. The AI agent needs to know:',
    )
    doc.add_paragraph('Historical patterns: Is demand growing or stable?', style='List Bullet')
    doc.add_paragraph('Current inventory: Do we have enough stock?', style='List Bullet')
    doc.add_paragraph('Market context: Why did demand spike (promotion? viral video?)', style='List Bullet')
    doc.add_paragraph('Production capacity: Can we even produce more if needed?', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Without this context, the AI would be "forecasting blind."').italic = True

    doc.add_page_break()

    # ============================================================
    # TOOL 2: SIMULATE DEMAND SCENARIOS
    # ============================================================
    doc.add_heading('TOOL 2: simulate_demand_scenarios()', 1)

    doc.add_heading('What This Tool Does', 2)
    p = doc.add_paragraph(
        'This tool generates 3 different demand forecast scenarios (Optimistic, Base, Pessimistic) '
        'with probability weights. It uses mathematical formulas to project future demand based on '
        'historical patterns, trends, and market context.'
    )

    # Inputs
    doc.add_heading('Inputs', 2)
    table2 = doc.add_table(rows=3, cols=3)
    table2.style = 'Light Grid Accent 1'

    hdr = table2.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Type'
    hdr[2].text = 'Example Value'
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True

    row1 = table2.rows[1].cells
    row1[0].text = 'product_id'
    row1[1].text = 'string'
    row1[2].text = '"PROD-A"'

    row2 = table2.rows[2].cells
    row2[0].text = 'horizon_weeks'
    row2[1].text = 'integer'
    row2[2].text = '4 (forecast 4 weeks ahead)'

    # Step by step algorithm
    doc.add_heading('Algorithm (Step-by-Step Calculation)', 2)

    p = doc.add_paragraph()
    p.add_run('STEP 1: Calculate Historical Statistics').bold = True

    calc1 = """From historical data (12 weeks):
  demands = [920, 985, 1450, 1050, 1080, 1100, 1120, 1150, 1180, 1200, 1220, 1240]

  Calculate:
  - Average demand = sum(demands) / 12 = 1137.9 units/week
  - Recent 4-week avg = (1180 + 1200 + 1220 + 1240) / 4 = 1210 units/week
  - Trend % = (1210 - 1137.9) / 1137.9 * 100 = +6.3% growth
  - Standard deviation = 127.5 units (volatility measure)"""

    p = doc.add_paragraph(calc1)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('STEP 2: Detect Anomalies').bold = True

    calc2 = """Check if recent week was unusually high:
  last_week = 1240 units
  spike_threshold = average * 1.3 = 1137.9 * 1.3 = 1479 units
  spike_detected = 1240 > 1479? NO (no anomaly)

  Note: Week 3 had spike (1450 units) - had promotion that week"""

    p = doc.add_paragraph(calc2)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('STEP 3: Generate OPTIMISTIC Scenario (20% probability)').bold = True

    calc3 = """Formula: recent_avg * (1.15 + 0.03 * week) + random_noise

  Week 1: 1210 * 1.15 + noise = 1391 units
  Week 2: 1210 * 1.18 + noise = 1427 units
  Week 3: 1210 * 1.21 + noise = 1464 units
  Week 4: 1210 * 1.24 + noise = 1500 units

  Total: 5,782 units over 4 weeks

  Assumptions:
  - Viral social media momentum continues (+15% boost)
  - Trade show generates new enterprise contracts
  - Competitor supply issues persist"""

    p = doc.add_paragraph(calc3)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('STEP 4: Generate BASE Scenario (55% probability - most likely)').bold = True

    calc4 = """Formula: recent_avg * (1.0 + 0.01 * week) + random_noise

  Week 1: 1210 * 1.00 + noise = 1210 units
  Week 2: 1210 * 1.01 + noise = 1222 units
  Week 3: 1210 * 1.02 + noise = 1234 units
  Week 4: 1210 * 1.03 + noise = 1246 units

  Total: 4,912 units over 4 weeks

  Assumptions:
  - Demand returns to normal trend after temporary spike
  - Normal seasonal patterns hold
  - No major market disruptions"""

    p = doc.add_paragraph(calc4)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('STEP 5: Generate PESSIMISTIC Scenario (25% probability)').bold = True

    calc5 = """Formula: recent_avg * (0.82 - 0.02 * week) + random_noise

  Week 1: 1210 * 0.82 + noise = 992 units
  Week 2: 1210 * 0.80 + noise = 968 units
  Week 3: 1210 * 0.78 + noise = 943 units
  Week 4: 1210 * 0.76 + noise = 919 units

  Total: 3,822 units over 4 weeks

  Assumptions:
  - Raw material price increases passed to customers
  - Competitor's new product captures market share
  - Economic slowdown reduces enterprise orders"""

    p = doc.add_paragraph(calc5)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('STEP 6: Calculate Expected Weighted Demand').bold = True

    calc6 = """Formula: (Optimistic * 0.20) + (Base * 0.55) + (Pessimistic * 0.25)

  = (5782 * 0.20) + (4912 * 0.55) + (3822 * 0.25)
  = 1156 + 2702 + 955
  = 4813 units over 4 weeks
  = 1203 units/week average

  This is the "most likely" overall forecast accounting for all scenarios"""

    p = doc.add_paragraph(calc6)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # Output
    doc.add_heading('Output (Scenario JSON)', 2)

    output2 = """{
  "scenarios": {
    "optimistic": {
      "probability_weight": 0.20,
      "weekly_forecast": [1391, 1427, 1464, 1500],
      "total_units": 5782
    },
    "base": {
      "probability_weight": 0.55,
      "weekly_forecast": [1210, 1222, 1234, 1246],
      "total_units": 4912
    },
    "pessimistic": {
      "probability_weight": 0.25,
      "weekly_forecast": [992, 968, 943, 919],
      "total_units": 3822
    }
  },
  "expected_weighted_demand": 4813,
  "confidence_interval_95pct": {
    "lower": 907,
    "upper": 1512
  }
}"""

    p = doc.add_paragraph(output2)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 100, 0)

    # Why it matters
    doc.add_heading('Why This Tool Matters', 2)

    doc.add_paragraph(
        'This tool gives us UNCERTAINTY QUANTIFICATION - not just one number, but a range of possibilities:',
    )
    doc.add_paragraph('20% chance demand is VERY HIGH (optimistic) - prepare for extra capacity', style='List Bullet')
    doc.add_paragraph('55% chance demand is NORMAL (base case) - our default planning target', style='List Bullet')
    doc.add_paragraph('25% chance demand is LOW (pessimistic) - don\'t overproduce and waste money', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Better than single-point forecast: ').italic = True
    p.add_run('We know the risks and can plan accordingly.')

    doc.add_page_break()

    # ============================================================
    # TOOL 3: ANALYZE DEMAND TRENDS
    # ============================================================
    doc.add_heading('TOOL 3: analyze_demand_trends()', 1)

    doc.add_heading('What This Tool Does', 2)
    p = doc.add_paragraph(
        'This tool performs deep statistical analysis on historical demand to understand patterns, '
        'detect anomalies, and calculate correlations. It answers: "WHY is demand changing?"'
    )

    # Inputs
    doc.add_heading('Inputs', 2)
    table3 = doc.add_table(rows=2, cols=3)
    table3.style = 'Light Grid Accent 1'

    hdr = table3.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Type'
    hdr[2].text = 'Example Value'
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True

    row1 = table3.rows[1].cells
    row1[0].text = 'product_id'
    row1[1].text = 'string'
    row1[2].text = '"PROD-A"'

    # Step by step
    doc.add_heading('Analysis Steps', 2)

    p = doc.add_paragraph()
    p.add_run('STEP 1: Linear Regression (Find Trend Slope)').bold = True

    trend1 = """Uses least-squares method to find best-fit line:

  Formula: slope = Σ((x - x_mean) * (y - y_mean)) / Σ((x - x_mean)²)

  Where:
    x = week number (0, 1, 2, ..., 11)
    y = demand values

  Calculation:
    x_mean = 5.5 (midpoint of 0-11)
    y_mean = 1138 (average demand)

    numerator = sum of (week - 5.5) * (demand - 1138) = 3382.5
    denominator = sum of (week - 5.5)² = 143

    slope = 3382.5 / 143 = 23.65 units/week

  Interpretation:
    Demand is growing by 23.65 units per week
    Growth rate = (23.65 / 1138) * 100 = 2.08% per week"""

    p = doc.add_paragraph(trend1)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('STEP 2: Week-over-Week Changes').bold = True

    trend2 = """Calculate percentage change from one week to next:

  Week 1 to 2: (985 - 920) / 920 * 100 = +7.1%
  Week 2 to 3: (1450 - 985) / 985 * 100 = +47.2% (SPIKE!)
  Week 3 to 4: (1050 - 1450) / 1450 * 100 = -27.6% (spike ends)
  Week 4 to 5: (1080 - 1050) / 1050 * 100 = +2.9%
  ...

  Result: [7.1%, 47.2%, -27.6%, 2.9%, 1.9%, 1.8%, ...]"""

    p = doc.add_paragraph(trend2)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('STEP 3: Detect Anomalies (Z-Score Method)').bold = True

    trend3 = """Find weeks where demand is unusually high or low:

  Formula: z_score = (demand - average) / std_deviation

  Week 3 example:
    demand = 1450 units
    average = 1138 units
    std_dev = 127.5 units

    z_score = (1450 - 1138) / 127.5 = 2.45

  Interpretation:
    z > 1.5 = Anomaly detected!
    z = 2.45 means this week was 2.45 standard deviations above normal
    This happens in only 0.8% of normal weeks (very rare)

  Additional context:
    Week 3 had promotion = TRUE
    Conclusion: Spike explained by promotion (not organic growth)"""

    p = doc.add_paragraph(trend3)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('STEP 4: Price-Demand Correlation').bold = True

    trend4 = """Check if price changes affect demand:

  Formula: correlation = covariance / (std_dev_demand * std_dev_price)

  Result: -0.12

  Interpretation:
    -1.0 = Strong negative (higher price = much lower demand)
    0.0 = No correlation
    +1.0 = Strong positive (higher price = higher demand?!)

    -0.12 = Weak negative correlation
    Meaning: Price changes don't significantly affect demand
    This product might be essential/required (not price-sensitive)"""

    p = doc.add_paragraph(trend4)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    # Output
    doc.add_heading('Output (Trend Analysis JSON)', 2)

    output3 = """{
  "trend_analysis": {
    "direction": "Upward",
    "slope_units_per_week": 23.65,
    "growth_rate_pct_per_week": 2.08
  },
  "demand_statistics": {
    "mean": 1138,
    "std_deviation": 127,
    "min": 920,
    "max": 1450,
    "coefficient_of_variation": 11.2
  },
  "anomalies_detected": [
    {
      "week": "2026-W03",
      "demand": 1450,
      "z_score": 2.45,
      "had_promotion": true
    }
  ],
  "price_demand_correlation": {
    "correlation_coefficient": -0.12,
    "interpretation": "Weak/no correlation"
  }
}"""

    p = doc.add_paragraph(output3)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 100, 0)

    # Why it matters
    doc.add_heading('Why This Tool Matters', 2)

    doc.add_paragraph(
        'This tool provides the "WHY" behind the numbers:',
    )
    doc.add_paragraph('Trend direction: Are we growing or declining? (Upward at 2.08%/week)', style='List Bullet')
    doc.add_paragraph('Anomaly detection: Which weeks were unusual? (Week 3 spike due to promotion)', style='List Bullet')
    doc.add_paragraph('Price sensitivity: Can we raise prices without losing sales? (Weak correlation = yes)', style='List Bullet')
    doc.add_paragraph('Volatility: How predictable is demand? (CV=11.2% = moderate, manageable)', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('The AI agent uses this to EXPLAIN its recommendations, not just state numbers.').italic = True

    doc.add_page_break()

    # ============================================================
    # SUMMARY: ALL 3 TOOLS TOGETHER
    # ============================================================
    doc.add_heading('SUMMARY: How All 3 Tools Work Together', 1)

    doc.add_heading('The Complete Picture', 2)

    p = doc.add_paragraph()
    p.add_run('Tool 1 (get_demand_data_summary)').bold = True
    p.add_run(' provides CONTEXT')
    doc.add_paragraph('Historical demand, inventory, market events, production capacity', style='List Bullet')
    doc.add_paragraph('Answer: "What data do we have?"', style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tool 2 (simulate_demand_scenarios)').bold = True
    p.add_run(' provides FORECASTS')
    doc.add_paragraph('3 scenarios with probabilities (optimistic, base, pessimistic)', style='List Bullet')
    doc.add_paragraph('Answer: "What WILL happen?" (with uncertainty)', style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tool 3 (analyze_demand_trends)').bold = True
    p.add_run(' provides EXPLANATION')
    doc.add_paragraph('Trend analysis, anomaly detection, correlations', style='List Bullet')
    doc.add_paragraph('Answer: "WHY is it happening?"', style='List Bullet')

    doc.add_heading('Example: How AI Agent Uses All 3 Tools', 2)

    example_box = """SCENARIO: User asks "What's the demand outlook for PROD-A?"

AGENT PROCESS:

1. Calls Tool 1 (get_demand_data_summary)
   Result: Gets 12 weeks of history, sees recent spike to 1450 units

2. Calls Tool 2 (simulate_demand_scenarios)
   Result: Generates 3 forecasts (4813 expected weighted demand)

3. Calls Tool 3 (analyze_demand_trends)
   Result: Detects anomaly in Week 3 (z-score 2.45, had promotion)

AI AGENT REASONING:
"I see demand spiked in Week 3 (1450 units, +47% wow). However, Tool 3
shows this was an anomaly (z-score 2.45) AND that week had a promotion.
The underlying trend is still upward at 2.08%/week (Tool 3 regression).

My base case forecast (Tool 2) of 1228 units/week assumes spike normalizes
but trend continues. I'm confident because:
- Tool 1 shows social media mentions up 180% (viral momentum)
- Tool 3 shows price correlation is weak (demand not price-sensitive)
- Tool 2 weighted average (4813 units) accounts for all scenarios

RECOMMENDATION: Plan for base case (1228/week) + 10% buffer for safety."

This is AGENTIC AI - the agent REASONS across multiple tools, not just
runs them mechanically."""

    p = doc.add_paragraph(example_box)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_heading('What Happens to Tool Results?', 2)

    table4 = doc.add_table(rows=4, cols=3)
    table4.style = 'Light Grid Accent 1'

    hdr = table4.rows[0].cells
    hdr[0].text = 'Tool Output'
    hdr[1].text = 'Sent To'
    hdr[2].text = 'Used For'
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True

    results = [
        ('All 3 tools combined', 'AI Agent (Claude LLM)', 'Generate human-readable analysis with recommendations'),
        ('Tool 2 scenarios', 'Database (demand_forecasts table)', 'Store optimistic/base/pessimistic forecasts for UI charts'),
        ('Tool 3 trends', 'Frontend UI', 'Display trend direction, growth rate % in metric cards'),
    ]

    for i, (output, sent_to, used_for) in enumerate(results, start=1):
        row = table4.rows[i].cells
        row[0].text = output
        row[1].text = sent_to
        row[2].text = used_for

    doc.add_page_break()

    # ============================================================
    # CONCLUSION
    # ============================================================
    doc.add_heading('CONCLUSION', 1)

    p = doc.add_paragraph(
        'You now understand EXACTLY what each tool does when the Demand Forecasting Agent runs:'
    )

    doc.add_paragraph('Tool 1 gathers context (historical data, inventory, market, capacity)', style='List Bullet')
    doc.add_paragraph('Tool 2 generates 3 probabilistic forecasts with mathematical formulas', style='List Bullet')
    doc.add_paragraph('Tool 3 analyzes trends, detects anomalies, finds correlations', style='List Bullet')
    doc.add_paragraph('AI agent combines all 3 to make intelligent, explainable recommendations', style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Key Takeaway: ').bold = True
    p.add_run(
        'This is not "AI magic" - every calculation is traceable, every formula is documented, '
        'and every recommendation has clear reasoning. The tools are ALGORITHMS, the agent is the INTELLIGENCE.'
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    p = doc.add_paragraph()
    p.add_run('Document Version: ').bold = True
    p.add_run('1.0 Enhanced\n')
    p.add_run('Last Updated: ').bold = True
    p.add_run('March 4, 2026\n')
    p.add_run('Focus: ').bold = True
    p.add_run('Complete Tool-by-Tool Explanation')

    # Save
    output_path = 'DEMAND_INTELLIGENCE_WITH_TOOL_DETAILS.docx'
    doc.save(output_path)
    print(f"[OK] Enhanced Word document created: {output_path}")
    print("[OK] Includes detailed 'What Each Tool Does' section")
    return output_path

if __name__ == "__main__":
    create_enhanced_doc()
