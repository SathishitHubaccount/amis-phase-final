from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colours
C_DARK_BLUE = RGBColor(0x1F, 0x35, 0x64)
C_MID_BLUE  = RGBColor(0x2E, 0x74, 0xB5)
C_ACCENT    = RGBColor(0x17, 0x75, 0xD2)
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_BODY      = RGBColor(0x1A, 0x1A, 0x1A)

# ── Set cell background
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Add horizontal rule
def add_hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '2E74B5')
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

# ── Configure heading styles
def cfg_style(styles, name, size, bold, color, sb=6, sa=3):
    try:
        s = styles[name]
        s.font.size       = Pt(size)
        s.font.bold       = bold
        s.font.color.rgb  = color
        s.paragraph_format.space_before = Pt(sb)
        s.paragraph_format.space_after  = Pt(sa)
    except Exception:
        pass

styles = doc.styles
cfg_style(styles, 'Heading 1', 20, True,  C_DARK_BLUE, 14, 6)
cfg_style(styles, 'Heading 2', 15, True,  C_MID_BLUE,  12, 4)
cfg_style(styles, 'Heading 3', 12, True,  C_MID_BLUE,   8, 3)
cfg_style(styles, 'Heading 4', 11, True,  C_ACCENT,     6, 2)
cfg_style(styles, 'Normal',    10, False, C_BODY,        2, 4)

# ── Title page
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('AMIS')
r.font.size      = Pt(36)
r.font.bold      = True
r.font.color.rgb = C_DARK_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Autonomous Manufacturing Intelligence System')
r2.font.size      = Pt(16)
r2.font.color.rgb = C_MID_BLUE

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = sub2.add_run('Real-World Implementation Blueprint')
r3.font.size      = Pt(14)
r3.font.bold      = True
r3.font.color.rgb = C_DARK_BLUE

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ['Version: 1.0', 'Date: February 2026',
             'Audience: Developers, Architects, Engineering Leads']:
    rm = meta.add_run(line + '\n')
    rm.font.size      = Pt(10)
    rm.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

add_hr(doc)
doc.add_paragraph()

# ── Inline formatting (bold + inline code)
def apply_inline(para, text):
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    for part in pattern.split(text):
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            r.font.name      = 'Courier New'
            r.font.size      = Pt(9)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            if part:
                para.add_run(part)

# ── Read markdown
md_path = r'c:\Users\user\Downloads\amis_phase1__data_flow_crises_output\amis_phase_final\AMIS_RealWorld_Blueprint.md'
with open(md_path, encoding='utf-8') as f:
    lines = f.readlines()

in_code  = False
code_buf = []
tbl_rows = []

# Skip the document title block (first ~6 lines we already rendered)
start_index = 0
for idx, l in enumerate(lines):
    if l.strip().startswith('## Table of Contents'):
        start_index = idx
        break

lines = lines[start_index:]

i = 0
while i < len(lines):
    raw  = lines[i].rstrip('\n')
    line = raw.strip()

    # ── Code block toggle
    if line.startswith('```'):
        if not in_code:
            in_code  = True
            code_buf = []
        else:
            in_code = False
            if code_buf:
                code_text = '\n'.join(code_buf)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent  = Cm(0.8)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(4)
                run = p.add_run(code_text)
                run.font.name      = 'Courier New'
                run.font.size      = Pt(8)
                run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  'F3F3F3')
                pPr.append(shd)
            code_buf = []
        i += 1
        continue

    if in_code:
        code_buf.append(raw)
        i += 1
        continue

    # ── Table row accumulation
    if line.startswith('|'):
        tbl_rows.append(line)
        i += 1
        continue
    else:
        if tbl_rows:
            # Flush table to document
            data_rows = [r for r in tbl_rows
                         if not re.match(r'^\|[-| :]+\|$', r.strip())]
            if data_rows:
                cols = [c.strip() for c in data_rows[0].split('|') if c.strip()]
                ncols = len(cols)
                if ncols > 0:
                    tbl = doc.add_table(rows=1, cols=ncols)
                    tbl.style = 'Table Grid'
                    # Header
                    hdr = tbl.rows[0].cells
                    for ci, cv in enumerate(cols):
                        hdr[ci].text = cv
                        run = hdr[ci].paragraphs[0].runs
                        if run:
                            run[0].font.bold      = True
                            run[0].font.color.rgb = C_WHITE
                            run[0].font.size      = Pt(9)
                        set_cell_bg(hdr[ci], '1F3564')
                    # Data rows
                    for dr in data_rows[1:]:
                        parts = dr.split('|')
                        if parts and parts[0].strip() == '':
                            parts = parts[1:]
                        if parts and parts[-1].strip() == '':
                            parts = parts[:-1]
                        row = tbl.add_row().cells
                        for ci in range(min(len(parts), ncols)):
                            p2 = row[ci].paragraphs[0]
                            apply_inline(p2, parts[ci].strip())
                            for r2 in p2.runs:
                                r2.font.size = Pt(9)
                    doc.add_paragraph()
            tbl_rows = []

    # ── Horizontal rule
    if re.match(r'^-{3,}$', line):
        add_hr(doc)
        i += 1
        continue

    # ── Headings
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        level   = len(m.group(1))
        heading = re.sub(r'\*\*([^*]+)\*\*', r'\1', m.group(2))
        heading = re.sub(r'`([^`]+)`', r'\1', heading)
        doc.add_heading(heading, level=level)
        i += 1
        continue

    # ── Blockquote
    if line.startswith('> '):
        text = line[2:]
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1.0)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        pPr = p._p.get_or_add_pPr()
        pb  = OxmlElement('w:pBdr')
        lft = OxmlElement('w:left')
        lft.set(qn('w:val'),   'single')
        lft.set(qn('w:sz'),    '12')
        lft.set(qn('w:space'), '8')
        lft.set(qn('w:color'), '2E74B5')
        pb.append(lft)
        pPr.append(pb)
        apply_inline(p, text)
        for r2 in p.runs:
            r2.font.size   = Pt(10)
            r2.font.italic = True
        i += 1
        continue

    # ── Bullet list
    m2 = re.match(r'^(\s*)[-*]\s+(.*)', raw)
    if m2:
        indent = len(m2.group(1)) // 2
        text   = m2.group(2)
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent  = Cm(0.6 + indent * 0.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        apply_inline(p, text)
        for r2 in p.runs:
            r2.font.size = Pt(10)
        i += 1
        continue

    # ── Numbered list
    m3 = re.match(r'^\s*\d+\.\s+(.*)', raw)
    if m3:
        text = m3.group(1)
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        apply_inline(p, text)
        for r2 in p.runs:
            r2.font.size = Pt(10)
        i += 1
        continue

    # ── Empty line
    if not line:
        i += 1
        continue

    # ── Normal paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    apply_inline(p, line)
    for r2 in p.runs:
        r2.font.size = Pt(10)
    i += 1

# ── Save
out = r'c:\Users\user\Downloads\amis_phase1__data_flow_crises_output\amis_phase_final\AMIS_RealWorld_Blueprint.docx'
doc.save(out)
print(f"Saved: {out}")
