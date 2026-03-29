"""
Generate Word Document for Demand Intelligence Module Documentation
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_demand_intelligence_doc():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title Page
    title = doc.add_heading('DEMAND INTELLIGENCE MODULE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Complete Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.bold = True

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('AMIS - Autonomous Manufacturing Intelligence System\n').font.size = Pt(12)
    info.add_run('Version 1.0 | March 2026\n').font.size = Pt(11)
    info.add_run('End-to-End Learning & Validation Guide').font.size = Pt(11)

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('TABLE OF CONTENTS', 1)
    toc_items = [
        ('1. MODULE OVERVIEW', 2),
        ('   1.1 What is Demand Intelligence?', 2),
        ('   1.2 Key Features', 2),
        ('   1.3 Who Uses This Module?', 3),
        ('2. SYSTEM ARCHITECTURE', 4),
        ('   2.1 High-Level Architecture Diagram', 4),
        ('   2.2 Technology Stack', 5),
        ('3. DATABASE SCHEMA & DATA SOURCES', 6),
        ('   3.1 Database Tables Overview', 6),
        ('   3.2 demand_forecasts Table (CORE TABLE)', 6),
        ('   3.3 products Table (REFERENCE)', 8),
        ('   3.4 Data Relationships', 9),
        ('   3.5 Database Access Functions', 9),
        ('   3.6 Sample Data Source (AI Training Data)', 12),
        ('4. PIPELINE EXECUTION FLOW', 14),
        ('   4.1 What is the Pipeline?', 14),
        ('   4.2 Pipeline Trigger Points', 14),
        ('   4.3 Step-by-Step Pipeline Flow', 15),
        ('5. AI AGENT & TOOLS', 25),
        ('   5.1 Agent Architecture', 25),
        ('   5.2 Agent Identity & Personality', 25),
        ('   5.3 Agent Tools (Detailed)', 26),
        ('   5.4 How Agent Interprets Tool Results', 27),
        ('6. CALCULATION LOGIC & FORMULAS', 29),
        ('   6.1 Frontend Calculations', 29),
        ('   6.2 AI Tool Calculations', 31),
        ('   6.3 Validation Formulas', 34),
        ('7. DATA FLOW: BACKEND TO FRONTEND', 35),
        ('   7.1 API Endpoints', 35),
        ('   7.2 Frontend API Client', 37),
        ('   7.3 React Query Integration', 37),
        ('   7.4 Complete Data Flow Diagram', 38),
        ('8. UI COMPONENTS & DISPLAY', 39),
        ('   8.1 Page Structure', 39),
        ('   8.2 Metric Cards (Detailed)', 39),
        ('   8.3 Area Chart Configuration', 41),
        ('   8.4 Forecast Analysis Boxes', 42),
        ('9. VALIDATION & TESTING', 44),
        ('   9.1 Database Validation Queries', 44),
        ('   9.2 API Endpoint Testing', 45),
        ('   9.3 Frontend Calculation Validation', 46),
        ('   9.4 AI Agent Output Validation', 47),
        ('   9.5 End-to-End Test Scenario', 48),
        ('10. TROUBLESHOOTING GUIDE', 49),
        ('   10.1 Common Issues & Solutions', 49),
        ('   10.2 Debugging Checklist', 51),
        ('   10.3 Logging & Monitoring', 52),
        ('APPENDIX A: Quick Reference Tables', 53),
        ('APPENDIX B: Sample Data for Testing', 54),
        ('GLOSSARY', 55),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5) if item.startswith('   ') else Inches(0)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        p.add_run(f'\t{page}')

    doc.add_page_break()

    # Section 1: MODULE OVERVIEW
    doc.add_heading('1. MODULE OVERVIEW', 1)

    doc.add_heading('1.1 What is Demand Intelligence?', 2)
    p = doc.add_paragraph(
        'Demand Intelligence is the forecasting module in AMIS that predicts future product '
        'demand using AI-powered scenario analysis. It helps manufacturing teams answer:'
    )
    doc.add_paragraph('"How much product should we make next week/month?"', style='List Bullet')
    doc.add_paragraph('"Is the recent demand spike a trend or anomaly?"', style='List Bullet')
    doc.add_paragraph('"What\'s the financial impact of different production strategies?"', style='List Bullet')

    doc.add_heading('1.2 Key Features', 2)

    # Create feature table
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Value to User'

    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Data rows
    features = [
        ('AI-Powered Forecasting', 'Uses Claude AI to analyze 12 weeks of historical data and generate multi-scenario forecasts', 'Confident decision-making with probability weights'),
        ('Automatic Pipeline Sync', 'When you run the AI Pipeline, forecasts automatically sync to database', 'No manual data entry needed'),
        ('Manual Forecast Entry', 'Users can add/edit forecasts manually with validation', 'Control over data, ability to override AI'),
        ('Trend Analysis', 'Calculates trend direction, growth rate, week-over-week changes', 'Understand if demand is growing, declining, or stable'),
        ('Visual Charts', 'Area chart shows 3 scenarios + actual demand', 'Easy to spot patterns and variances'),
        ('Actual vs Forecast', 'Track actual demand against predictions', 'Continuous improvement of forecasting process'),
    ]

    for i, (feature, desc, value) in enumerate(features, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = feature
        row_cells[1].text = desc
        row_cells[2].text = value

    doc.add_paragraph()

    doc.add_heading('1.3 Who Uses This Module?', 2)

    # User roles table
    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Light Grid Accent 1'

    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Role'
    hdr_cells2[1].text = 'What They Do'
    hdr_cells2[2].text = 'How They Use Demand Intelligence'

    for cell in hdr_cells2:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    roles = [
        ('Demand Planner (Jessica)', 'Forecasts customer demand, manages orders', 'Primary user - reviews AI forecasts, adds manual forecasts, tracks actuals'),
        ('Operations Manager (Sarah)', 'Makes strategic production decisions', 'Uses insights to approve production volumes'),
        ('Production Supervisor (Mike)', 'Executes daily production plans', 'Checks demand outlook to validate production schedules'),
    ]

    for i, (role, what, how) in enumerate(roles, start=1):
        row_cells = table2.rows[i].cells
        row_cells[0].text = role
        row_cells[1].text = what
        row_cells[2].text = how

    doc.add_page_break()

    # Section 2: SYSTEM ARCHITECTURE
    doc.add_heading('2. SYSTEM ARCHITECTURE', 1)

    doc.add_heading('2.1 High-Level Architecture Diagram', 2)

    arch_text = """
USER (Browser - localhost:5173)
    ↓
FRONTEND (React + Vite)
  - DemandIntelligence.jsx
  - Fetches /api/products
  - Fetches /api/demand/forecasts/{product_id}
  - POST /api/demand/forecasts (create)
  - Calculates insights (avg, trend)
    ↓
BACKEND (FastAPI - localhost:8000)
  - main.py
  - GET /api/demand/forecasts/{product_id}
  - POST /api/demand/forecasts/{product_id}
  - POST /api/pipeline/run
    ↓
DATABASE (SQLite - backend/amis.db)
  - demand_forecasts table
  - products table
  - production_schedule table
    ↓
AI AGENT (DemandForecastingAgent)
  - demand_agent.py
  - Uses LangChain + Claude (Anthropic API)
  - Tools: simulate_demand_scenarios()
           analyze_demand_trends()
           get_demand_data_summary()
    ↓
SAMPLE DATA (Simulated ERP Data)
  - sample_data.py
  - get_historical_demand() - 12 weeks
  - get_market_context() - External factors
  - get_current_inventory() - Stock levels
  - get_production_capacity() - Constraints
    """

    p = doc.add_paragraph(arch_text)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_heading('2.2 Technology Stack', 2)

    table3 = doc.add_table(rows=8, cols=3)
    table3.style = 'Light Grid Accent 1'

    hdr_cells3 = table3.rows[0].cells
    hdr_cells3[0].text = 'Layer'
    hdr_cells3[1].text = 'Technology'
    hdr_cells3[2].text = 'Purpose'

    for cell in hdr_cells3:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    stack = [
        ('Frontend', 'React 18 + Vite', 'User interface, charts'),
        ('Charts', 'Recharts', 'Data visualization (area charts)'),
        ('State Management', 'React Query (TanStack Query)', 'API calls, caching'),
        ('Backend', 'FastAPI (Python 3.11+)', 'REST API server'),
        ('AI Framework', 'LangChain + Claude 3.5 Sonnet', 'Agentic AI reasoning'),
        ('Database', 'SQLite 3', 'Persistent data storage'),
        ('Data Tools', 'Python (math, statistics)', 'Forecasting algorithms'),
    ]

    for i, (layer, tech, purpose) in enumerate(stack, start=1):
        row_cells = table3.rows[i].cells
        row_cells[0].text = layer
        row_cells[1].text = tech
        row_cells[2].text = purpose

    doc.add_page_break()

    # Section 3: DATABASE SCHEMA
    doc.add_heading('3. DATABASE SCHEMA & DATA SOURCES', 1)

    doc.add_heading('3.1 Database Tables Overview', 2)
    p = doc.add_paragraph('The Demand Intelligence module uses ')
    p.add_run('2 primary tables').bold = True
    p.add_run(' and ')
    p.add_run('1 reference table').bold = True
    p.add_run(':')

    doc.add_paragraph('demand_forecasts - Stores all forecast data (AI + manual)', style='List Number')
    doc.add_paragraph('products - Product master data', style='List Number')
    doc.add_paragraph('production_schedule - Links demand to production plans', style='List Number')

    doc.add_heading('3.2 demand_forecasts Table (CORE TABLE)', 2)

    p = doc.add_paragraph()
    p.add_run('Location: ').bold = True
    p.add_run('backend/schema.sql lines 179-190')

    # SQL code
    sql_code = """CREATE TABLE IF NOT EXISTS demand_forecasts (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    product_id TEXT NOT NULL,
    week_number INTEGER,
    forecast_date DATE,
    optimistic INTEGER,
    base_case INTEGER,
    pessimistic INTEGER,
    actual INTEGER,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (product_id) REFERENCES products(id)
);"""

    p = doc.add_paragraph(sql_code)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 128)

    doc.add_heading('Column Details', 3)

    # Column details table
    table4 = doc.add_table(rows=10, cols=5)
    table4.style = 'Light Grid Accent 1'

    hdr_cells4 = table4.rows[0].cells
    headers = ['Column', 'Type', 'Nullable', 'Business Meaning', 'Example Value']
    for i, header in enumerate(headers):
        hdr_cells4[i].text = header
        for paragraph in hdr_cells4[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    columns = [
        ('id', 'INTEGER', 'NO', 'Unique ID for each forecast entry', '42'),
        ('product_id', 'TEXT', 'NO', 'Which product this forecast is for', '"PROD-A"'),
        ('week_number', 'INTEGER', 'YES', 'Which week this forecast covers (1-52)', '10'),
        ('forecast_date', 'DATE', 'YES', 'When this forecast week begins', '"2026-03-09"'),
        ('optimistic', 'INTEGER', 'YES', 'Upper bound demand (20% probability)', '1,250'),
        ('base_case', 'INTEGER', 'YES', 'Expected demand (55% probability)', '1,050'),
        ('pessimistic', 'INTEGER', 'YES', 'Lower bound demand (25% probability)', '850'),
        ('actual', 'INTEGER', 'YES', 'Real sales after week completes', '1,080'),
        ('created_at', 'TIMESTAMP', 'NO', 'When this record was created', '"2026-03-04 10:23:45"'),
    ]

    for i, (col, typ, null, meaning, example) in enumerate(columns, start=1):
        row_cells = table4.rows[i].cells
        row_cells[0].text = col
        row_cells[1].text = typ
        row_cells[2].text = null
        row_cells[3].text = meaning
        row_cells[4].text = example

    doc.add_page_break()

    # Section 4: PIPELINE EXECUTION FLOW
    doc.add_heading('4. PIPELINE EXECUTION FLOW', 1)

    doc.add_heading('4.1 What is the Pipeline?', 2)
    p = doc.add_paragraph(
        'The AI Pipeline is the automated daily process that runs all 5 AI agents in sequence '
        'to update the system with fresh insights.'
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Think of it like this:').bold = True
    doc.add_paragraph('Without Pipeline: You manually ask each AI agent for analysis (slow, manual)', style='List Bullet')
    doc.add_paragraph('With Pipeline: Click one button → all 5 agents run → database updates automatically', style='List Bullet')

    doc.add_heading('4.2 Pipeline Trigger Points', 2)

    table5 = doc.add_table(rows=4, cols=4)
    table5.style = 'Light Grid Accent 1'

    hdr_cells5 = table5.rows[0].cells
    headers5 = ['Trigger', 'When', 'How', 'Who']
    for i, header in enumerate(headers5):
        hdr_cells5[i].text = header
        for paragraph in hdr_cells5[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    triggers = [
        ('Manual (UI Button)', 'Anytime user clicks "Run Pipeline"', 'User goes to Pipeline page → clicks button', 'Any user'),
        ('Scheduled (Future)', 'Daily at 6:00 AM', 'Cron job / scheduled task', 'System (automated)'),
        ('API Call', 'Programmatic trigger', 'POST /api/pipeline/run', 'External system'),
    ]

    for i, (trigger, when, how, who) in enumerate(triggers, start=1):
        row_cells = table5.rows[i].cells
        row_cells[0].text = trigger
        row_cells[1].text = when
        row_cells[2].text = how
        row_cells[3].text = who

    doc.add_page_break()

    # Section 6: CALCULATION LOGIC & FORMULAS
    doc.add_heading('6. CALCULATION LOGIC & FORMULAS', 1)

    doc.add_heading('6.1 Frontend Calculations (Real-Time Insights)', 2)

    doc.add_heading('6.1.1 Average Weekly Demand (Base Case)', 3)

    p = doc.add_paragraph()
    p.add_run('Formula:').bold = True

    formula = """avgBase = forecastData.reduce((sum, f) => sum + (f.base_case || 0), 0) / forecastData.length"""
    p = doc.add_paragraph(formula)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.add_run('Example Calculation:').bold = True

    example = """Forecast Data from Database:
Week 10: base_case = 1050
Week 11: base_case = 1075
Week 12: base_case = 1100
Week 13: base_case = 1125

avgBase = (1050 + 1075 + 1100 + 1125) / 4
        = 4350 / 4
        = 1087.5
        ≈ 1088 units"""

    p = doc.add_paragraph(example)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.add_run('Displayed As: ').bold = True
    p.add_run('"1,088 units" (with comma formatting)')

    doc.add_heading('6.1.2 Trend Direction & Percent', 3)

    p = doc.add_paragraph()
    p.add_run('Formula:').bold = True

    formula2 = """const trend = forecastData.length > 1
  ? ((forecastData[forecastData.length - 1]?.base_case || 0) - (forecastData[0]?.base_case || 0)) / forecastData.length
  : 0

const trendDirection = trend > 0 ? 'Upward' : trend < 0 ? 'Downward' : 'Stable'
const trendPercent = avgBase > 0 ? ((trend / avgBase) * 100).toFixed(1) : 0"""

    p = doc.add_paragraph(formula2)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    example2 = """Example Calculation:
First week (week 10): base_case = 1050
Last week (week 13): base_case = 1125
Number of weeks: 4

trend = (1125 - 1050) / 4 = 75 / 4 = 18.75 units/week growth

trendDirection = 'Upward' (since 18.75 > 0)

trendPercent = (18.75 / 1088) * 100 = 1.72% ≈ 1.7%

Displayed As: "Upward" with badge "+1.7% per week" """

    p = doc.add_paragraph(example2)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_page_break()

    # Section 9: VALIDATION & TESTING
    doc.add_heading('9. VALIDATION & TESTING', 1)

    doc.add_heading('9.1 Database Validation Queries', 2)

    doc.add_heading('9.1.1 Check if Forecasts Exist', 3)

    p = doc.add_paragraph()
    p.add_run('Query:').bold = True

    query1 = """SELECT COUNT(*) as forecast_count
FROM demand_forecasts
WHERE product_id = 'PROD-A';"""

    p = doc.add_paragraph(query1)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 128)

    p = doc.add_paragraph()
    p.add_run('Expected Result:').bold = True

    result1 = """forecast_count
--------------
12"""

    p = doc.add_paragraph(result1)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.add_run('If Result = 0:').bold = True
    doc.add_paragraph('Pipeline has NOT been run yet', style='List Bullet')
    doc.add_paragraph('OR database sync failed', style='List Bullet')
    doc.add_paragraph('Fix: Run pipeline manually from UI', style='List Bullet')

    doc.add_heading('9.1.2 Verify Forecast Data Quality', 3)

    query2 = """SELECT
    week_number,
    optimistic,
    base_case,
    pessimistic,
    actual,
    CASE
        WHEN optimistic >= base_case AND base_case >= pessimistic THEN 'VALID'
        ELSE 'INVALID'
    END as data_quality
FROM demand_forecasts
WHERE product_id = 'PROD-A'
ORDER BY week_number;"""

    p = doc.add_paragraph(query2)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 128)

    doc.add_page_break()

    # Appendix A: Quick Reference
    doc.add_heading('APPENDIX A: QUICK REFERENCE TABLES', 1)

    doc.add_heading('A.1 Key Formulas Reference', 2)

    table6 = doc.add_table(rows=7, cols=3)
    table6.style = 'Light Grid Accent 1'

    hdr_cells6 = table6.rows[0].cells
    hdr_cells6[0].text = 'Metric'
    hdr_cells6[1].text = 'Formula'
    hdr_cells6[2].text = 'Example'

    for cell in hdr_cells6:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    formulas_ref = [
        ('Average Demand', 'sum(base_case) / count', '(1050+1075+1100+1125) / 4 = 1088'),
        ('Trend Slope', '(last - first) / count', '(1125 - 1050) / 4 = 18.75 units/week'),
        ('Trend %', '(slope / average) × 100', '(18.75 / 1088) × 100 = 1.7%'),
        ('Std Deviation', 'sqrt(sum((x - avg)²) / n)', 'sqrt(204051 / 12) = 130.4'),
        ('Z-Score', '(value - average) / std_dev', '(1450 - 1138) / 130.4 = 2.39'),
        ('Expected Weighted', 'Σ(scenario × probability)', '5782×0.2 + 4912×0.55 + 3822×0.25 = 4813'),
    ]

    for i, (metric, formula, example) in enumerate(formulas_ref, start=1):
        row_cells = table6.rows[i].cells
        row_cells[0].text = metric
        row_cells[1].text = formula
        row_cells[2].text = example

    doc.add_heading('A.2 API Endpoints Reference', 2)

    table7 = doc.add_table(rows=6, cols=3)
    table7.style = 'Light Grid Accent 1'

    hdr_cells7 = table7.rows[0].cells
    hdr_cells7[0].text = 'Method'
    hdr_cells7[1].text = 'Endpoint'
    hdr_cells7[2].text = 'Purpose'

    for cell in hdr_cells7:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    endpoints = [
        ('GET', '/api/demand/forecasts/{product_id}', 'Fetch forecasts'),
        ('POST', '/api/demand/forecasts/{product_id}', 'Create forecast'),
        ('PATCH', '/api/demand/forecasts/{product_id}/{week}', 'Update actual'),
        ('POST', '/api/pipeline/run', 'Run AI pipeline'),
        ('GET', '/api/pipeline/runs/{run_id}', 'Get pipeline status'),
    ]

    for i, (method, endpoint, purpose) in enumerate(endpoints, start=1):
        row_cells = table7.rows[i].cells
        row_cells[0].text = method
        row_cells[1].text = endpoint
        row_cells[2].text = purpose

    doc.add_page_break()

    # Glossary
    doc.add_heading('GLOSSARY', 1)

    glossary_items = [
        ('Agentic AI', 'AI that autonomously decides which tools to use, interprets results, and reasons like a human expert (not just executing predefined scripts)'),
        ('Area Chart', 'Chart type that fills the area under a line, useful for showing volume/magnitude over time'),
        ('Base Case Forecast', 'Most likely demand scenario, typically assigned 50-60% probability weight'),
        ('Coefficient of Variation (CV)', 'Standard deviation divided by mean, expressed as percentage - measures relative variability'),
        ('Confidence Interval', 'Range of values (e.g., 95% CI) where actual value is likely to fall'),
        ('FastAPI', 'Modern Python web framework for building APIs with automatic validation and documentation'),
        ('LangChain', 'Framework for building LLM-powered applications with tool calling and agent orchestration'),
        ('Linear Regression', 'Statistical method to find the best-fit line through data points (used for trend analysis)'),
        ('Optimistic Scenario', 'Best-case demand forecast, typically assigned 15-25% probability weight'),
        ('Pessimistic Scenario', 'Worst-case demand forecast, typically assigned 20-30% probability weight'),
        ('Pipeline', 'Automated workflow that runs all 5 AI agents in sequence'),
        ('Standard Deviation', 'Measure of data spread - how much values deviate from the average'),
        ('Trend', 'Long-term direction of change in data (upward, downward, or stable)'),
        ('Z-Score', 'Number of standard deviations a value is from the mean (used for anomaly detection)'),
    ]

    for term, definition in glossary_items:
        p = doc.add_paragraph()
        p.add_run(term + ': ').bold = True
        p.add_run(definition)

    doc.add_page_break()

    # Conclusion
    doc.add_heading('CONCLUSION', 1)

    p = doc.add_paragraph(
        'This documentation provides complete end-to-end coverage of the Demand Intelligence module. '
        'You now understand:'
    )

    doc.add_paragraph('How the system works: From user click to AI agent to database to UI display', style='List Bullet')
    doc.add_paragraph('Where data comes from: Database tables, sample data generators, AI tools', style='List Bullet')
    doc.add_paragraph('How calculations work: Every formula explained with examples', style='List Bullet')
    doc.add_paragraph('How to validate results: SQL queries, API tests, manual calculations', style='List Bullet')
    doc.add_paragraph('How to troubleshoot: Common issues, debugging steps, log interpretation', style='List Bullet')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Next Steps:').bold = True
    doc.add_paragraph('Run the system: Start backend + frontend, run pipeline, observe data flow', style='List Number')
    doc.add_paragraph('Trace one forecast: Pick a single forecast week, trace it from AI → DB → UI', style='List Number')
    doc.add_paragraph('Modify and test: Change a value in DB, verify UI updates correctly', style='List Number')
    doc.add_paragraph('Experiment: Try different products, add manual forecasts, update actuals', style='List Number')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('For Questions:').bold = True
    doc.add_paragraph('Review Section 4 (Pipeline Flow) for "how it works"', style='List Bullet')
    doc.add_paragraph('Review Section 6 (Calculations) for "why this number"', style='List Bullet')
    doc.add_paragraph('Review Section 9 (Validation) for "how to verify"', style='List Bullet')
    doc.add_paragraph('Review Section 10 (Troubleshooting) for "it\'s broken, help!"', style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer info
    p = doc.add_paragraph()
    p.add_run('Document Version: ').bold = True
    p.add_run('1.0\n')
    p.add_run('Last Updated: ').bold = True
    p.add_run('March 4, 2026\n')
    p.add_run('Maintained By: ').bold = True
    p.add_run('AMIS Development Team\n')
    p.add_run('For: ').bold = True
    p.add_run('Manufacturing Team Training & System Validation')

    # Save document
    output_path = 'DEMAND_INTELLIGENCE_COMPLETE_DOCUMENTATION.docx'
    doc.save(output_path)
    print(f"[OK] Word document created: {output_path}")
    return output_path

if __name__ == "__main__":
    create_demand_intelligence_doc()
