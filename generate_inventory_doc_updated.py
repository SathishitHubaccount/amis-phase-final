"""
Generate comprehensive Inventory Management Complete Guide as Word document.
This script creates detailed documentation of all 6 inventory management tools.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_heading(doc, text, level=1):
    """Add a formatted heading."""
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_paragraph(doc, text, bold=False):
    """Add a formatted paragraph."""
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    return p

def add_table_from_data(doc, headers, rows):
    """Create a formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

def create_inventory_guide():
    """Create the complete Inventory Management guide."""

    doc = Document()

    # Title
    title = doc.add_heading('AMIS Inventory Management Agent', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Complete Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # Section 1: Overview
    add_heading(doc, '1. Agent Overview', 1)
    add_paragraph(doc, 'The Inventory Management Agent is responsible for maintaining optimal inventory levels, preventing stockouts, and minimizing holding costs. It uses advanced statistical methods and Monte Carlo simulations to provide data-driven replenishment recommendations.')

    add_paragraph(doc, 'Key Capabilities:', bold=True)
    capabilities = [
        'Real-time inventory status monitoring across multiple warehouse zones',
        'Reorder point calculations using statistical safety stock formulas',
        'Safety stock optimization balancing service levels and costs',
        'Monte Carlo simulations for stockout risk assessment',
        'Economic Order Quantity (EOQ) analysis for holding cost optimization',
        'Multi-supplier replenishment planning with optimal allocation'
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')

    # Section 2: Database Tables
    add_heading(doc, '2. Database Tables Used', 1)
    add_paragraph(doc, 'The agent queries 12 database tables to gather comprehensive inventory intelligence:')

    table_data = [
        ['products', 'Product master data (ID, name, category, status)'],
        ['inventory', 'Current stock levels, safety stock, reorder points, stockout risk'],
        ['warehouse_zones', 'Zone capacity, utilization, temperature control status'],
        ['stockout_events', 'Historical stockout incidents with root cause analysis'],
        ['purchase_orders', 'Open and historical POs with delivery tracking'],
        ['suppliers', 'Supplier ratings, lead times, reliability scores'],
        ['supplier_contracts', 'Contract terms, volumes, validity periods'],
        ['bom_items', 'Bill of materials with component costs and stock levels'],
        ['historical_demand_data', 'Weekly demand history for trend analysis'],
        ['market_context_data', 'Market conditions, seasonality, economic indicators'],
        ['production_schedule', 'Planned vs actual production, capacity gaps'],
        ['activity_log', 'Audit trail of all inventory transactions']
    ]

    add_table_from_data(doc, ['Table Name', 'Purpose'], table_data)

    # Section 3: Tool 1 - get_inventory_status
    add_heading(doc, '3. Tool #1: get_inventory_status()', 1)
    add_paragraph(doc, 'Purpose: Provides a comprehensive snapshot of current inventory health.')

    add_heading(doc, '3.1 Data Sources', 2)
    data_sources = [
        ['get_current_inventory()', 'Current stock, safety stock, avg daily usage, stockout risk'],
        ['get_warehouse_details()', 'Zone capacities, utilization rates, bin locations'],
        ['get_open_purchase_orders()', 'Incoming shipments with expected delivery dates'],
        ['get_supplier_performance()', 'Supplier reliability, on-time delivery rates'],
        ['get_historical_stockouts()', 'Past stockout events, duration, revenue impact'],
        ['get_product_info()', 'Product category, status, active/discontinued flag']
    ]
    add_table_from_data(doc, ['Database Function', 'Data Retrieved'], data_sources)

    add_heading(doc, '3.2 Health Indicators', 2)
    indicators = [
        ['Days of Supply', 'current_stock / avg_daily_usage', '13.0 days'],
        ['Stockout Risk', 'Monte Carlo simulation result', '12.3%'],
        ['Stock Status', 'Based on days supply thresholds', 'Adequate/Low/Critical'],
        ['Warehouse Utilization', 'Sum of zone utilization percentages', '36.7%'],
        ['Incoming Units', 'Sum of open PO quantities', '5,000 units'],
        ['Supplier Reliability', 'Average on-time delivery %', '92.5%'],
        ['Recent Stockouts', 'Count in last 90 days', '2 events'],
        ['Product Status', 'Active, discontinued, or phase-out', 'Active']
    ]
    add_table_from_data(doc, ['Indicator', 'Calculation', 'Example Value'], indicators)

    add_heading(doc, '3.3 SQL Validation Query', 2)
    sql1 = doc.add_paragraph()
    sql1.add_run('SELECT\n    i.product_id,\n    i.current_stock,\n    i.safety_stock,\n    i.avg_daily_usage,\n    i.stockout_risk,\n    p.name,\n    p.status\nFROM inventory i\nJOIN products p ON i.product_id = p.id\nWHERE i.product_id = ?').font.name = 'Courier New'

    # Section 4: Tool 2 - calculate_reorder_point
    add_heading(doc, '4. Tool #2: calculate_reorder_point()', 1)
    add_paragraph(doc, 'Purpose: Calculates the optimal reorder point using statistical safety stock methodology.')

    add_heading(doc, '4.1 Formula', 2)
    add_paragraph(doc, 'Reorder Point (ROP) = (Average Daily Demand × Lead Time) + Safety Stock')
    doc.add_paragraph()
    add_paragraph(doc, 'Safety Stock = Z-score × √(LT × σ_demand² + Demand² × σ_LT²)')
    doc.add_paragraph()
    add_paragraph(doc, 'Where:')
    variables = [
        'Z-score: Statistical multiplier based on desired service level',
        'LT: Lead time in days',
        'σ_demand: Standard deviation of daily demand',
        'σ_LT: Standard deviation of lead time',
        'Demand: Average daily demand'
    ]
    for var in variables:
        doc.add_paragraph(var, style='List Bullet')

    add_heading(doc, '4.2 Z-Score Lookup Table', 2)
    zscores = [
        ['85%', '1.04', 'Lower inventory, higher stockout risk'],
        ['90%', '1.28', 'Balanced approach'],
        ['95%', '1.65', 'Standard target for most products'],
        ['97%', '1.88', 'High-value items'],
        ['99%', '2.33', 'Critical items, expensive stockouts']
    ]
    add_table_from_data(doc, ['Service Level', 'Z-Score', 'Use Case'], zscores)

    add_heading(doc, '4.3 Example Calculation', 2)
    add_paragraph(doc, 'Given:')
    inputs = [
        'Average daily demand: 142 units',
        'Demand std dev: 25 units',
        'Lead time: 5 days',
        'Lead time std dev: 1 day',
        'Service level: 95%'
    ]
    for inp in inputs:
        doc.add_paragraph(inp, style='List Bullet')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 1: Get Z-score for 95% service level = 1.65')
    add_paragraph(doc, 'Step 2: Calculate safety stock')
    add_paragraph(doc, '  = 1.65 × √(5 × 25² + 142² × 1²)')
    add_paragraph(doc, '  = 1.65 × √(3,125 + 20,164)')
    add_paragraph(doc, '  = 1.65 × √23,289')
    add_paragraph(doc, '  = 1.65 × 152.6')
    add_paragraph(doc, '  = 252 units')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 3: Calculate ROP')
    add_paragraph(doc, '  = (142 × 5) + 252')
    add_paragraph(doc, '  = 710 + 252')
    add_paragraph(doc, '  = 962 units')

    # Section 5: Tool 3 - optimize_safety_stock
    add_heading(doc, '5. Tool #3: optimize_safety_stock()', 1)
    add_paragraph(doc, 'Purpose: Evaluates multiple service levels to find the optimal balance between stockout costs and holding costs.')

    add_heading(doc, '5.1 Service Levels Tested', 2)
    add_paragraph(doc, 'The tool tests 7 service levels: 85%, 90%, 95%, 97%, 99%, 99.5%, 99.9%')

    add_heading(doc, '5.2 Cost Calculations', 2)
    cost_calcs = [
        ['Safety Stock', 'Z × √(LT × σ_demand² + Demand² × σ_LT²)', 'Units held as buffer'],
        ['Holding Cost', 'Safety Stock × Unit Cost × Holding Rate', '$/year to hold safety stock'],
        ['Expected Stockouts', '(1 - Service Level) × Demand × Days', 'Units short per year'],
        ['Stockout Cost', 'Expected Stockouts × Stockout Cost/Unit', '$/year from stockouts'],
        ['Total Cost', 'Holding Cost + Stockout Cost', 'Total annual cost']
    ]
    add_table_from_data(doc, ['Cost Component', 'Formula', 'Description'], cost_calcs)

    add_heading(doc, '5.3 Example Results Table', 2)
    example_results = [
        ['85%', '1.04', '158', '$790', '$12,775', '$13,565'],
        ['90%', '1.28', '195', '$975', '$8,517', '$9,492'],
        ['95%', '1.65', '252', '$1,260', '$4,258', '$5,518'],
        ['97%', '1.88', '287', '$1,435', '$2,555', '$3,990'],
        ['99%', '2.33', '356', '$1,780', '$851', '$2,631'],
        ['99.5%', '2.58', '394', '$1,970', '$426', '$2,396'],
        ['99.9%', '3.09', '$472', '$2,360', '$85', '$2,445']
    ]
    add_table_from_data(doc, ['Service Level', 'Z-Score', 'Safety Stock', 'Holding Cost', 'Stockout Cost', 'Total Cost'], example_results)

    add_paragraph(doc, 'Recommendation: 99.5% service level provides lowest total cost at $2,396/year')

    # Section 6: Tool 4 - simulate_stockout_risk
    add_heading(doc, '6. Tool #4: simulate_stockout_risk()', 1)
    add_paragraph(doc, 'Purpose: Uses Monte Carlo simulation to estimate stockout probability over a planning horizon.')

    add_heading(doc, '6.1 Simulation Parameters', 2)
    sim_params = [
        ['Number of Runs', '1,000 simulations', 'Statistical significance'],
        ['Planning Horizon', '14 days (default)', 'Configurable up to 30 days'],
        ['Demand Distribution', 'Normal(mean, std_dev)', 'Daily demand variability'],
        ['Lead Time Distribution', 'Normal(mean, std_dev)', 'Supplier delivery variability'],
        ['Initial Stock', 'current_stock', 'Starting inventory level'],
        ['Reorder Trigger', 'reorder_point', 'When to place new order']
    ]
    add_table_from_data(doc, ['Parameter', 'Value', 'Purpose'], sim_params)

    add_heading(doc, '6.2 Simulation Logic', 2)
    logic_steps = [
        'Day 1: Start with current_stock',
        'Each day: Generate random demand from Normal(avg_demand, std_dev)',
        'Subtract demand from stock',
        'If stock < reorder_point and no order pending: Place order',
        'When order arrives (after lead_time days): Add order_quantity to stock',
        'If stock < 0 at any point: Record stockout event',
        'Repeat for all days in horizon',
        'Run 1,000 times and calculate % of runs with stockouts'
    ]
    for step in logic_steps:
        doc.add_paragraph(step, style='List Number')

    add_heading(doc, '6.3 Output Metrics', 2)
    output_metrics = [
        ['Stockout Risk %', 'Percentage of simulations with at least one stockout', '12.3%'],
        ['Avg Stockout Days', 'Average days out of stock when stockouts occur', '2.1 days'],
        ['Worst Case', 'Maximum stock deficit observed across all runs', '-180 units'],
        ['Best Case', 'Minimum stock level (if no stockouts)', '320 units'],
        ['Avg Ending Stock', 'Mean stock level at end of horizon', '1,245 units']
    ]
    add_table_from_data(doc, ['Metric', 'Description', 'Example Value'], output_metrics)

    # Section 7: Tool 5 - evaluate_holding_costs
    add_heading(doc, '7. Tool #5: evaluate_holding_costs()', 1)
    add_paragraph(doc, 'Purpose: Analyzes inventory holding costs and calculates Economic Order Quantity (EOQ).')

    add_heading(doc, '7.1 EOQ Formula', 2)
    add_paragraph(doc, 'EOQ = √(2 × Annual Demand × Ordering Cost / Holding Cost per Unit)')

    add_heading(doc, '7.2 Example Calculation', 2)
    add_paragraph(doc, 'Given:')
    eoq_inputs = [
        'Annual demand: 51,830 units (142/day × 365 days)',
        'Ordering cost: $150 per order',
        'Unit cost: $5.00',
        'Holding cost rate: 20% per year',
        'Holding cost per unit: $1.00 ($5.00 × 0.20)'
    ]
    for inp in eoq_inputs:
        doc.add_paragraph(inp, style='List Bullet')

    doc.add_paragraph()
    add_paragraph(doc, 'EOQ = √(2 × 51,830 × $150 / $1.00)')
    add_paragraph(doc, '    = √(15,549,000 / 1.00)')
    add_paragraph(doc, '    = √15,549,000')
    add_paragraph(doc, '    = 3,943 units per order')

    add_heading(doc, '7.3 Cost Analysis', 2)
    cost_analysis = [
        ['Current Order Quantity', '2,000 units', 'Existing practice'],
        ['EOQ Recommendation', '3,943 units', 'Optimal order size'],
        ['Orders per Year (Current)', '26 orders', '51,830 / 2,000'],
        ['Orders per Year (EOQ)', '13 orders', '51,830 / 3,943'],
        ['Annual Ordering Cost (Current)', '$3,900', '26 × $150'],
        ['Annual Ordering Cost (EOQ)', '$1,950', '13 × $150'],
        ['Annual Holding Cost (Current)', '$1,000', 'Avg inventory × holding cost'],
        ['Annual Holding Cost (EOQ)', '$1,972', 'Avg inventory × holding cost'],
        ['Total Cost (Current)', '$4,900', 'Sum of ordering + holding'],
        ['Total Cost (EOQ)', '$3,922', 'Sum of ordering + holding'],
        ['Annual Savings', '$978', 'Difference']
    ]
    add_table_from_data(doc, ['Metric', 'Value', 'Calculation'], cost_analysis)

    # Section 8: Tool 6 - generate_replenishment_plan
    add_heading(doc, '8. Tool #6: generate_replenishment_plan()', 1)
    add_paragraph(doc, 'Purpose: Creates a multi-week replenishment schedule with optimal supplier allocation.')

    add_heading(doc, '8.1 Planning Parameters', 2)
    plan_params = [
        ['Planning Horizon', '4 weeks (default)', 'Can extend to 12 weeks'],
        ['Order Quantity', 'EOQ or custom', 'Based on evaluate_holding_costs()'],
        ['Supplier Split', '65% / 35%', 'Primary vs backup supplier'],
        ['Lead Times', 'Supplier-specific', 'From suppliers table'],
        ['Safety Stock Target', 'From optimize_safety_stock()', 'Maintains buffer']
    ]
    add_table_from_data(doc, ['Parameter', 'Value', 'Notes'], plan_params)

    add_heading(doc, '8.2 Allocation Logic', 2)
    allocation_steps = [
        'Calculate total order quantity (e.g., 4,000 units)',
        'Allocate 65% to primary supplier (Supplier A): 2,600 units',
        'Allocate 35% to backup supplier (Supplier B): 1,400 units',
        'Adjust for supplier MOQ (minimum order quantity)',
        'Schedule orders based on lead times and expected stockout dates',
        'Ensure safety stock is maintained throughout horizon'
    ]
    for step in allocation_steps:
        doc.add_paragraph(step, style='List Number')

    add_heading(doc, '8.3 Example 4-Week Plan', 2)
    plan_example = [
        ['Week 1', '1,850', '710', 'No', '-', '-', '-', '1,850'],
        ['Week 2', '856', '710', 'Yes', '2,600 (Supplier A)', '5 days', 'Week 3', '856'],
        ['Week 3', '3,312', '710', 'Yes', '1,400 (Supplier B)', '7 days', 'Week 4', '3,312'],
        ['Week 4', '4,430', '710', 'No', '-', '-', '-', '4,430']
    ]
    add_table_from_data(doc, ['Week', 'Projected Stock', 'Reorder Point', 'Order?', 'Quantity', 'Lead Time', 'Arrives', 'Ending Stock'], plan_example)

    # Section 9: Complete Execution Flow
    add_heading(doc, '9. Complete Agent Execution Flow', 1)
    add_paragraph(doc, 'When the Inventory Management Agent receives a query, it follows this workflow:')

    workflow = [
        'User Query: "What is the inventory status for Product A?"',
        'Agent invokes: get_inventory_status(product_id="PROD-A")',
        'Tool queries 6 database functions to gather data',
        'Tool calculates 8 health indicators',
        'Tool returns comprehensive status report',
        'If stock is low, agent automatically invokes: calculate_reorder_point()',
        'Agent evaluates: simulate_stockout_risk() to assess urgency',
        'If reorder needed, agent invokes: generate_replenishment_plan()',
        'Agent presents recommendations with supporting calculations',
        'User can request: optimize_safety_stock() for cost analysis',
        'Agent logs all actions to activity_log table for audit trail'
    ]
    for step in workflow:
        doc.add_paragraph(step, style='List Number')

    # Section 10: SQL Validation Queries
    add_heading(doc, '10. SQL Validation Queries', 1)
    add_paragraph(doc, 'Use these queries to validate agent calculations:')

    add_heading(doc, '10.1 Check Current Inventory Status', 2)
    q1 = doc.add_paragraph()
    q1.add_run('SELECT product_id, current_stock, safety_stock,\n       reorder_point, avg_daily_usage, stockout_risk\nFROM inventory\nWHERE product_id = \'PROD-A\';').font.name = 'Courier New'

    add_heading(doc, '10.2 Verify Warehouse Zone Utilization', 2)
    q2 = doc.add_paragraph()
    q2.add_run('SELECT zone_id, capacity, current_units,\n       ROUND(current_units * 100.0 / capacity, 1) AS utilization_pct\nFROM warehouse_zones\nWHERE product_id = \'PROD-A\'\nORDER BY utilization_pct DESC;').font.name = 'Courier New'

    add_heading(doc, '10.3 Check Open Purchase Orders', 2)
    q3 = doc.add_paragraph()
    q3.add_run('SELECT po_id, supplier_id, quantity,\n       order_date, expected_delivery_date,\n       status\nFROM purchase_orders\nWHERE product_id = \'PROD-A\'\n  AND status = \'open\'\nORDER BY expected_delivery_date;').font.name = 'Courier New'

    add_heading(doc, '10.4 Review Recent Stockout Events', 2)
    q4 = doc.add_paragraph()
    q4.add_run('SELECT event_date, duration_days, units_short,\n       root_cause, revenue_lost\nFROM stockout_events\nWHERE product_id = \'PROD-A\'\n  AND event_date >= date(\'now\', \'-90 days\')\nORDER BY event_date DESC;').font.name = 'Courier New'

    add_heading(doc, '10.5 Calculate Days of Supply', 2)
    q5 = doc.add_paragraph()
    q5.add_run('SELECT product_id,\n       current_stock,\n       avg_daily_usage,\n       ROUND(current_stock * 1.0 / avg_daily_usage, 1) AS days_supply\nFROM inventory\nWHERE product_id = \'PROD-A\';').font.name = 'Courier New'

    # Section 11: Formula Reference
    add_heading(doc, '11. Formula Reference Sheet', 1)

    formulas = [
        ['Reorder Point (ROP)', 'ROP = (Avg Daily Demand × Lead Time) + Safety Stock'],
        ['Safety Stock', 'SS = Z × √(LT × σ_demand² + Demand² × σ_LT²)'],
        ['Economic Order Quantity', 'EOQ = √(2 × Annual Demand × Ordering Cost / Holding Cost)'],
        ['Days of Supply', 'DOS = Current Stock / Average Daily Usage'],
        ['Stockout Risk', 'Monte Carlo: % of simulations with stock < 0'],
        ['Holding Cost', 'HC = Average Inventory × Unit Cost × Holding Rate'],
        ['Stockout Cost', 'SC = Expected Shortfall × Stockout Cost per Unit'],
        ['Total Inventory Cost', 'TIC = Holding Cost + Stockout Cost + Ordering Cost'],
        ['Warehouse Utilization', 'Util% = (Current Units / Capacity) × 100'],
        ['Service Level', 'SL% = 1 - (Expected Stockouts / Total Demand)']
    ]
    add_table_from_data(doc, ['Formula Name', 'Equation'], formulas)

    # Section 12: Summary
    add_heading(doc, '12. Summary', 1)
    add_paragraph(doc, 'The AMIS Inventory Management Agent provides comprehensive inventory optimization through:')

    summary_points = [
        'Real-time monitoring of 12 database tables with 281 data points',
        '6 specialized tools covering status, reorder points, safety stock, risk simulation, EOQ analysis, and replenishment planning',
        'Statistical methods including Z-scores, normal distributions, and Monte Carlo simulations',
        'Cost optimization balancing holding costs, ordering costs, and stockout costs',
        'Multi-supplier allocation with 65/35 split for supply chain resilience',
        'Complete audit trail in activity_log table for compliance',
        'Validation SQL queries for manual verification of calculations'
    ]
    for point in summary_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_paragraph()
    add_paragraph(doc, 'Technology Stack:', bold=True)
    tech_stack = [
        'LangChain Framework: Agent orchestration and tool management',
        'Claude 3.5 Sonnet: Natural language understanding and reasoning',
        'SQLite Database: Persistent data storage (amis.db)',
        'Python: Backend logic and calculations',
        'Statistical Libraries: numpy, scipy for advanced calculations'
    ]
    for tech in tech_stack:
        doc.add_paragraph(tech, style='List Bullet')

    # Save document
    output_path = os.path.join(os.getcwd(), 'INVENTORY_MANAGEMENT_COMPLETE_GUIDE_UPDATED.docx')
    doc.save(output_path)
    print(f"\n[SUCCESS] Document created: {output_path}")
    print(f"[INFO] Document contains 12 comprehensive sections covering all 6 tools")
    return output_path

if __name__ == "__main__":
    create_inventory_guide()
